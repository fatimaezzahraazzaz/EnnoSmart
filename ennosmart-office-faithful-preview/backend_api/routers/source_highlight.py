# -*- coding: utf-8 -*-
from __future__ import annotations

"""
EnnoSmart — Source Highlight V170
=================================

Objectif
--------
Afficher les documents sources de façon fidèle dans EnnoDiagnostic / Comparaison docs.

Principe principal :
- PDF : affichage PDF natif + surlignage PyMuPDF.
- DOC/DOCX/DOCM/ODT/RTF : conversion LibreOffice -> PDF -> surlignage PDF.
- PPT/PPTX/PPTM/ODP : conversion LibreOffice -> PDF -> surlignage PDF.
- XLS/XLSX/XLSM/ODS : conversion LibreOffice -> PDF -> surlignage PDF.
- Images : affichage natif.
- TXT/MD/JSON/CSV/XML/LOG/HTML : aperçu HTML texte.
- MSG/EML : aperçu HTML.

La conversion Office -> PDF conserve bien mieux que l'ancien rendu texte :
- figures et images ;
- tableaux ;
- graphiques ;
- polices / tailles ;
- pagination ;
- en-têtes / pieds de page ;
- position générale des objets.

Aucun organisme, projet, fichier ou année n'est hardcodé.

Dépendances Python :
    pip install pymupdf python-docx openpyxl extract-msg

Dépendance système recommandée :
    LibreOffice / soffice

Linux / OVH :
    sudo apt update
    sudo apt install -y libreoffice libreoffice-writer libreoffice-calc libreoffice-impress fonts-liberation

Windows :
    Installer LibreOffice normalement.
    Le code cherche automatiquement soffice.exe dans Program Files.
"""

import hashlib
import html
import json
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
import zipfile
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple
from urllib.parse import quote

from fastapi import APIRouter, Body, Depends, HTTPException, Query, Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, PlainTextResponse
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from core.deps import get_current_user, get_db
from db.models import User
from services.diagnostic_service import get_project_store
from services.project_service import get_project_for_user

try:
    from db.models import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None  # type: ignore

try:
    from sqlalchemy.orm import undefer
except Exception:  # pragma: no cover
    undefer = None  # type: ignore


router = APIRouter(tags=["source-preview"])

PROJECT_ROOT = Path(
    os.getenv("ENNOSMART_ROOT") or Path(__file__).resolve().parents[2]
).resolve()
STORAGE_ROOT = Path(
    os.getenv("ENNOSMART_STORAGE_ROOT") or PROJECT_ROOT / "storage"
).resolve()
PREVIEW_ROOT = STORAGE_ROOT / "previews" / "source_highlight"
OFFICE_PDF_CACHE = PREVIEW_ROOT / "office_pdf"

PDF_EXTENSIONS = {".pdf"}

WORD_EXTENSIONS = {
    ".doc",
    ".docx",
    ".docm",
    ".odt",
    ".rtf",
}

PRESENTATION_EXTENSIONS = {
    ".ppt",
    ".pptx",
    ".pptm",
    ".odp",
}

SPREADSHEET_EXTENSIONS = {
    ".xls",
    ".xlsx",
    ".xlsm",
    ".ods",
}

OFFICE_TO_PDF_EXTENSIONS = (
    WORD_EXTENSIONS
    | PRESENTATION_EXTENSIONS
    | SPREADSHEET_EXTENSIONS
)

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".json",
    ".csv",
    ".xml",
    ".log",
    ".html",
    ".htm",
}

EMAIL_EXTENSIONS = {".msg", ".eml"}

IMAGE_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".webp",
    ".bmp",
    ".tif",
    ".tiff",
}

ALLOWED_EXTENSIONS = (
    PDF_EXTENSIONS
    | OFFICE_TO_PDF_EXTENSIONS
    | TEXT_EXTENSIONS
    | EMAIL_EXTENSIONS
    | IMAGE_EXTENSIONS
)

MAX_EXCERPT_CHARS = 8_000
MAX_HTML_TEXT_CHARS = 2_000_000
MAX_PROJECT_SEARCH_FILES = 2_500

OFFICE_CONVERT_TIMEOUT_SECONDS = int(
    os.getenv("ENNOSMART_OFFICE_CONVERT_TIMEOUT", "120")
)

OFFICE_PDF_CACHE_ENABLED = str(
    os.getenv("ENNOSMART_OFFICE_PDF_CACHE", "1")
).strip().lower() not in {"0", "false", "no", "off"}


class SourceHighlightRequest(BaseModel):
    excerpt: str = Field(default="", max_length=MAX_EXCERPT_CHARS)
    document_id: Optional[int] = None
    source_path: Optional[str] = None
    source_name: Optional[str] = None
    document_name: Optional[str] = None
    title: Optional[str] = None
    role: Optional[str] = None
    passage_id: Optional[str] = None
    page_number: Optional[int] = None
    paragraph_index: Optional[int] = None
    char_start: Optional[int] = None
    char_end: Optional[int] = None
    year: Optional[str] = None
    return_json: bool = False


# ---------------------------------------------------------------------------
# Utilitaires
# ---------------------------------------------------------------------------


def _ensure_preview_root() -> None:
    PREVIEW_ROOT.mkdir(parents=True, exist_ok=True)
    OFFICE_PDF_CACHE.mkdir(parents=True, exist_ok=True)


def _safe_filename(value: Any, fallback: str = "document") -> str:
    name = Path(str(value or fallback).replace("\\", "/")).name
    name = re.sub(r'[\\/:*?"<>|]+', "_", name).strip(" ._")
    return name or fallback


def _strip_accents(value: str) -> str:
    return "".join(
        char
        for char in unicodedata.normalize("NFD", str(value or ""))
        if unicodedata.category(char) != "Mn"
    )


def normalize_text(value: Any) -> str:
    text = _strip_accents(str(value or "")).lower()
    text = text.replace("\u00a0", " ").replace("’", "'")
    text = re.sub(r"[^a-z0-9%.,;:/+\-_' ]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def normalize_filename(value: Any) -> str:
    name = Path(str(value or "").replace("\\", "/")).name
    stem = Path(name).stem
    stem = re.sub(r"_[a-f0-9]{10,64}$", "", stem, flags=re.I)
    return normalize_text(stem)


def clean_excerpt(value: Any) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"#{1,6}\s*", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()[:MAX_EXCERPT_CHARS]


def short_hash(*parts: Any) -> str:
    raw = "||".join(str(part or "") for part in parts)
    return hashlib.sha256(
        raw.encode("utf-8", errors="ignore")
    ).hexdigest()[:24]


def is_inside(child: Path, parent: Path) -> bool:
    try:
        child.resolve().relative_to(parent.resolve())
        return True
    except Exception:
        return False


def _append_unique_path(
    output: List[Path],
    seen: set[str],
    candidate: Path,
) -> None:
    try:
        resolved = candidate.resolve()
    except Exception:
        return

    key = str(resolved).lower()
    if key in seen:
        return

    output.append(resolved)
    seen.add(key)


# ---------------------------------------------------------------------------
# Résolution sécurisée des documents
# ---------------------------------------------------------------------------


def _same_project_history_roots(project: Any) -> List[Path]:
    project_store = get_project_store(project)

    seeds = [
        Path(project_store.project_dir),
        Path(project_store.documents_raw_dir),
    ]

    output: List[Path] = []
    seen: set[str] = set()

    for seed in seeds:
        try:
            resolved = seed.resolve()
        except Exception:
            continue

        _append_unique_path(output, seen, resolved)

        lineage = [resolved, *resolved.parents]

        for node in lineage:
            if node.name.lower() == "years":
                project_root = node.parent
                if is_inside(project_root, STORAGE_ROOT):
                    _append_unique_path(output, seen, project_root)
                break

        for node in lineage:
            if node.parent.name.lower() == "projects":
                if is_inside(node, STORAGE_ROOT):
                    _append_unique_path(output, seen, node)
                break

    return output


def _project_allowed_roots(project: Any) -> List[Path]:
    output: List[Path] = []
    seen: set[str] = set()

    for candidate in [
        *_same_project_history_roots(project),
        PREVIEW_ROOT,
    ]:
        _append_unique_path(output, seen, candidate)

    return output


def _safe_project_path(
    project: Any,
    raw_value: Optional[str],
) -> Optional[Path]:
    if not raw_value:
        return None

    value = str(raw_value).strip().strip('"').strip("'")
    if not value:
        return None

    path = Path(value)
    project_store = get_project_store(project)

    candidates = (
        [path]
        if path.is_absolute()
        else [
            Path(project_store.project_dir) / path,
            Path(project_store.documents_raw_dir) / path,
            PROJECT_ROOT / path,
        ]
    )

    allowed_roots = _project_allowed_roots(project)

    for candidate in candidates:
        try:
            resolved = candidate.resolve()
        except Exception:
            continue

        if not resolved.exists() or not resolved.is_file():
            continue

        if resolved.suffix.lower() not in ALLOWED_EXTENSIONS:
            continue

        if any(is_inside(resolved, root) for root in allowed_roots):
            return resolved

    return None


def _query_project_document(
    db: Session,
    project_id: int,
    document_id: int,
) -> Any:
    if Document is None:
        return None

    query = db.query(Document)

    if undefer is not None and hasattr(Document, "file_data"):
        try:
            query = query.options(undefer(Document.file_data))
        except Exception:
            pass

    return (
        query.filter(
            getattr(Document, "id") == document_id,
            getattr(Document, "project_id") == project_id,
        )
        .first()
    )


def _materialize_db_document(
    document: Any,
    project_id: int,
) -> Optional[Path]:
    if document is None:
        return None

    filename = _safe_filename(
        getattr(document, "stored_filename", None)
        or getattr(document, "filename", None)
        or f"document_{getattr(document, 'id', 'unknown')}"
    )

    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        return None

    file_path = getattr(document, "file_path", None)

    if file_path and not str(file_path).startswith("db://"):
        path = Path(str(file_path))
        if path.exists() and path.is_file():
            return path.resolve()

    file_data = getattr(document, "file_data", None)
    if not file_data:
        return None

    _ensure_preview_root()

    document_id = getattr(document, "id", "unknown")
    target_dir = (
        PREVIEW_ROOT
        / "materialized"
        / str(project_id)
        / str(document_id)
    )
    target_dir.mkdir(parents=True, exist_ok=True)

    target = target_dir / filename
    raw = bytes(file_data)

    if not target.exists() or target.stat().st_size != len(raw):
        target.write_bytes(raw)

    return target.resolve()


def _file_name_score(
    path: Path,
    wanted_names: Iterable[str],
) -> int:
    candidate = normalize_filename(path.name)
    if not candidate:
        return 0

    best = 0

    for raw in wanted_names:
        wanted = normalize_filename(raw)
        if not wanted:
            continue

        if candidate == wanted:
            best = max(best, 100)
            continue

        if candidate in wanted or wanted in candidate:
            best = max(best, 92)
            continue

        wanted_words = set(wanted.split())
        candidate_words = set(candidate.split())

        common = len(wanted_words & candidate_words)
        ratio = common / max(
            1,
            min(len(wanted_words), len(candidate_words)),
        )

        if common >= 2 and ratio >= 0.7:
            best = max(best, int(70 + ratio * 20))

    return best


def _search_project_file(
    project: Any,
    wanted_names: List[str],
    requested_year: Optional[str] = None,
) -> Optional[Path]:
    project_store = get_project_store(project)

    roots = [
        Path(project_store.documents_raw_dir),
        Path(project_store.project_dir),
        *_same_project_history_roots(project),
    ]

    year_text = str(requested_year or "").strip()

    if year_text:
        year_roots: List[Path] = []

        for root in list(roots):
            try:
                resolved = root.resolve()
            except Exception:
                continue

            candidate = resolved / "years" / year_text
            if candidate.exists():
                year_roots.append(candidate)

            if resolved.name.lower() == "years":
                candidate = resolved / year_text
                if candidate.exists():
                    year_roots.append(candidate)

        roots = [*year_roots, *roots]

    best_path: Optional[Path] = None
    best_score = 0
    seen: set[str] = set()
    examined = 0

    for root in roots:
        if not root.exists():
            continue

        try:
            for path in root.rglob("*"):
                if examined >= MAX_PROJECT_SEARCH_FILES:
                    break

                examined += 1

                if (
                    not path.is_file()
                    or path.suffix.lower() not in ALLOWED_EXTENSIONS
                ):
                    continue

                key = str(path.resolve()).lower()
                if key in seen:
                    continue

                seen.add(key)

                score = _file_name_score(path, wanted_names)

                if score > best_score:
                    best_path = path.resolve()
                    best_score = score

                if score >= 100:
                    return best_path
        except Exception:
            continue

    return best_path if best_score >= 70 else None


def resolve_document_path(
    *,
    db: Session,
    project: Any,
    payload: SourceHighlightRequest,
) -> Path:
    # document_id peut être un ID PostgreSQL OU un ID provenant du NLP/RAG.
    # Ne jamais lever 404 trop tôt : on doit pouvoir continuer par nom/path.
    if payload.document_id:
        document = _query_project_document(
            db,
            project.id,
            payload.document_id,
        )

        if document is not None:
            materialized = _materialize_db_document(
                document,
                project.id,
            )
            if materialized:
                return materialized

    direct = _safe_project_path(
        project,
        payload.source_path,
    )

    wanted_names = [
        payload.source_name or "",
        payload.document_name or "",
        (
            Path(payload.source_path or "").name
            if payload.source_path
            else ""
        ),
    ]
    wanted_names = [
        item
        for item in wanted_names
        if str(item).strip()
    ]

    direct_is_requested_document = False

    if direct:
        requested_extensions = {
            Path(str(name).replace("\\", "/")).suffix.lower()
            for name in wanted_names
            if Path(str(name).replace("\\", "/")).suffix
        }

        direct_is_requested_document = (
            not requested_extensions
            or direct.suffix.lower() in requested_extensions
            or normalize_filename(direct.name)
            in {
                normalize_filename(name)
                for name in wanted_names
            }
        )

        if direct_is_requested_document:
            return direct

    found = _search_project_file(
        project,
        wanted_names,
        requested_year=payload.year,
    )

    if found:
        return found

    if direct:
        return direct

    requested = (
        ", ".join(
            _safe_filename(name)
            for name in wanted_names[:4]
        )
        or "nom absent"
    )

    year_hint = (
        f" pour l'année {payload.year}"
        if payload.year
        else ""
    )

    raise HTTPException(
        status_code=404,
        detail=(
            "Document source introuvable dans PostgreSQL "
            "ou dans les dossiers du même projet"
            f"{year_hint}. Noms recherchés : {requested}."
        ),
    )


# ---------------------------------------------------------------------------
# LibreOffice : détection + conversion fidèle Office -> PDF
# ---------------------------------------------------------------------------


def _libreoffice_candidates() -> List[str]:
    env_bin = str(
        os.getenv("LIBREOFFICE_BIN")
        or os.getenv("SOFFICE_BIN")
        or ""
    ).strip()

    candidates: List[str] = []

    if env_bin:
        candidates.append(env_bin)

    for binary in (
        shutil.which("libreoffice"),
        shutil.which("soffice"),
        shutil.which("soffice.exe"),
    ):
        if binary:
            candidates.append(binary)

    program_files = [
        os.getenv("PROGRAMFILES"),
        os.getenv("PROGRAMFILES(X86)"),
        os.getenv("LOCALAPPDATA"),
    ]

    for root in program_files:
        if not root:
            continue

        candidates.extend(
            [
                str(
                    Path(root)
                    / "LibreOffice"
                    / "program"
                    / "soffice.exe"
                ),
                str(
                    Path(root)
                    / "LibreOffice"
                    / "program"
                    / "soffice.com"
                ),
            ]
        )

    output: List[str] = []
    seen = set()

    for candidate in candidates:
        value = str(candidate or "").strip()
        if not value:
            continue

        key = value.lower()
        if key in seen:
            continue

        seen.add(key)

        path = Path(value)

        if path.exists():
            output.append(str(path))
            continue

        resolved = shutil.which(value)
        if resolved:
            output.append(resolved)

    return output


def find_libreoffice_binary() -> Optional[str]:
    candidates = _libreoffice_candidates()
    return candidates[0] if candidates else None


def _office_cache_target(
    source: Path,
    project_id: int,
) -> Path:
    stat = source.stat()

    fingerprint = short_hash(
        source.resolve(),
        stat.st_size,
        stat.st_mtime_ns,
    )

    target_dir = (
        OFFICE_PDF_CACHE
        / str(project_id)
        / fingerprint
    )
    target_dir.mkdir(parents=True, exist_ok=True)

    return target_dir / f"{source.stem}.pdf"


def convert_office_to_pdf(
    source: Path,
    project_id: int,
) -> Path:
    """
    Convertit un document Office en PDF avec LibreOffice.

    Le PDF est mis en cache sur :
        storage/previews/source_highlight/office_pdf/<project>/<hash>/...

    Ainsi, un même DOCX/PPTX n'est pas reconverti à chaque clic.
    """
    suffix = source.suffix.lower()

    if suffix not in OFFICE_TO_PDF_EXTENSIONS:
        raise RuntimeError(
            f"Format Office non pris en charge : {suffix}"
        )

    _ensure_preview_root()

    target = _office_cache_target(
        source,
        project_id,
    )

    if (
        OFFICE_PDF_CACHE_ENABLED
        and target.exists()
        and target.stat().st_size > 0
    ):
        return target.resolve()

    libreoffice = find_libreoffice_binary()

    if not libreoffice:
        raise RuntimeError(
            "LibreOffice / soffice est introuvable. "
            "Installe LibreOffice ou définis LIBREOFFICE_BIN."
        )

    target.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    # LibreOffice peut verrouiller son profil utilisateur si plusieurs
    # conversions sont lancées en parallèle. Chaque conversion utilise donc
    # son propre profil temporaire.
    profile_dir = Path(
        tempfile.mkdtemp(
            prefix="ennosmart_lo_",
            dir=str(PREVIEW_ROOT),
        )
    )

    try:
        profile_uri = profile_dir.resolve().as_uri()

        command = [
            libreoffice,
            f"-env:UserInstallation={profile_uri}",
            "--headless",
            "--nologo",
            "--nodefault",
            "--nolockcheck",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(target.parent),
            str(source.resolve()),
        ]

        completed = subprocess.run(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=OFFICE_CONVERT_TIMEOUT_SECONDS,
            check=False,
            env={
                **os.environ,
                "HOME": os.environ.get(
                    "HOME",
                    str(profile_dir),
                ),
            },
        )

        produced = target.parent / f"{source.stem}.pdf"

        if (
            completed.returncode != 0
            or not produced.exists()
            or produced.stat().st_size <= 0
        ):
            message = (
                completed.stderr.strip()
                or completed.stdout.strip()
                or "LibreOffice n'a pas produit de PDF."
            )

            raise RuntimeError(
                "Conversion Office -> PDF impossible : "
                + message[:1_500]
            )

        if produced.resolve() != target.resolve():
            if target.exists():
                target.unlink()

            produced.replace(target)

        metadata = {
            "source": str(source.resolve()),
            "source_size": source.stat().st_size,
            "source_mtime_ns": source.stat().st_mtime_ns,
            "pdf": str(target.resolve()),
            "libreoffice": libreoffice,
            "stdout": completed.stdout[-2_000:],
        }

        target.with_suffix(".conversion.json").write_text(
            json.dumps(
                metadata,
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )

        return target.resolve()

    finally:
        shutil.rmtree(
            profile_dir,
            ignore_errors=True,
        )


# ---------------------------------------------------------------------------
# Extraction texte de fallback
# ---------------------------------------------------------------------------


def read_text_file(path: Path) -> str:
    for encoding in (
        "utf-8",
        "utf-8-sig",
        "cp1252",
        "latin-1",
    ):
        try:
            return path.read_text(
                encoding=encoding,
                errors="ignore",
            )
        except Exception:
            continue

    return ""


def read_docx_text(path: Path) -> str:
    # Fallback uniquement. Le rendu visuel Word utilise LibreOffice -> PDF.
    try:
        from docx import Document as DocxDocument  # type: ignore

        doc = DocxDocument(str(path))
        parts: List[str] = []

        for paragraph in doc.paragraphs:
            value = (paragraph.text or "").strip()
            if value:
                parts.append(value)

        for table in doc.tables:
            for row in table.rows:
                cells = [
                    (cell.text or "").strip()
                    for cell in row.cells
                ]
                cells = [
                    cell
                    for cell in cells
                    if cell
                ]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n\n".join(parts)

    except Exception:
        pass

    try:
        with zipfile.ZipFile(path) as archive:
            xml = archive.read(
                "word/document.xml"
            ).decode(
                "utf-8",
                errors="ignore",
            )

        xml = re.sub(
            r"<w:tab[^>]*>",
            "\t",
            xml,
        )
        xml = re.sub(
            r"</w:p>",
            "\n",
            xml,
        )
        xml = re.sub(
            r"<[^>]+>",
            "",
            xml,
        )

        return html.unescape(xml)

    except Exception:
        return ""


def read_pdf_text(
    path: Path,
    max_pages: int = 200,
) -> str:
    try:
        import fitz  # type: ignore
    except Exception:
        return ""

    try:
        document = fitz.open(str(path))
        chunks: List[str] = []

        for page_index, page in enumerate(document):
            if page_index >= max_pages:
                break

            chunks.append(
                page.get_text("text") or ""
            )

        document.close()
        return "\n".join(chunks)

    except Exception:
        return ""


def read_msg_text(path: Path) -> str:
    try:
        import extract_msg  # type: ignore

        message = extract_msg.Message(str(path))

        parts = [
            (
                f"Objet : {message.subject}"
                if message.subject
                else ""
            ),
            (
                f"De : {message.sender}"
                if message.sender
                else ""
            ),
            (
                f"À : {message.to}"
                if message.to
                else ""
            ),
            message.body or "",
        ]

        try:
            message.close()
        except Exception:
            pass

        return "\n\n".join(
            part
            for part in parts
            if str(part).strip()
        )

    except Exception:
        return ""


def read_eml_text(path: Path) -> str:
    try:
        from email import policy
        from email.parser import BytesParser

        message = BytesParser(
            policy=policy.default
        ).parsebytes(
            path.read_bytes()
        )

        parts = [
            (
                f"Objet : {message.get('subject', '')}"
                if message.get("subject")
                else ""
            ),
            (
                f"De : {message.get('from', '')}"
                if message.get("from")
                else ""
            ),
            (
                f"À : {message.get('to', '')}"
                if message.get("to")
                else ""
            ),
        ]

        body = message.get_body(
            preferencelist=("plain", "html")
        )

        if body:
            content = body.get_content()

            if body.get_content_type() == "text/html":
                content = re.sub(
                    r"<[^>]+>",
                    " ",
                    content,
                )
                content = html.unescape(content)

            parts.append(content)

        return "\n\n".join(
            part
            for part in parts
            if str(part).strip()
        )

    except Exception:
        return ""


def read_document_text(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix in TEXT_EXTENSIONS:
        return read_text_file(path)

    if suffix in WORD_EXTENSIONS:
        return read_docx_text(path)

    if suffix in PDF_EXTENSIONS:
        return read_pdf_text(path)

    if suffix == ".msg":
        return read_msg_text(path)

    if suffix == ".eml":
        return read_eml_text(path)

    return ""


# ---------------------------------------------------------------------------
# Recherche / surlignage PDF
# ---------------------------------------------------------------------------


def overlap_score(
    excerpt: str,
    text: str,
) -> float:
    normalized_excerpt = normalize_text(excerpt)
    normalized_text = normalize_text(text)

    if not normalized_excerpt or not normalized_text:
        return 0.0

    if normalized_excerpt in normalized_text:
        return 1.0

    words = list(
        dict.fromkeys(
            re.findall(
                r"[a-z0-9]{4,}",
                normalized_excerpt,
            )
        )
    )[:100]

    if not words:
        return 0.0

    hits = sum(
        1
        for word in words
        if word in normalized_text
    )

    return hits / max(1, len(words))


def make_search_queries(excerpt: str) -> List[str]:
    text = clean_excerpt(excerpt)

    if not text:
        return []

    candidates = [text[:300]]

    candidates.extend(
        sentence.strip()
        for sentence in re.split(
            r"(?<=[.!?;:])\s+",
            text,
        )
        if 25 <= len(sentence.strip()) <= 220
    )

    words = text.split()

    for size in (20, 16, 12, 9, 7, 5):
        if len(words) < size:
            continue

        step = max(1, size // 2)

        for start in range(
            0,
            min(
                len(words) - size + 1,
                60,
            ),
            step,
        ):
            query = " ".join(
                words[start : start + size]
            ).strip()

            if len(query) >= 18:
                candidates.append(query)

    output: List[str] = []
    seen: set[str] = set()

    for candidate in candidates:
        candidate = re.sub(
            r"\s+",
            " ",
            candidate,
        ).strip()

        key = normalize_text(candidate)

        if candidate and key and key not in seen:
            output.append(candidate)
            seen.add(key)

    return output[:50]


def _normalized_page_hint(
    page_number: Optional[int],
    page_count: int,
) -> Optional[int]:
    if page_number is None or page_count <= 0:
        return None

    try:
        raw = int(page_number)
    except Exception:
        return None

    candidates = (
        [raw - 1, raw]
        if raw > 0
        else [0]
    )

    for candidate in candidates:
        if 0 <= candidate < page_count:
            return candidate

    return None


def _pdf_search_order(
    page_count: int,
    page_hint: Optional[int],
) -> List[int]:
    if page_hint is None:
        return list(range(page_count))

    output = [page_hint]

    for distance in range(1, page_count):
        for candidate in (
            page_hint - distance,
            page_hint + distance,
        ):
            if (
                0 <= candidate < page_count
                and candidate not in output
            ):
                output.append(candidate)

    return output


def best_pdf_page_by_overlap(
    document: Any,
    excerpt: str,
) -> int:
    best_page = 0
    best_score = -1.0

    for index, page in enumerate(document):
        try:
            score = overlap_score(
                excerpt,
                page.get_text("text") or "",
            )
        except Exception:
            score = 0.0

        if score > best_score:
            best_score = score
            best_page = index

    return best_page


def create_highlighted_pdf(
    path: Path,
    excerpt: str,
    project_id: int,
    page_number: Optional[int] = None,
) -> Tuple[Path, bool, int, int]:
    try:
        import fitz  # type: ignore
    except Exception as exc:
        raise RuntimeError(
            "PyMuPDF n'est pas installé. "
            "Lance : pip install pymupdf"
        ) from exc

    _ensure_preview_root()

    stat = path.stat()

    output = PREVIEW_ROOT / (
        "highlight_"
        f"p{project_id}_"
        f"{short_hash(path, stat.st_mtime_ns, excerpt, page_number)}"
        ".pdf"
    )

    metadata_path = output.with_suffix(
        ".metadata.json"
    )

    if (
        output.exists()
        and output.stat().st_size > 0
        and metadata_path.exists()
    ):
        try:
            metadata = json.loads(
                metadata_path.read_text(
                    encoding="utf-8"
                )
            )

            return (
                output,
                bool(metadata.get("exact")),
                int(metadata.get("page_index") or 0),
                int(metadata.get("found_count") or 0),
            )

        except Exception:
            pass

    document = fitz.open(str(path))

    page_hint = _normalized_page_hint(
        page_number,
        len(document),
    )

    queries = make_search_queries(excerpt)

    found_count = 0
    found_page = page_hint or 0

    try:
        for page_index in _pdf_search_order(
            len(document),
            page_hint,
        ):
            page = document[page_index]
            page_found = False

            for query in queries:
                try:
                    rectangles = page.search_for(query)
                except Exception:
                    rectangles = []

                if not rectangles:
                    continue

                for rectangle in rectangles[:20]:
                    highlight = page.add_highlight_annot(
                        rectangle
                    )
                    highlight.set_colors(
                        stroke=(1.0, 0.86, 0.20)
                    )
                    highlight.update()

                    border = page.add_rect_annot(
                        rectangle
                    )
                    border.set_colors(
                        stroke=(1.0, 0.42, 0.0)
                    )
                    border.set_border(width=1.0)
                    border.update()

                    found_count += 1
                    page_found = True

                if page_found:
                    found_page = page_index
                    break

            if page_found:
                break

        exact = found_count > 0

        # Important :
        # Si le passage n'est pas retrouvé exactement, on garde tout le vrai
        # document et on positionne une note sur la page la plus probable.
        if not exact and len(document) > 0:
            found_page = (
                page_hint
                if page_hint is not None
                else best_pdf_page_by_overlap(
                    document,
                    excerpt,
                )
            )

            page = document[found_page]
            page_rect = page.rect

            banner = fitz.Rect(
                page_rect.x0 + 24,
                page_rect.y0 + 24,
                page_rect.x1 - 24,
                min(
                    page_rect.y0 + 140,
                    page_rect.y1 - 24,
                ),
            )

            box = page.add_rect_annot(banner)
            box.set_colors(
                stroke=(1.0, 0.42, 0.0),
                fill=(1.0, 0.94, 0.72),
            )
            box.set_opacity(0.18)
            box.set_border(width=1.3)
            box.update()

            page.insert_textbox(
                banner,
                (
                    "Passage recherché — correspondance approximative\n"
                    + clean_excerpt(excerpt)[:450]
                ),
                fontsize=8.2,
                color=(0.20, 0.12, 0.02),
            )

        if output.exists():
            output.unlink()

        document.save(
            str(output),
            garbage=4,
            deflate=True,
        )

        metadata_path.write_text(
            json.dumps(
                {
                    "exact": exact,
                    "page_index": found_page,
                    "found_count": found_count,
                    "source_path": str(path),
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )

        return (
            output,
            exact,
            found_page,
            found_count,
        )

    finally:
        document.close()


# ---------------------------------------------------------------------------
# HTML fallback
# ---------------------------------------------------------------------------


def _find_text_range(
    full_text: str,
    excerpt: str,
) -> Optional[Tuple[int, int]]:
    clean = clean_excerpt(excerpt)

    if not full_text or not clean:
        return None

    exact = full_text.lower().find(
        clean.lower()
    )

    if exact >= 0:
        return (
            exact,
            exact + len(clean),
        )

    for query in make_search_queries(clean):
        index = full_text.lower().find(
            query.lower()
        )

        if index >= 0:
            return (
                index,
                index + len(query),
            )

    return None


def highlight_html_text(
    full_text: str,
    excerpt: str,
) -> Tuple[str, bool]:
    full_text = (
        full_text or ""
    )[:MAX_HTML_TEXT_CHARS]

    selected_range = _find_text_range(
        full_text,
        excerpt,
    )

    if selected_range:
        start, end = selected_range

        before = html.escape(
            full_text[:start]
        )
        selected = html.escape(
            full_text[start:end]
        )
        after = html.escape(
            full_text[end:]
        )

        return (
            "<div class='doc'>"
            f"{before}"
            "<mark id='selected-passage'>"
            f"{selected}"
            "</mark>"
            f"{after}"
            "</div>",
            True,
        )

    safe_document = html.escape(full_text)
    safe_excerpt = html.escape(
        clean_excerpt(excerpt)
        or "Aucun extrait disponible."
    )

    return (
        "<p class='muted'>"
        "Le document a été retrouvé, mais le passage exact "
        "n'a pas été localisé mot à mot."
        "</p>"
        f"<div class='excerpt' id='selected-passage'>{safe_excerpt}</div>"
        "<hr/>"
        f"<div class='doc'>{safe_document}</div>",
        False,
    )


def html_document(
    title: str,
    body: str,
    source: str = "",
    exact: bool = True,
) -> str:
    safe_title = html.escape(
        title or "Prévisualisation source"
    )
    safe_source = html.escape(
        source or ""
    )

    badge = (
        "Passage retrouvé"
        if exact
        else "Correspondance approximative"
    )
    badge_class = (
        "ok"
        if exact
        else "warn"
    )

    return f"""<!doctype html>
<html lang="fr">
<head>
  <meta charset="utf-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1"/>
  <title>{safe_title}</title>
  <style>
    html {{ scroll-behavior:smooth; }}
    body {{ margin:0; font-family:Inter,Arial,sans-serif; background:#f8fafc; color:#0f172a; }}
    .wrap {{ max-width:1180px; margin:0 auto; padding:22px; }}
    .card {{ background:#fff; border:1px solid #e2e8f0; border-radius:16px; overflow:hidden; box-shadow:0 8px 28px rgba(15,23,42,.06); }}
    .head {{ position:sticky; top:0; z-index:5; display:flex; justify-content:space-between; gap:16px; padding:16px 20px; background:#fff; border-bottom:1px solid #e2e8f0; }}
    h1 {{ margin:0; font-size:16px; }}
    .source {{ margin-top:5px; color:#64748b; font-size:12px; word-break:break-all; }}
    .badge {{ align-self:flex-start; border-radius:999px; padding:4px 10px; font-size:12px; font-weight:650; border:1px solid; }}
    .badge.ok {{ color:#047857; background:#ecfdf5; border-color:#a7f3d0; }}
    .badge.warn {{ color:#c2410c; background:#fff7ed; border-color:#fed7aa; }}
    .content {{ padding:22px; }}
    .doc {{ white-space:pre-wrap; line-height:1.72; font-size:14px; }}
    mark {{ background:#fef08a; outline:2px solid #f97316; border-radius:4px; padding:2px 3px; color:#111827; }}
    .excerpt {{ white-space:pre-wrap; line-height:1.7; background:#fff7ed; border:2px solid #f97316; border-radius:12px; padding:14px; }}
    .muted {{ color:#64748b; font-size:13px; }}
    hr {{ border:0; border-top:1px solid #e2e8f0; margin:20px 0; }}
  </style>
</head>
<body onload="document.getElementById('selected-passage')?.scrollIntoView({{block:'center'}})">
  <div class="wrap">
    <div class="card">
      <div class="head">
        <div>
          <h1>{safe_title}</h1>
          <div class="source">{safe_source}</div>
        </div>
        <span class="badge {badge_class}">{badge}</span>
      </div>
      <div class="content">{body}</div>
    </div>
  </div>
</body>
</html>"""


def create_html_preview(
    path: Path,
    excerpt: str,
    title: str,
) -> HTMLResponse:
    full_text = read_document_text(path)

    if not full_text:
        body = (
            "<p class='muted'>"
            "Le texte du document n'a pas pu être extrait."
            "</p>"
            "<div class='excerpt' id='selected-passage'>"
            f"{html.escape(clean_excerpt(excerpt))}"
            "</div>"
        )

        return HTMLResponse(
            html_document(
                title or path.name,
                body,
                str(path),
                exact=False,
            )
        )

    body, exact = highlight_html_text(
        full_text,
        excerpt,
    )

    return HTMLResponse(
        html_document(
            title or path.name,
            body,
            str(path),
            exact=exact,
        )
    )


# ---------------------------------------------------------------------------
# Rendu final
# ---------------------------------------------------------------------------


def inline_file_response(
    path: Path,
    media_type: Optional[str] = None,
    *,
    extra_headers: Optional[Dict[str, str]] = None,
) -> FileResponse:
    encoded = quote(path.name)

    headers = {
        "Content-Disposition": (
            f"inline; filename*=UTF-8''{encoded}"
        ),
        "Cache-Control": "no-store",
    }

    if extra_headers:
        headers.update(extra_headers)

    return FileResponse(
        path,
        media_type=media_type,
        headers=headers,
    )


def _render_pdf(
    pdf_path: Path,
    *,
    original_path: Path,
    payload: SourceHighlightRequest,
    project_id: int,
    converted_from_office: bool,
):
    excerpt = clean_excerpt(
        payload.excerpt
    )

    base_headers = {
        "X-EnnoSmart-Source-File": original_path.name,
        "X-EnnoSmart-Rendered-File": pdf_path.name,
        "X-EnnoSmart-Preview-Mode": (
            "office-pdf"
            if converted_from_office
            else "pdf"
        ),
        "X-EnnoSmart-Office-Converted": (
            "true"
            if converted_from_office
            else "false"
        ),
    }

    if not excerpt:
        return inline_file_response(
            pdf_path,
            "application/pdf",
            extra_headers=base_headers,
        )

    highlighted, exact, page_index, count = (
        create_highlighted_pdf(
            pdf_path,
            excerpt,
            project_id,
            payload.page_number,
        )
    )

    return inline_file_response(
        highlighted,
        "application/pdf",
        extra_headers={
            **base_headers,
            "X-EnnoSmart-Highlight-Exact": (
                "true"
                if exact
                else "false"
            ),
            "X-EnnoSmart-Highlight-Page": str(
                page_index + 1
            ),
            "X-EnnoSmart-Highlight-Count": str(
                count
            ),
        },
    )


def render_document_preview(
    path: Path,
    payload: SourceHighlightRequest,
    project_id: int,
):
    suffix = path.suffix.lower()

    if suffix in PDF_EXTENSIONS:
        return _render_pdf(
            path,
            original_path=path,
            payload=payload,
            project_id=project_id,
            converted_from_office=False,
        )

    # IMPORTANT :
    # C'est le changement principal par rapport à l'ancien code.
    # On ne transforme plus DOCX/PPTX/XLSX en texte HTML.
    # On convertit le vrai fichier Office en PDF pour préserver son rendu.
    if suffix in OFFICE_TO_PDF_EXTENSIONS:
        try:
            office_pdf = convert_office_to_pdf(
                path,
                project_id,
            )

            return _render_pdf(
                office_pdf,
                original_path=path,
                payload=payload,
                project_id=project_id,
                converted_from_office=True,
            )

        except Exception as exc:
            # Fallback lisible si LibreOffice n'est pas installé.
            # Pour Word/texte, on garde l'ancien rendu texte afin de ne pas
            # casser totalement l'interface, mais on indique clairement
            # que le rendu fidèle n'est pas disponible.
            fallback_text = read_document_text(path)

            if fallback_text:
                body, exact = highlight_html_text(
                    fallback_text,
                    payload.excerpt,
                )

                warning = (
                    "<div style='margin-bottom:14px;padding:12px;"
                    "border:1px solid #fdba74;background:#fff7ed;"
                    "border-radius:10px;color:#9a3412;font-size:13px'>"
                    "<strong>Rendu Office fidèle indisponible.</strong><br/>"
                    + html.escape(str(exc))
                    + "</div>"
                )

                return HTMLResponse(
                    html_document(
                        payload.title or path.name,
                        warning + body,
                        str(path),
                        exact=exact,
                    )
                )

            raise HTTPException(
                status_code=503,
                detail=(
                    "Prévisualisation Office fidèle indisponible : "
                    + str(exc)
                ),
            )

    if suffix in IMAGE_EXTENSIONS:
        media_type = (
            mimetypes.guess_type(path.name)[0]
            or "application/octet-stream"
        )

        return inline_file_response(
            path,
            media_type,
            extra_headers={
                "X-EnnoSmart-Preview-Mode": "image",
                "X-EnnoSmart-Source-File": path.name,
            },
        )

    if suffix in TEXT_EXTENSIONS | EMAIL_EXTENSIONS:
        return create_html_preview(
            path,
            clean_excerpt(payload.excerpt),
            payload.title or path.name,
        )

    return PlainTextResponse(
        (
            "Aperçu indisponible pour ce format : "
            f"{suffix}"
        ),
        status_code=415,
    )


def build_preview_url(
    project_id: int,
    payload: SourceHighlightRequest,
) -> str:
    params = {
        "excerpt": payload.excerpt,
        "document_id": payload.document_id,
        "source_path": payload.source_path,
        "source_name": (
            payload.source_name
            or payload.document_name
        ),
        "document_name": payload.document_name,
        "title": payload.title,
        "page_number": payload.page_number,
        "passage_id": payload.passage_id,
        "year": payload.year,
    }

    query = "&".join(
        f"{key}={quote(str(value))}"
        for key, value in params.items()
        if value not in (None, "")
    )

    return (
        f"/projects/{project_id}/"
        f"source-highlight/preview?{query}"
    )


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------


@router.post(
    "/projects/{project_id}/source-highlight/preview"
)
def source_highlight_preview_post(
    project_id: int,
    payload: SourceHighlightRequest = Body(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    project = get_project_for_user(
        db,
        project_id,
        current_user,
    )

    if payload.return_json:
        return JSONResponse(
            {
                "ok": True,
                "preview_url": build_preview_url(
                    project_id,
                    payload,
                ),
            }
        )

    path = resolve_document_path(
        db=db,
        project=project,
        payload=payload,
    )

    return render_document_preview(
        path,
        payload,
        project_id,
    )


@router.get(
    "/projects/{project_id}/source-highlight/preview"
)
def source_highlight_preview_get(
    project_id: int,
    request: Request,
    excerpt: str = Query(
        "",
        max_length=MAX_EXCERPT_CHARS,
    ),
    document_id: Optional[int] = Query(None),
    source_path: Optional[str] = Query(None),
    source_name: Optional[str] = Query(None),
    document_name: Optional[str] = Query(None),
    title: Optional[str] = Query(None),
    passage_id: Optional[str] = Query(None),
    page_number: Optional[int] = Query(None),
    paragraph_index: Optional[int] = Query(None),
    char_start: Optional[int] = Query(None),
    char_end: Optional[int] = Query(None),
    year: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    del request

    project = get_project_for_user(
        db,
        project_id,
        current_user,
    )

    payload = SourceHighlightRequest(
        excerpt=excerpt,
        document_id=document_id,
        source_path=source_path,
        source_name=source_name,
        document_name=document_name,
        title=title,
        passage_id=passage_id,
        page_number=page_number,
        paragraph_index=paragraph_index,
        char_start=char_start,
        char_end=char_end,
        year=year,
    )

    path = resolve_document_path(
        db=db,
        project=project,
        payload=payload,
    )

    return render_document_preview(
        path,
        payload,
        project_id,
    )


@router.get(
    "/projects/{project_id}/source-highlight/health"
)
def source_highlight_health(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    project = get_project_for_user(
        db,
        project_id,
        current_user,
    )

    try:
        import fitz  # type: ignore

        pymupdf_ok = True
    except Exception:
        pymupdf_ok = False

    libreoffice = find_libreoffice_binary()

    try:
        import docx  # type: ignore

        python_docx_ok = True
    except Exception:
        python_docx_ok = False

    return {
        "ok": True,
        "version": "V170-office-pdf",
        "project_id": project.id,
        "project_dir": str(
            get_project_store(project).project_dir
        ),
        "preview_root": str(PREVIEW_ROOT),
        "office_pdf_cache": str(
            OFFICE_PDF_CACHE
        ),
        "pymupdf_ok": pymupdf_ok,
        "libreoffice_ok": bool(libreoffice),
        "libreoffice_binary": (
            libreoffice or ""
        ),
        "python_docx_ok": python_docx_ok,
        "document_model_available": (
            Document is not None
        ),
        "office_formats": sorted(
            OFFICE_TO_PDF_EXTENSIONS
        ),
    }


# ---------------------------------------------------------------------------
# Compatibilité CIR source view
# ---------------------------------------------------------------------------


def _pick_first(
    payload: Dict[str, Any],
    keys: Iterable[str],
    default: str = "",
) -> str:
    for key in keys:
        value = payload.get(key)

        if (
            value is not None
            and str(value).strip()
        ):
            return str(value)

    return default


def _payload_to_request(
    payload: Dict[str, Any],
) -> SourceHighlightRequest:
    current = (
        payload.get("current_item")
        if isinstance(
            payload.get("current_item"),
            dict,
        )
        else {}
    )

    previous = (
        payload.get("previous_candidate")
        if isinstance(
            payload.get("previous_candidate"),
            dict,
        )
        else {}
    )

    best_match = (
        payload.get("best_match")
        if isinstance(
            payload.get("best_match"),
            dict,
        )
        else {}
    )

    best_previous = (
        best_match.get("previous_candidate")
        if isinstance(
            best_match.get("previous_candidate"),
            dict,
        )
        else {}
    )

    excerpt = _pick_first(
        payload,
        [
            "excerpt",
            "passage",
            "text",
            "highlight",
            "selected_text",
            "source_excerpt",
            "content",
        ],
    )

    if not excerpt:
        excerpt = _pick_first(
            current,
            [
                "text",
                "source_text",
                "excerpt",
                "description",
                "title",
            ],
        )

    if not excerpt:
        excerpt = _pick_first(
            previous,
            [
                "text",
                "source_text",
                "excerpt",
                "description",
                "title",
            ],
        )

    if not excerpt:
        excerpt = _pick_first(
            best_previous,
            [
                "text",
                "source_text",
                "excerpt",
                "description",
                "title",
            ],
        )

    source_path = _pick_first(
        payload,
        [
            "source_path",
            "path",
            "file_path",
            "document_path",
            "sourceFilePath",
            "filePath",
        ],
    )

    if not source_path:
        source_path = _pick_first(
            current,
            [
                "source_path",
                "path",
                "file_path",
                "document_path",
            ],
        )

    if not source_path:
        source_path = _pick_first(
            previous,
            [
                "source_path",
                "path",
                "file_path",
                "document_path",
            ],
        )

    if not source_path:
        source_path = _pick_first(
            best_previous,
            [
                "source_path",
                "path",
                "file_path",
                "document_path",
            ],
        )

    source_name = _pick_first(
        payload,
        [
            "source_name",
            "document_name",
            "filename",
            "file_name",
            "name",
            "source",
            "document",
        ],
    )

    if not source_name:
        source_name = _pick_first(
            current,
            [
                "source_name",
                "document_name",
                "filename",
                "file_name",
                "source",
                "document",
            ],
        )

    if not source_name:
        source_name = _pick_first(
            previous,
            [
                "source_name",
                "document_name",
                "filename",
                "file_name",
                "source",
                "document",
            ],
        )

    if not source_name:
        source_name = _pick_first(
            best_previous,
            [
                "source_name",
                "document_name",
                "filename",
                "file_name",
                "source",
                "document",
            ],
        )

    document_id_raw = (
        payload.get("document_id")
        or current.get("document_id")
        or previous.get("document_id")
        or best_previous.get("document_id")
    )

    try:
        document_id = (
            int(document_id_raw)
            if document_id_raw not in (None, "")
            else None
        )
    except Exception:
        document_id = None

    return SourceHighlightRequest(
        excerpt=excerpt,
        document_id=document_id,
        source_path=(
            source_path or None
        ),
        source_name=(
            source_name or None
        ),
        document_name=(
            _pick_first(
                payload,
                [
                    "document_name",
                    "document",
                    "filename",
                    "file_name",
                ],
            )
            or None
        ),
        title=(
            _pick_first(
                payload,
                [
                    "title",
                    "card_title",
                    "label",
                ],
                "Prévisualisation source",
            )
            or None
        ),
        role=(
            _pick_first(
                payload,
                [
                    "role",
                    "type",
                    "category",
                ],
            )
            or None
        ),
        passage_id=(
            _pick_first(
                payload,
                [
                    "passage_id",
                    "rag_chunk_id",
                    "evidence_id",
                ],
            )
            or None
        ),
        page_number=payload.get(
            "page_number"
        ),
        paragraph_index=payload.get(
            "paragraph_index"
        ),
        char_start=payload.get(
            "char_start"
        ),
        char_end=payload.get(
            "char_end"
        ),
        year=(
            _pick_first(
                payload,
                ["year", "previous_year"],
            )
            or None
        ),
        return_json=bool(
            payload.get(
                "return_json",
                True,
            )
        ),
    )


@router.post(
    "/projects/{project_id}/cir-source-view/open-passage"
)
def cir_source_view_open_passage(
    project_id: int,
    payload: Dict[str, Any] = Body(
        default_factory=dict
    ),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    project = get_project_for_user(
        db,
        project_id,
        current_user,
    )

    request_payload = _payload_to_request(
        payload or {}
    )

    if (
        bool(payload.get("return_file", False))
        or bool(payload.get("direct", False))
    ):
        path = resolve_document_path(
            db=db,
            project=project,
            payload=request_payload,
        )

        return render_document_preview(
            path,
            request_payload,
            project_id,
        )

    preview_url = build_preview_url(
        project_id,
        request_payload,
    )

    return {
        "ok": True,
        "project_id": project_id,
        "preview_url": preview_url,
        "open_url": preview_url,
        "url": preview_url,
        "display_url": preview_url,
        "method": "GET",
    }


@router.get(
    "/projects/{project_id}/cir-source-view/open-passage"
)
def cir_source_view_open_passage_get(
    project_id: int,
    excerpt: str = Query(
        "",
        max_length=MAX_EXCERPT_CHARS,
    ),
    document_id: Optional[int] = Query(None),
    source_path: Optional[str] = Query(None),
    source_name: Optional[str] = Query(None),
    document_name: Optional[str] = Query(None),
    title: Optional[str] = Query(None),
    page_number: Optional[int] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    project = get_project_for_user(
        db,
        project_id,
        current_user,
    )

    payload = SourceHighlightRequest(
        excerpt=excerpt,
        document_id=document_id,
        source_path=source_path,
        source_name=source_name,
        document_name=document_name,
        title=title,
        page_number=page_number,
    )

    path = resolve_document_path(
        db=db,
        project=project,
        payload=payload,
    )

    return render_document_preview(
        path,
        payload,
        project_id,
    )


@router.get(
    "/projects/{project_id}/cir-source-view/health"
)
def cir_source_view_health(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    return source_highlight_health(
        project_id,
        db,
        current_user,
    )


# ---------------------------------------------------------------------------
# CIR précédent
# ---------------------------------------------------------------------------


def _find_previous_cir_file(
    project: Any,
    year: str,
) -> Optional[Path]:
    project_store = get_project_store(project)

    current_project_dir = Path(
        project_store.project_dir
    )

    years_root = current_project_dir.parent
    target_year_root = (
        years_root / str(year)
    )

    candidates = [
        (
            target_year_root
            / "cir_final_consultant"
            / "current"
        ),
        (
            target_year_root
            / "cir_final"
            / "raw"
        ),
        (
            target_year_root
            / "cir_final"
        ),
        target_year_root,
    ]

    extensions = sorted(
        {
            f"*{ext}"
            for ext in (
                PDF_EXTENSIONS
                | WORD_EXTENSIONS
                | PRESENTATION_EXTENSIONS
                | TEXT_EXTENSIONS
            )
        }
    )

    for base in candidates:
        if not base.exists():
            continue

        files: List[Path] = []

        for extension in extensions:
            files.extend(
                base.glob(extension)
            )

        files = [
            path
            for path in files
            if path.is_file()
        ]

        if files:
            return max(
                files,
                key=lambda path: path.stat().st_mtime,
            )

    return None


def _render_previous_memory(
    project: Any,
    year: str,
    highlight: str,
) -> HTMLResponse:
    project_store = get_project_store(project)

    target_year_root = (
        Path(project_store.project_dir).parent
        / str(year)
    )

    paths = [
        target_year_root / "cir_final_memory.json",
        (
            target_year_root
            / "cir_final_consultant"
            / "current"
            / "cir_final_memory.json"
        ),
    ]

    for path in paths:
        if not path.exists():
            continue

        try:
            data = json.loads(
                path.read_text(
                    encoding="utf-8",
                    errors="ignore",
                )
            )
        except Exception:
            continue

        sections = (
            data.get("sections_full")
            or data.get("sections")
            or {}
        )

        text = "\n\n".join(
            f"{key}\n{value}"
            for key, value in sections.items()
            if (
                isinstance(value, str)
                and value.strip()
            )
        )

        body, exact = highlight_html_text(
            text,
            highlight,
        )

        return HTMLResponse(
            html_document(
                f"CIR {year} — mémoire extraite",
                body,
                str(path),
                exact,
            )
        )

    raise HTTPException(
        status_code=404,
        detail=(
            "Aucun CIR précédent trouvé "
            "pour cette année."
        ),
    )


@router.get(
    "/projects/{project_id}/cir-previous/preview"
)
def preview_previous_cir(
    project_id: int,
    year: str = Query(...),
    highlight: str = Query(
        "",
        max_length=MAX_EXCERPT_CHARS,
    ),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    project = get_project_for_user(
        db,
        project_id,
        current_user,
    )

    file_path = _find_previous_cir_file(
        project,
        year,
    )

    if file_path:
        payload = SourceHighlightRequest(
            excerpt=highlight,
            source_path=str(file_path),
            title=f"CIR précédent {year}",
        )

        return render_document_preview(
            file_path,
            payload,
            project_id,
        )

    return _render_previous_memory(
        project,
        year,
        highlight,
    )
