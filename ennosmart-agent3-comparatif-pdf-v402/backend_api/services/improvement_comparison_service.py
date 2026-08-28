from __future__ import annotations

"""
EnnoSmart Agent 3 — comparatif documentaire V4.02.

- document original fidèle, Office -> PDF via source_highlight ;
- proposition DOCX reconstruite sur une copie du document source ;
- suppression/remplacement en rouge côté original ;
- ajout/modification en vert côté proposition ;
- page de la modification renvoyée au frontend pour auto-positionnement.
"""

import hashlib
import json
import os
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from fastapi import HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session, undefer

from db.models import Document, ImprovementSession, ImprovementVersion

try:
    from routers.source_highlight import (
        OFFICE_TO_PDF_EXTENSIONS,
        PDF_EXTENSIONS,
        convert_office_to_pdf,
        make_search_queries,
        normalize_text,
    )
except Exception as exc:  # pragma: no cover
    raise RuntimeError(
        "Le comparatif Agent 3 nécessite source_highlight V171/V172 "
        "avec convert_office_to_pdf()."
    ) from exc


BACKEND_ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = BACKEND_ROOT.parent
STORAGE_ROOT = Path(
    os.getenv("ENNOSMART_STORAGE_ROOT") or PROJECT_ROOT / "storage"
).resolve()
COMPARE_ROOT = STORAGE_ROOT / "previews" / "improvement_comparison"
CANONICAL_SOURCE_ROOT = COMPARE_ROOT / "canonical_source"
CANONICAL_PDF_ROOT = COMPARE_ROOT / "canonical_pdf"

DOCX_EXTENSIONS = {".docx", ".docm"}
TEXT_EXTENSIONS = {".txt", ".md"}

WORD_TO_PDF_TIMEOUT_SECONDS = int(
    os.getenv("ENNOSMART_WORD_CONVERT_TIMEOUT", "240")
)

RED = (0.90, 0.12, 0.18)
RED_FILL = (1.0, 0.88, 0.89)
GREEN = (0.05, 0.58, 0.32)
GREEN_FILL = (0.86, 0.98, 0.91)
AMBER = (0.93, 0.48, 0.06)
AMBER_FILL = (1.0, 0.95, 0.82)


@dataclass
class ComparisonPreview:
    path: Path
    page: int
    matched: bool
    mode: str


def _ensure_root() -> None:
    for root in (
        COMPARE_ROOT,
        CANONICAL_SOURCE_ROOT,
        CANONICAL_PDF_ROOT,
    ):
        root.mkdir(parents=True, exist_ok=True)


def _hash(*parts: Any) -> str:
    raw = "||".join(str(part or "") for part in parts)
    return hashlib.sha256(
        raw.encode("utf-8", errors="ignore")
    ).hexdigest()[:24]


def _clean(value: Any) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(
        r"\[BLOC DOCUMENT IMMUTABLE[\s\S]*?\[/BLOC DOCUMENT IMMUTABLE\]",
        " ",
        text,
        flags=re.I,
    )
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _session(
    db: Session,
    project_id: int,
    session_id: str,
) -> ImprovementSession:
    row = (
        db.query(ImprovementSession)
        .filter(
            ImprovementSession.id == session_id,
            ImprovementSession.project_id == project_id,
        )
        .first()
    )
    if row is None:
        raise HTTPException(
            status_code=404,
            detail="Conversation EnnoAmelioration introuvable.",
        )
    return row


def _version(
    db: Session,
    session_id: str,
    version_id: str,
) -> ImprovementVersion:
    row = (
        db.query(ImprovementVersion)
        .filter(
            ImprovementVersion.id == version_id,
            ImprovementVersion.session_id == session_id,
        )
        .first()
    )
    if row is None:
        raise HTTPException(
            status_code=404,
            detail="Version d'amélioration introuvable.",
        )
    return row


def _source_document(
    db: Session,
    project_id: int,
    document_id: int,
) -> Document:
    query = db.query(Document)
    try:
        query = query.options(undefer(Document.file_data))
    except Exception:
        pass

    row = (
        query.filter(
            Document.id == document_id,
            Document.project_id == project_id,
        )
        .first()
    )
    if row is None:
        raise HTTPException(
            status_code=404,
            detail="Document source de la conversation introuvable.",
        )
    return row


def _safe_filename(value: Any) -> str:
    name = Path(str(value or "document").replace("\\", "/")).name
    name = re.sub(r'[\\/:*?"<>|]+', "_", name).strip(" .")
    return name or "document"


def _document_identity(document: Document) -> str:
    """Clé stable indépendante d'un chemin temporaire."""
    sha = str(getattr(document, "file_sha256", None) or "").strip().lower()
    if sha:
        return sha
    return _hash(
        "document",
        getattr(document, "id", ""),
        getattr(document, "file_size", ""),
        getattr(document, "filename", ""),
    )


def _read_exact_document_bytes(document: Document) -> bytes:
    """
    Source de vérité : le Document lié par ImprovementSession.source_document_id.
    Pour les nouveaux uploads, file_data/PostgreSQL est prioritaire.
    file_path n'est qu'un fallback legacy.
    """
    raw = bytes(getattr(document, "file_data", None) or b"")
    if raw:
        return raw

    file_path = str(getattr(document, "file_path", "") or "").strip()
    if file_path and not file_path.startswith("db://"):
        candidate = Path(file_path)
        if candidate.exists() and candidate.is_file():
            return candidate.read_bytes()

    raise HTTPException(
        status_code=404,
        detail=(
            "Le CIR est bien rattaché à la conversation, mais son contenu "
            "binaire n'est plus disponible."
        ),
    )


def _materialize_document(document: Document, project_id: int) -> Path:
    """
    Matérialise EXACTEMENT le CIR sélectionné.
    Aucun scan par nom de fichier.
    """
    _ensure_root()

    raw = _read_exact_document_bytes(document)
    identity = _document_identity(document)
    filename = _safe_filename(
        getattr(document, "filename", None)
        or getattr(document, "stored_filename", None)
        or f"document_{document.id}.bin"
    )

    output = (
        CANONICAL_SOURCE_ROOT
        / str(project_id)
        / str(document.id)
        / identity[:24]
        / filename
    )
    output.parent.mkdir(parents=True, exist_ok=True)

    if not output.exists() or output.stat().st_size != len(raw):
        staging = output.with_suffix(output.suffix + ".staging")
        staging.write_bytes(raw)
        os.replace(staging, output)

    return output.resolve()


def _pdf_from_text(source: Path, project_id: int) -> Path:
    try:
        import fitz  # type: ignore
    except Exception as exc:
        raise RuntimeError("PyMuPDF est requis pour le comparatif Agent 3.") from exc

    text = source.read_text(encoding="utf-8", errors="ignore")
    stat = source.stat()
    output = (
        COMPARE_ROOT
        / "text_pdf"
        / str(project_id)
        / f"{_hash(source, stat.st_mtime_ns)}.pdf"
    )
    if output.exists() and output.stat().st_size > 0:
        return output.resolve()

    output.parent.mkdir(parents=True, exist_ok=True)
    document = fitz.open()
    try:
        page = document.new_page(width=595, height=842)
        y = 48.0
        for raw_line in text.splitlines():
            if not raw_line.strip():
                y += 8
                continue
            remaining = raw_line
            while remaining:
                chunk = remaining[:100]
                remaining = remaining[100:]
                if y > 790:
                    page = document.new_page(width=595, height=842)
                    y = 48.0
                page.insert_text((48, y), chunk, fontsize=9.5)
                y += 14
        document.save(str(output), deflate=True)
    finally:
        document.close()
    return output.resolve()



def _valid_pdf(path: Path) -> bool:
    try:
        if not path.exists() or not path.is_file() or path.stat().st_size <= 5:
            return False
        with path.open("rb") as stream:
            return stream.read(5) == b"%PDF-"
    except Exception:
        return False


def _powershell_executable() -> str:
    for candidate in ("powershell.exe", "powershell", "pwsh.exe", "pwsh"):
        resolved = shutil.which(candidate)
        if resolved:
            return resolved
    raise RuntimeError("PowerShell est introuvable.")


def _convert_with_microsoft_word(source: Path, target: Path) -> Path:
    """
    Windows : utilise Microsoft Word COM pour préserver au mieux un gros DOCX.
    Aucun pywin32 requis.
    """
    if os.name != "nt":
        raise RuntimeError("Microsoft Word COM est disponible uniquement sous Windows.")

    target.parent.mkdir(parents=True, exist_ok=True)

    ps = _powershell_executable()
    source_ps = str(source.resolve()).replace("'", "''")
    target_ps = str(target.resolve()).replace("'", "''")

    script = f"""
$ErrorActionPreference = 'Stop'
$word = $null
$doc = $null
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    try {{ $word.AutomationSecurity = 3 }} catch {{}}
    try {{ $word.Options.UpdateLinksAtOpen = $false }} catch {{}}

    $doc = $word.Documents.Open('{source_ps}', $false, $true)
    $doc.ExportAsFixedFormat('{target_ps}', 17)
}}
finally {{
    if ($doc -ne $null) {{
        $doc.Close($false)
        [void][System.Runtime.InteropServices.Marshal]::ReleaseComObject($doc)
    }}
    if ($word -ne $null) {{
        $word.Quit()
        [void][System.Runtime.InteropServices.Marshal]::ReleaseComObject($word)
    }}
    [GC]::Collect()
    [GC]::WaitForPendingFinalizers()
}}
"""

    completed = subprocess.run(
        [
            ps,
            "-NoProfile",
            "-NonInteractive",
            "-ExecutionPolicy",
            "Bypass",
            "-Command",
            script,
        ],
        capture_output=True,
        text=True,
        timeout=max(60, WORD_TO_PDF_TIMEOUT_SECONDS),
        creationflags=(
            subprocess.CREATE_NO_WINDOW
            if os.name == "nt" and hasattr(subprocess, "CREATE_NO_WINDOW")
            else 0
        ),
    )

    if completed.returncode != 0:
        raise RuntimeError(
            "Microsoft Word n'a pas pu convertir le document. "
            f"stderr={completed.stderr[-1200:]!r}"
        )

    if not _valid_pdf(target):
        raise RuntimeError("Microsoft Word n'a pas produit de PDF valide.")

    return target.resolve()


def _stable_word_pdf(
    source: Path,
    project_id: int,
    cache_key: str,
    *,
    target_stem: str | None = None,
) -> tuple[Path, str]:
    """
    Word Windows en priorité, LibreOffice en fallback.
    Le cache dépend d'une clé stable, pas du chemin temporaire.
    """
    target = (
        CANONICAL_PDF_ROOT
        / str(project_id)
        / cache_key[:24]
        / f"{target_stem or source.stem}.pdf"
    )

    if _valid_pdf(target):
        return target.resolve(), "canonical-cache"

    target.parent.mkdir(parents=True, exist_ok=True)
    word_error = None

    if os.name == "nt":
        try:
            return _convert_with_microsoft_word(source, target), "microsoft-word"
        except Exception as exc:
            word_error = str(exc)

    try:
        lo_pdf = convert_office_to_pdf(source, project_id).resolve()
        if not _valid_pdf(lo_pdf):
            raise RuntimeError("LibreOffice n'a pas produit de PDF valide.")

        staging = target.with_suffix(".staging.pdf")
        shutil.copy2(lo_pdf, staging)
        os.replace(staging, target)
        return target.resolve(), "libreoffice"

    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=(
                "Conversion PDF impossible pour le CIR sélectionné. "
                f"Microsoft Word: {word_error or 'non tenté / indisponible'}. "
                f"LibreOffice: {exc}"
            ),
        ) from exc

def _as_pdf(
    source: Path,
    project_id: int,
    document: Document | None = None,
) -> Path:
    suffix = source.suffix.lower()

    # PDF sélectionné par le consultant : aucune conversion Office.
    if suffix in PDF_EXTENSIONS:
        if not _valid_pdf(source):
            raise HTTPException(
                status_code=422,
                detail="Le CIR sélectionné est déclaré PDF mais son contenu est invalide.",
            )
        return source.resolve()

    if suffix in OFFICE_TO_PDF_EXTENSIONS:
        cache_key = (
            _document_identity(document)
            if document is not None
            else _hash(source.resolve(), source.stat().st_size, source.stat().st_mtime_ns)
        )
        pdf, _mode = _stable_word_pdf(
            source,
            project_id,
            cache_key,
            target_stem=source.stem,
        )
        return pdf.resolve()

    if suffix in TEXT_EXTENSIONS:
        return _pdf_from_text(source, project_id)

    raise HTTPException(
        status_code=415,
        detail=f"Format non pris en charge pour le comparatif PDF : {suffix or 'inconnu'}.",
    )


def _structured_result(version: ImprovementVersion) -> dict[str, Any]:
    generation = (
        version.generation_json
        if isinstance(version.generation_json, dict)
        else {}
    )
    for key in ("structured_result", "trace"):
        value = generation.get(key)
        if isinstance(value, dict):
            return value
    return {}


def _changes(version: ImprovementVersion) -> list[dict[str, Any]]:
    structured = _structured_result(version)
    rows = structured.get("changes")
    if isinstance(rows, list) and rows:
        return [row for row in rows if isinstance(row, dict)]

    diff = version.diff_json if isinstance(version.diff_json, dict) else {}
    rows = diff.get("changes")
    if not isinstance(rows, list):
        return []
    return [row for row in rows if isinstance(row, dict)]


def _change_at(version: ImprovementVersion, index: int) -> dict[str, Any]:
    rows = _changes(version)
    if index < 0 or index >= len(rows):
        raise HTTPException(
            status_code=404,
            detail="Modification introuvable dans cette proposition.",
        )
    return rows[index]


def _change_before(change: dict[str, Any]) -> str:
    return _clean(change.get("before"))


def _change_after(change: dict[str, Any]) -> str:
    return _clean(change.get("after"))


def _section_hints(change: dict[str, Any]) -> list[str]:
    values = [
        change.get("section_title"),
        change.get("section_ref"),
        change.get("label"),
    ]
    return [_clean(value) for value in values if _clean(value)]


def _search_queries(text: str) -> list[str]:
    value = _clean(text)
    if not value:
        return []

    try:
        rows = list(make_search_queries(value))
    except Exception:
        rows = []

    output: list[str] = []
    seen: set[str] = set()

    for row in [value[:350], *rows]:
        row = re.sub(r"\s+", " ", str(row or "")).strip()
        if len(row) < 4:
            continue
        key = normalize_text(row)
        if key and key not in seen:
            seen.add(key)
            output.append(row)

    words = value.split()
    for size in (14, 10, 8, 6, 4):
        if len(words) < size:
            continue
        step = max(1, size // 2)
        for start in range(
            0,
            min(len(words) - size + 1, 24),
            step,
        ):
            row = " ".join(words[start : start + size])
            key = normalize_text(row)
            if key and key not in seen:
                seen.add(key)
                output.append(row)

    return output[:60]


def _find_rectangles(
    pdf: Any,
    texts: Iterable[str],
) -> tuple[int, list[Any], bool]:
    queries: list[str] = []
    for text in texts:
        queries.extend(_search_queries(text))

    for page_index, page in enumerate(pdf):
        for query in queries:
            try:
                rectangles = page.search_for(query)
            except Exception:
                rectangles = []
            if rectangles:
                return page_index, rectangles[:20], True
    return 0, [], False


def _best_anchor(
    pdf: Any,
    change: dict[str, Any],
    *,
    prefer_after: bool,
) -> tuple[int, list[Any], bool]:
    primary = _change_after(change) if prefer_after else _change_before(change)
    secondary = _change_before(change) if prefer_after else _change_after(change)
    texts = [primary, secondary, *_section_hints(change)]
    return _find_rectangles(pdf, [text for text in texts if text])


def _add_label(
    page: Any,
    text: str,
    color: tuple[float, float, float],
    fill: tuple[float, float, float],
) -> None:
    try:
        rect = page.rect
        box = (
            rect.x1 - 185,
            rect.y0 + 18,
            rect.x1 - 22,
            rect.y0 + 58,
        )
        annot = page.add_freetext_annot(
            box,
            text,
            fontsize=8,
            text_color=color,
            fill_color=fill,
            border_color=color,
            align=1,
        )
        annot.set_border(width=1)
        annot.update()
    except Exception:
        pass


def _mark_red(page: Any, rectangles: list[Any]) -> None:
    for rectangle in rectangles[:20]:
        try:
            highlight = page.add_highlight_annot(rectangle)
            highlight.set_colors(stroke=RED)
            highlight.set_opacity(0.40)
            highlight.update()
        except Exception:
            pass

        try:
            strike = page.add_strikeout_annot(rectangle)
            strike.set_colors(stroke=RED)
            strike.update()
        except Exception:
            pass

        try:
            border = page.add_rect_annot(rectangle)
            border.set_colors(stroke=RED)
            border.set_border(width=1.2)
            border.update()
        except Exception:
            pass


def _mark_green(page: Any, rectangles: list[Any]) -> None:
    for rectangle in rectangles[:20]:
        try:
            highlight = page.add_highlight_annot(rectangle)
            highlight.set_colors(stroke=GREEN)
            highlight.set_opacity(0.34)
            highlight.update()
        except Exception:
            pass

        try:
            border = page.add_rect_annot(rectangle)
            border.set_colors(stroke=GREEN)
            border.set_border(width=1.25)
            border.update()
        except Exception:
            pass


def _paragraphs_in_document(document: Any) -> list[Any]:
    output = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                output.extend(cell.paragraphs)
    return output


def _paragraph_contains_drawing(paragraph: Any) -> bool:
    try:
        xml = paragraph._p.xml
        return (
            "<w:drawing" in xml
            or "<w:pict" in xml
            or "<a:graphic" in xml
        )
    except Exception:
        return False


def _normalized(value: Any) -> str:
    return normalize_text(_clean(value))


def _replace_single_paragraph(
    paragraph: Any,
    before: str,
    after: str,
) -> bool:
    current = str(paragraph.text or "")
    if not current.strip() or _paragraph_contains_drawing(paragraph):
        return False

    if before in current:
        paragraph.text = current.replace(before, after, 1)
        return True

    wanted = _normalized(before)
    current_normalized = _normalized(current)

    if wanted and current_normalized == wanted:
        paragraph.text = after
        return True

    if wanted and len(wanted) >= 18 and wanted in current_normalized:
        paragraph.text = after
        return True

    return False


def _replace_multi_paragraph(
    paragraphs: list[Any],
    before: str,
    after: str,
) -> bool:
    wanted = _normalized(before)
    if len(wanted) < 30:
        return False

    for start in range(len(paragraphs)):
        combined: list[str] = []
        indices: list[int] = []

        for end in range(start, min(len(paragraphs), start + 8)):
            paragraph = paragraphs[end]
            if _paragraph_contains_drawing(paragraph):
                break

            value = str(paragraph.text or "").strip()
            if not value:
                continue

            combined.append(value)
            indices.append(end)
            normalized_combined = _normalized(" ".join(combined))

            if normalized_combined == wanted:
                paragraphs[indices[0]].text = after
                for index in indices[1:]:
                    paragraphs[index].text = ""
                return True

            if len(normalized_combined) > len(wanted) * 1.35:
                break

    return False


def _insert_after_paragraph(paragraph: Any, text: str) -> bool:
    try:
        from docx.oxml import OxmlElement  # type: ignore
        from docx.text.paragraph import Paragraph  # type: ignore

        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_paragraph = Paragraph(new_p, paragraph._parent)
        new_paragraph.style = paragraph.style
        new_paragraph.add_run(text)
        return True
    except Exception:
        return False


def _apply_change_to_docx(document: Any, change: dict[str, Any]) -> bool:
    before = _change_before(change)
    after = _change_after(change)

    if before:
        paragraphs = _paragraphs_in_document(document)
        for paragraph in paragraphs:
            if _replace_single_paragraph(paragraph, before, after):
                return True
        return _replace_multi_paragraph(paragraphs, before, after)

    if after:
        hints = _section_hints(change)
        paragraphs = _paragraphs_in_document(document)
        for hint in hints:
            hint_norm = _normalized(hint)
            if not hint_norm:
                continue
            for paragraph in paragraphs:
                text_norm = _normalized(paragraph.text)
                if text_norm and (
                    text_norm == hint_norm or hint_norm in text_norm
                ):
                    return _insert_after_paragraph(paragraph, after)
    return False


def _candidate_docx(
    source: Path,
    version: ImprovementVersion,
    project_id: int,
) -> tuple[Path, dict[str, Any]]:
    try:
        from docx import Document as DocxDocument  # type: ignore
    except Exception as exc:
        raise RuntimeError(
            "python-docx est requis pour créer l'aperçu de la proposition."
        ) from exc

    stat = source.stat()
    rows = _changes(version)
    key = _hash(
        source.resolve(),
        stat.st_size,
        stat.st_mtime_ns,
        version.id,
        json.dumps(rows, ensure_ascii=False, sort_keys=True, default=str),
    )

    directory = (
        COMPARE_ROOT
        / "candidate_docx"
        / str(project_id)
        / str(version.id)
        / key
    )
    directory.mkdir(parents=True, exist_ok=True)
    target = directory / f"{source.stem}_proposition{source.suffix}"
    manifest_path = directory / "manifest.json"

    if target.exists() and target.stat().st_size > 0 and manifest_path.exists():
        try:
            return target.resolve(), json.loads(
                manifest_path.read_text(encoding="utf-8")
            )
        except Exception:
            pass

    document = DocxDocument(str(source))
    applied: list[int] = []
    missed: list[int] = []

    for index, change in enumerate(rows):
        if _apply_change_to_docx(document, change):
            applied.append(index)
        else:
            missed.append(index)

    document.save(str(target))
    manifest = {
        "version": "agent3-docx-patch-v402",
        "source": str(source),
        "version_id": version.id,
        "changes_total": len(rows),
        "applied": applied,
        "missed": missed,
        "applied_count": len(applied),
        "missed_count": len(missed),
    }
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return target.resolve(), manifest


def _base_proposed_pdf(
    source: Path,
    original_pdf: Path,
    version: ImprovementVersion,
    project_id: int,
) -> tuple[Path, str, dict[str, Any]]:
    if source.suffix.lower() in DOCX_EXTENSIONS:
        try:
            candidate_docx, manifest = _candidate_docx(
                source, version, project_id
            )
            candidate_key = _hash(
                "candidate",
                version.id,
                candidate_docx.stat().st_size,
                candidate_docx.stat().st_mtime_ns,
            )
            candidate_pdf, candidate_mode = _stable_word_pdf(
                candidate_docx,
                project_id,
                candidate_key,
                target_stem=f"{candidate_docx.stem}_{version.id[:8]}",
            )
            return candidate_pdf, candidate_mode, manifest
        except Exception as exc:
            return (
                original_pdf,
                "review-overlay",
                {
                    "fallback_error": str(exc),
                    "applied": [],
                    "missed": list(range(len(_changes(version)))),
                },
            )

    return (
        original_pdf,
        (
            "review-overlay-pdf"
            if source.suffix.lower() in PDF_EXTENSIONS
            else "review-overlay-office"
        ),
        {
            "source_format": source.suffix.lower().lstrip("."),
            "applied": [],
            "missed": list(range(len(_changes(version)))),
        },
    )


def _original_preview(
    original_pdf: Path,
    version: ImprovementVersion,
    change: dict[str, Any],
    change_index: int,
    project_id: int,
) -> ComparisonPreview:
    try:
        import fitz  # type: ignore
    except Exception as exc:
        raise RuntimeError("PyMuPDF est requis pour le comparatif Agent 3.") from exc

    stat = original_pdf.stat()
    target = (
        COMPARE_ROOT
        / "rendered"
        / str(project_id)
        / str(version.id)
        / f"original_{change_index}_{_hash(original_pdf, stat.st_mtime_ns, change)}.pdf"
    )
    metadata_path = target.with_suffix(".json")

    if target.exists() and target.stat().st_size > 0 and metadata_path.exists():
        try:
            metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
            return ComparisonPreview(
                path=target.resolve(),
                page=int(metadata.get("page") or 1),
                matched=bool(metadata.get("matched")),
                mode="original",
            )
        except Exception:
            pass

    target.parent.mkdir(parents=True, exist_ok=True)
    pdf = fitz.open(str(original_pdf))
    try:
        before = _change_before(change)
        page_index, rectangles, matched = _best_anchor(
            pdf, change, prefer_after=False
        )

        if before and matched and rectangles:
            _mark_red(pdf[page_index], rectangles)
            _add_label(
                pdf[page_index],
                "SUPPRIMÉ / REMPLACÉ",
                RED,
                RED_FILL,
            )
        else:
            _add_label(
                pdf[page_index],
                "EMPLACEMENT DE L'AJOUT",
                AMBER,
                AMBER_FILL,
            )

        pdf.save(str(target), garbage=4, deflate=True)
    finally:
        pdf.close()

    metadata = {
        "page": page_index + 1,
        "matched": matched,
        "side": "original",
    }
    metadata_path.write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return ComparisonPreview(
        path=target.resolve(),
        page=page_index + 1,
        matched=matched,
        mode="original",
    )


def _green_overlay_note(page: Any, text: str) -> None:
    try:
        rect = page.rect
        box = (
            rect.x0 + 35,
            rect.y0 + 70,
            rect.x1 - 35,
            min(rect.y0 + 225, rect.y1 - 35),
        )
        note = page.add_freetext_annot(
            box,
            "PROPOSITION\n\n" + text[:1800],
            fontsize=8.2,
            text_color=GREEN,
            fill_color=GREEN_FILL,
            border_color=GREEN,
            align=0,
        )
        note.set_border(width=1.3)
        note.update()
    except Exception:
        pass


def _proposed_preview(
    source: Path,
    original_pdf: Path,
    version: ImprovementVersion,
    change: dict[str, Any],
    change_index: int,
    project_id: int,
) -> ComparisonPreview:
    try:
        import fitz  # type: ignore
    except Exception as exc:
        raise RuntimeError("PyMuPDF est requis pour le comparatif Agent 3.") from exc

    proposed_pdf, mode, manifest = _base_proposed_pdf(
        source, original_pdf, version, project_id
    )
    stat = proposed_pdf.stat()
    target = (
        COMPARE_ROOT
        / "rendered"
        / str(project_id)
        / str(version.id)
        / f"proposed_{change_index}_{_hash(proposed_pdf, stat.st_mtime_ns, change, mode)}.pdf"
    )
    metadata_path = target.with_suffix(".json")

    if target.exists() and target.stat().st_size > 0 and metadata_path.exists():
        try:
            metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
            return ComparisonPreview(
                path=target.resolve(),
                page=int(metadata.get("page") or 1),
                matched=bool(metadata.get("matched")),
                mode=str(metadata.get("mode") or mode),
            )
        except Exception:
            pass

    target.parent.mkdir(parents=True, exist_ok=True)
    pdf = fitz.open(str(proposed_pdf))
    try:
        after = _change_after(change)
        page_index, rectangles, matched = _best_anchor(
            pdf, change, prefer_after=True
        )

        if after and matched and rectangles:
            _mark_green(pdf[page_index], rectangles)
            _add_label(
                pdf[page_index],
                "AJOUTÉ / MODIFIÉ",
                GREEN,
                GREEN_FILL,
            )

        elif after:
            anchor_index, anchor_rects, anchor_found = _best_anchor(
                pdf, change, prefer_after=False
            )
            page_index = anchor_index
            matched = anchor_found
            page = pdf[page_index]

            if anchor_rects:
                _mark_green(page, anchor_rects)

            _green_overlay_note(page, after)
            _add_label(
                page,
                "AJOUTÉ / MODIFIÉ",
                GREEN,
                GREEN_FILL,
            )

        else:
            _add_label(
                pdf[page_index],
                "PASSAGE SUPPRIMÉ",
                GREEN,
                GREEN_FILL,
            )

        pdf.save(str(target), garbage=4, deflate=True)
    finally:
        pdf.close()

    metadata = {
        "page": page_index + 1,
        "matched": matched,
        "side": "proposed",
        "mode": mode,
        "docx_patch": manifest,
    }
    metadata_path.write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return ComparisonPreview(
        path=target.resolve(),
        page=page_index + 1,
        matched=matched,
        mode=mode,
    )


def build_comparison_preview(
    db: Session,
    *,
    project_id: int,
    session_id: str,
    version_id: str,
    side: str,
    change_index: int,
) -> ComparisonPreview:
    _ensure_root()

    session = _session(db, project_id, session_id)
    version = _version(db, session_id, version_id)

    if not session.source_document_id:
        raise HTTPException(
            status_code=400,
            detail=(
                "Cette conversation n'est pas rattachée à un document source. "
                "Le comparatif PDF est disponible pour les CIR importés."
            ),
        )

    source_document = _source_document(
        db,
        project_id,
        int(session.source_document_id),
    )
    source = _materialize_document(source_document, project_id)
    # V4.02 : le document vient directement de session.source_document_id.
    # Un PDF est utilisé tel quel ; un Word est converti via Word/LibreOffice.
    original_pdf = _as_pdf(
        source,
        project_id,
        source_document,
    )
    change = _change_at(version, change_index)

    normalized_side = str(side or "").strip().lower()
    if normalized_side == "original":
        return _original_preview(
            original_pdf,
            version,
            change,
            change_index,
            project_id,
        )
    if normalized_side == "proposed":
        return _proposed_preview(
            source,
            original_pdf,
            version,
            change,
            change_index,
            project_id,
        )

    raise HTTPException(
        status_code=400,
        detail="side doit être 'original' ou 'proposed'.",
    )


def comparison_file_response(preview: ComparisonPreview) -> FileResponse:
    headers = {
        "Content-Disposition": "inline; filename=comparatif.pdf",
        "Cache-Control": "no-store",
        "X-EnnoSmart-Comparison-Page": str(max(1, int(preview.page or 1))),
        "X-EnnoSmart-Comparison-Match": "true" if preview.matched else "false",
        "X-EnnoSmart-Comparison-Mode": preview.mode,
        "Access-Control-Expose-Headers": (
            "X-EnnoSmart-Comparison-Page, "
            "X-EnnoSmart-Comparison-Match, "
            "X-EnnoSmart-Comparison-Mode"
        ),
    }
    return FileResponse(
        preview.path,
        media_type="application/pdf",
        headers=headers,
    )
