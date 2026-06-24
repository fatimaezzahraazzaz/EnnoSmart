# -*- coding: utf-8 -*-
from __future__ import annotations

import html
import json
import re
import zipfile
from pathlib import Path
from typing import Optional
from urllib.parse import quote

from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import FileResponse, HTMLResponse, PlainTextResponse

router = APIRouter(tags=["source-preview"])

ROOT = Path(r"C:\EnnoSmart")
STORAGE_ROOT = ROOT / "storage"


def slugify(value: str) -> str:
    value = str(value or "").strip().lower()
    value = value.replace("\\", "/").split("/")[-1]
    value = re.sub(r"[^\w\-]+", "_", value, flags=re.UNICODE)
    value = re.sub(r"_+", "_", value).strip("_")
    return value or "unknown"


def is_inside(child: Path, parent: Path) -> bool:
    try:
        child.resolve().relative_to(parent.resolve())
        return True
    except Exception:
        return False


def safe_storage_path(value: str) -> Path:
    if not value:
        raise HTTPException(status_code=400, detail="source_path manquant.")

    path = Path(value)

    if not path.is_absolute():
        path = ROOT / value

    path = path.resolve()

    if not is_inside(path, STORAGE_ROOT):
        raise HTTPException(
            status_code=403,
            detail="Chemin refusé : la prévisualisation est limitée au dossier storage.",
        )

    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=404, detail=f"Fichier introuvable : {path}")

    return path


def normalize_spaces(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def highlight_html(text: str, highlight: str) -> str:
    safe_text = html.escape(text or "")
    needle = normalize_spaces(highlight)

    if len(needle) < 8:
        return safe_text.replace("\n", "<br/>")

    words = needle.split()
    probe = " ".join(words[: min(len(words), 14)])

    if len(probe) < 8:
        return safe_text.replace("\n", "<br/>")

    escaped_probe = html.escape(probe)
    pattern = re.compile(re.escape(escaped_probe), re.IGNORECASE)

    if pattern.search(safe_text):
        safe_text = pattern.sub(
            lambda m: f'<mark class="hit">{m.group(0)}</mark>',
            safe_text,
            count=1,
        )
        return safe_text.replace("\n", "<br/>")

    passage = html.escape(highlight or "")
    return (
        f'<div class="selected"><b>Passage sélectionné</b><br/>{passage}</div>'
        + safe_text.replace("\n", "<br/>")
    )


def html_page(title: str, body: str, source: str = "") -> str:
    return f"""<!doctype html>
<html lang="fr">
<head>
  <meta charset="utf-8"/>
  <title>{html.escape(title)}</title>
  <style>
    body {{ margin: 0; padding: 24px; font-family: Arial, Helvetica, sans-serif; background: #fbfbfc; color: #1f2937; line-height: 1.65; font-size: 14px; }}
    .top {{ position: sticky; top: 0; background: #fbfbfc; border-bottom: 1px solid #e5e7eb; padding-bottom: 12px; margin-bottom: 18px; z-index: 5; }}
    h1 {{ font-size: 16px; margin: 0 0 6px; }}
    .source {{ color: #6b7280; font-size: 12px; word-break: break-all; }}
    .doc {{ background: white; border: 1px solid #e5e7eb; border-radius: 12px; padding: 22px; box-shadow: 0 1px 2px rgba(0,0,0,.04); }}
    mark.hit {{ background: #fde68a; color: #111827; border-radius: 4px; padding: 2px 4px; }}
    .selected {{ background: #fff7cc; border: 1px solid #facc15; border-radius: 10px; padding: 12px; margin-bottom: 18px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
    td, th {{ border: 1px solid #e5e7eb; padding: 6px 8px; vertical-align: top; }}
  </style>
</head>
<body>
  <div class="top"><h1>{html.escape(title)}</h1><div class="source">{html.escape(source)}</div></div>
  <div class="doc">{body}</div>
</body>
</html>"""


def read_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
        doc = Document(str(path))
        parts = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    except Exception:
        pass

    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        xml = re.sub(r"<w:tab[^>]*>", "\t", xml)
        xml = re.sub(r"</w:p>", "\n", xml)
        xml = re.sub(r"<[^>]+>", "", xml)
        return html.unescape(xml)
    except Exception as exc:
        return f"Aperçu DOCX indisponible.\n\nErreur : {exc}"


def read_xlsx_as_html(path: Path) -> str:
    try:
        from openpyxl import load_workbook  # type: ignore
        wb = load_workbook(path, read_only=True, data_only=True)
        blocks = []
        for ws in wb.worksheets[:4]:
            rows = []
            for row in ws.iter_rows(max_row=80, max_col=12, values_only=True):
                values = ["" if cell is None else str(cell) for cell in row]
                if any(v.strip() for v in values):
                    rows.append(values)
            if not rows:
                continue
            trs = []
            for row in rows:
                tds = "".join(f"<td>{html.escape(v)}</td>" for v in row)
                trs.append(f"<tr>{tds}</tr>")
            blocks.append(f"<h2>{html.escape(ws.title)}</h2><table>{''.join(trs)}</table>")
        return "".join(blocks) or "<p>Aucune donnée lisible dans ce fichier Excel.</p>"
    except Exception as exc:
        return f"<p>Aperçu Excel indisponible : {html.escape(str(exc))}</p>"


def render_file_preview(path: Path, highlight: str = ""):
    suffix = path.suffix.lower()
    if suffix in {".pdf", ".png", ".jpg", ".jpeg", ".gif", ".webp"}:
        media_type = "application/pdf" if suffix == ".pdf" else None
        return FileResponse(path, media_type=media_type, filename=path.name, headers={"Content-Disposition": f'inline; filename="{quote(path.name)}"'})
    if suffix == ".docx":
        text = read_docx_text(path)
        return HTMLResponse(html_page(path.name, highlight_html(text, highlight), str(path)))
    if suffix in {".txt", ".md", ".json", ".csv"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return HTMLResponse(html_page(path.name, highlight_html(text, highlight), str(path)))
    if suffix in {".xlsx", ".xlsm"}:
        return HTMLResponse(html_page(path.name, read_xlsx_as_html(path), str(path)))
    return PlainTextResponse(f"Aperçu non disponible pour ce type de fichier : {suffix}\n\nFichier : {path}", media_type="text/plain; charset=utf-8")


def find_previous_cir_file(organisme: str, project: str, year: str) -> Optional[Path]:
    org = slugify(organisme)
    proj = slugify(project)
    y = str(year).strip()
    bases = [
        STORAGE_ROOT / "organismes" / org / "projects" / proj / "years" / y / "cir_final_consultant" / "current",
        STORAGE_ROOT / "organismes" / org / "projects" / proj / "years" / y / "cir_final",
        STORAGE_ROOT / "organismes" / org / "projects" / proj / "years" / y,
    ]
    for base in bases:
        if not base.exists():
            continue
        for ext in ("*.pdf", "*.docx", "*.txt", "*.md"):
            files = sorted(base.glob(ext), key=lambda p: p.stat().st_mtime, reverse=True)
            if files:
                return files[0]
    return None


def render_memory_preview(organisme: str, project: str, year: str, highlight: str):
    org = slugify(organisme)
    proj = slugify(project)
    memory_paths = [
        STORAGE_ROOT / "organismes" / org / "projects" / proj / "years" / str(year) / "cir_final_memory.json",
        STORAGE_ROOT / "organismes" / org / "projects" / proj / "years" / str(year) / "cir_final_consultant" / "current" / "cir_final_memory.json",
    ]
    for path in memory_paths:
        if path.exists():
            data = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
            sections = data.get("sections_full") or data.get("sections") or {}
            text = "\n\n".join(f"{key}\n{value}" for key, value in sections.items() if isinstance(value, str) and value.strip())
            return HTMLResponse(html_page(f"CIR {year} — mémoire extraite", highlight_html(text, highlight), str(path)))
    raise HTTPException(status_code=404, detail="Aucun fichier ou mémoire CIR précédent trouvée pour cette année.")


@router.get("/projects/{project_id}/source-preview")
def preview_source_file(project_id: int, source_path: str = Query(...), highlight: str = Query("", max_length=5000)):
    path = safe_storage_path(source_path)
    return render_file_preview(path, highlight)


@router.get("/projects/{project_id}/cir-previous/preview")
def preview_previous_cir(project_id: int, organisme: str = Query(...), project: str = Query(...), year: str = Query(...), highlight: str = Query("", max_length=5000)):
    file_path = find_previous_cir_file(organisme, project, year)
    if file_path:
        return render_file_preview(file_path, highlight)
    return render_memory_preview(organisme, project, year, highlight)
