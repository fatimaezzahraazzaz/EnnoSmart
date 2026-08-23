DIAGNOSTIC_ROUTER_VERSION = "v143_complete_db_persistence"

from fastapi import APIRouter, Depends, HTTPException, Query, status, UploadFile, File, Form
from sqlalchemy.orm import Session, defer
from typing import Any, Dict
from pathlib import Path
from datetime import datetime
import re
import json
import importlib.util
import shutil
import hashlib

from core.deps import get_current_user, get_db, require_agent_enabled
from db.models import DiagnosticRun, User, Verrou
from schemas.diagnostic import (
    DiagnosticRead,
    VerrouDecisionRequest,
    VerrouManualCreate,
    VerrouRead,
)
from services.diagnostic_display_service import (
    build_compact_diagnostic_display,
    build_diagnostic_display,
)
from services import diagnostic_service as diagnostic_service_module
from services.consultant_verrou_service import (
    create_or_reuse_consultant_verrou,
    get_latest_diagnostic_verrous,
)
from services.diagnostic_service import (
    create_diagnostic_run_from_files,
    prepare_ennodiagnostic_sources,
    read_diagnostic_bundle,
    run_ennodiagnostic,
    run_ennodiagnostic_agent_only,
    sanitize_json_value,
    sync_verrous_from_diagnostic,
    ensure_ennosmart_imports,
)
from services.project_service import get_project_for_user


router = APIRouter(tags=["ennodiagnostic"], dependencies=[Depends(require_agent_enabled("diagnostic"))])


# ============================================================
# V147 - Prechargement des dependances RAG natives
# ============================================================
# Les routes FastAPI synchrones sont executees dans un thread AnyIO.
# Sur cet environnement Windows/Python 3.12, le premier import de
# sentence_transformers -> datasets -> pyarrow depuis ce worker peut provoquer
# une violation d'acces native. On charge donc RAG au moment de l'import du
# router, dans le thread principal d'Uvicorn.
RAG_NATIVE_PRELOAD_VERSION = "v147_main_thread"
ensure_ennosmart_imports()
from modules.RAG.indexer import index_nlp_result as _PRELOADED_INDEX_NLP_RESULT
print(
    "[EnnoDiagnostic][V147] RAG / SentenceTransformers / PyArrow "
    "precharges dans le thread principal.",
    flush=True,
)



def _latest_run_for_project(db: Session, project_id: int) -> DiagnosticRun | None:
    return (
        db.query(DiagnosticRun)
        .options(defer(DiagnosticRun.raw_result_json))
        .filter(DiagnosticRun.project_id == project_id)
        .order_by(DiagnosticRun.created_at.desc())
        .first()
    )


def _as_dict(value: Any) -> Dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _clean(value: Any) -> str:
    return "" if value is None else str(value).strip()


def _section_key(title: str) -> str:
    title = title.lower().strip()
    table = str.maketrans("àâäéèêëîïôöùûüç’'", "aaaeeeeiioouuuc__")
    title = title.translate(table)
    title = re.sub(r"[^a-z0-9]+", "_", title)
    return re.sub(r"_+", "_", title).strip("_")


def _extract_sections(markdown: str) -> Dict[str, str]:
    markdown = _clean(markdown)
    matches = list(re.finditer(r"(?m)^##\s+(.+?)\s*$", markdown))
    if not matches:
        return {}

    sections: Dict[str, str] = {}
    for i, match in enumerate(matches):
        title = _section_key(match.group(1))
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(markdown)
        sections[title] = markdown[start:end].strip()

    return sections



def extract_complete_diagnostic_snapshot(report: Dict[str, Any]) -> Dict[str, Any]:
    """
    Adaptateur robuste V143.

    La version officielle se trouve dans services.diagnostic_service. Le router
    n'importe plus cette fonction directement au chargement, afin qu'une copie
    momentanément désynchronisée des fichiers ne fasse pas tomber Uvicorn.

    Le fallback local conserve tout de même :
    - le Markdown complet ;
    - les sections par titre et par clé ;
    - les sections canoniques du frontend ;
    - les cartes et les verrous finaux ;
    - Frascati, IA, mémoire, Chroma et traçabilité.
    """
    builder = getattr(
        diagnostic_service_module,
        "extract_complete_diagnostic_snapshot",
        None,
    )
    if callable(builder):
        return sanitize_json_value(builder(report))

    report = sanitize_json_value(report if isinstance(report, dict) else {})
    diagnostic = _as_dict(report.get("diagnostic"))
    static = _as_dict(report.get("static_diagnostic"))

    markdown = _clean(
        diagnostic.get("content")
        or report.get("report_markdown")
        or report.get("content")
        or static.get("markdown")
        or ""
    )

    markdown_sections = _extract_sections(markdown)
    sections_by_key: Dict[str, str] = {}
    sections_by_title: Dict[str, str] = {}

    def merge_keyed(value: Any) -> None:
        if not isinstance(value, dict):
            return
        for key, body in value.items():
            text = _clean(body)
            normalized = _section_key(str(key))
            if normalized and text and not sections_by_key.get(normalized):
                sections_by_key[normalized] = text

    def merge_titled(value: Any) -> None:
        if not isinstance(value, dict):
            return
        for title, body in value.items():
            text = _clean(body)
            clean_title = _clean(title)
            if clean_title and text and not sections_by_title.get(clean_title):
                sections_by_title[clean_title] = text

    merge_keyed(report.get("diagnostic_sections_by_key"))
    merge_keyed(static.get("sections_by_key"))
    merge_keyed(report.get("report_sections"))
    merge_keyed(markdown_sections)

    merge_titled(report.get("diagnostic_sections"))
    merge_titled(static.get("sections"))

    for title, body in sections_by_title.items():
        key = _section_key(title)
        if key and body and not sections_by_key.get(key):
            sections_by_key[key] = body

    aliases = {
        "lecture_frascati": ["lecture_frascati", "lecture_frascati_du_dossier"],
        "justification_frascati": [
            "justification_frascati",
            "justification_frascati_du_score",
            "justification_du_score_frascati",
        ],
        "memoire_v2": ["memoire_v2"],
        "synthese_strategique": [
            "synthese_strategique",
            "synthese_strategique_du_projet",
            "synthese",
        ],
        "objectif_global": [
            "objectif_global",
            "objectif_global_reformule",
            "objectif_global_du_projet",
            "objectif",
        ],
        "verrous_rnd": [
            "verrous_rnd",
            "verrous_cir",
            "verrous_cir_consolides",
            "verrous_r_d_signaux_de_verrous",
            "verrous",
        ],
        "demarche_detectee": [
            "demarche_detectee",
            "demarche_experimentale_detectee",
            "demarche_experimentale",
            "demarche",
        ],
        "resultats_metriques": [
            "resultats_metriques",
            "resultats_et_metriques_disponibles",
            "resultats_metriques_disponibles",
            "resultats",
        ],
        "parametres_contraintes": [
            "parametres_contraintes",
            "parametres_et_contraintes_techniques",
            "parametres_techniques",
            "parametres",
        ],
        "points_validation": [
            "points_validation",
            "points_a_valider",
            "points_a_valider_par_le_consultant",
            "validation",
        ],
    }

    canonical: Dict[str, str] = {}
    for canonical_key, candidates in aliases.items():
        for candidate in candidates:
            value = sections_by_key.get(candidate)
            if isinstance(value, str) and value.strip():
                canonical[canonical_key] = value.strip()
                break

    cards = report.get("diagnostic_cards") or static.get("cards") or []
    if not isinstance(cards, list):
        cards = []

    final_verrous = _extract_final_accepted_verrous_from_report(report)

    return sanitize_json_value({
        "snapshot_version": "v143_complete_db_persistence_router_fallback",
        "generated_at": report.get("generated_at") or datetime.utcnow().isoformat(),
        "mode": report.get("mode"),
        "status": report.get("status") or diagnostic.get("status"),
        "report_markdown": markdown,
        "sections_by_key": sections_by_key,
        "sections_by_title": sections_by_title,
        "canonical_sections": canonical,
        "sections_count": len(sections_by_key),
        "section_titles_count": len(sections_by_title),
        "diagnostic_cards": cards,
        "diagnostic_cards_count": len(cards),
        "final_verrous": final_verrous,
        "final_verrous_count": len(final_verrous),
        "frascati_summary": report.get("frascati_summary") or {},
        "frascati_justification": report.get("frascati_justification") or {},
        "ai_detection_report": (
            report.get("ai_detection_report_runtime")
            or report.get("ai_detection_report")
            or {}
        ),
        "style_memory_report": report.get("style_memory_report") or {},
        "cir_memory_report": report.get("cir_memory_report") or {},
        "inputs_status": report.get("inputs_status") or {},
        "pipeline_before_agent": report.get("pipeline_before_agent") or {},
        "verrou_synthesis_report": report.get("verrou_synthesis_report") or {},
        "chroma_sections": report.get("chroma_sections") or {},
        "source_paths": {"output_path": report.get("output_path")},
        "fallback_reason": (
            "services.diagnostic_service.extract_complete_diagnostic_snapshot "
            "indisponible au chargement"
        ),
    })


def _replace_section(markdown: str, section_title: str, body: str) -> str:
    """
    Remplace une section Markdown de manière robuste :
    - accepte espaces après le titre ;
    - accepte CRLF ;
    - conserve toutes les autres sections.
    """
    markdown = _clean(markdown).replace("\r\n", "\n").replace("\r", "\n")
    section_title_escaped = re.escape(section_title)
    pattern = re.compile(
        rf"(?ms)^##\s+{section_title_escaped}[^\n]*\n.*?(?=^##\s+|\Z)"
    )
    new_block = f"## {section_title}  \n{body.strip()}\n\n"

    if pattern.search(markdown):
        return pattern.sub(new_block, markdown).strip()

    return (markdown.rstrip() + "\n\n" + new_block).strip()


def _normalize_agent_report_content(content: str) -> str:
    """Normalisation générique du rapport agent, sans forcer d'objectif projet."""
    return _clean(content)


def _extract_latest_run_report(latest_run: DiagnosticRun | None) -> tuple[Dict[str, Any], str]:
    """
    Récupère le rapport EnnoDiagnostic depuis le dernier run.

    Correction V136 :
    selon les versions, raw_result_json peut contenir :
    1) {"script_or_pipeline_result": {"report": ...}}
    2) {"report": ...}
    3) directement le rapport final avec verrou_synthesis_report au premier niveau.

    Avant, seul le cas 1 était lu, donc sync-verrous pouvait ne rien trouver.
    """
    if not latest_run or not latest_run.raw_result_json:
        return {}, ""

    raw = sanitize_json_value(latest_run.raw_result_json)
    if not isinstance(raw, dict):
        return {}, ""

    report: Dict[str, Any] = {}

    pipeline = _as_dict(raw.get("script_or_pipeline_result"))
    if isinstance(pipeline.get("report"), dict):
        report = _as_dict(pipeline.get("report"))

    if not report and isinstance(raw.get("report"), dict):
        report = _as_dict(raw.get("report"))

    # Cas où raw_result_json EST déjà le rapport final.
    if not report and any(
        key in raw
        for key in [
            "verrou_synthesis_report",
            "static_diagnostic",
            "diagnostic_sections",
            "diagnostic_sections_by_key",
            "frascati_summary",
        ]
    ):
        report = raw

    diagnostic = _as_dict(report.get("diagnostic"))

    content = _clean(
        diagnostic.get("content")
        or report.get("report_markdown")
        or report.get("content")
        or ""
    )

    return report, content


# ============================================================
# V139 — choix d'une seule sortie officielle pour /diagnostic/latest
# ============================================================

def _parse_report_datetime(value: Any) -> float:
    if not value:
        return 0.0

    if isinstance(value, datetime):
        return value.timestamp()

    text = str(value).strip()
    if not text:
        return 0.0

    try:
        # Accepte "2026-07-10T17:34:17" et variantes ISO.
        return datetime.fromisoformat(text.replace("Z", "+00:00")).timestamp()
    except Exception:
        return 0.0


def _path_mtime(value: Any) -> float:
    if not value:
        return 0.0

    try:
        path = Path(str(value))
        if path.exists() and path.is_file():
            return path.stat().st_mtime
    except Exception:
        return 0.0

    return 0.0


def _report_timestamp(report: Dict[str, Any], fallback: float = 0.0) -> float:
    if not isinstance(report, dict) or not report:
        return fallback

    candidates = [
        report.get("generated_at"),
        report.get("completed_at"),
        report.get("created_at"),
        report.get("updated_at"),
        _as_dict(report.get("diagnostic")).get("generated_at"),
        _as_dict(report.get("static_diagnostic")).get("generated_at"),
    ]

    for candidate in candidates:
        ts = _parse_report_datetime(candidate)
        if ts > 0:
            return ts

    return fallback


def _run_timestamp(latest_run: DiagnosticRun | None, report: Dict[str, Any] | None = None) -> float:
    if not latest_run:
        return 0.0

    values = [
        getattr(latest_run, "completed_at", None),
        getattr(latest_run, "updated_at", None),
        getattr(latest_run, "created_at", None),
    ]

    for value in values:
        ts = _parse_report_datetime(value)
        if ts > 0:
            return ts

    return _report_timestamp(report or {}, 0.0)


def _choose_latest_report_source(
    bundle: Dict[str, Any],
    latest_run: DiagnosticRun | None,
    *,
    include_run_raw: bool = True,
) -> Dict[str, Any]:
    """
    Décide quelle sortie EnnoDiagnostic est officielle pour l'affichage.

    Problème corrigé :
    - si l'agent est lancé en CLI, il écrit bien
      storage/.../ennodiagnostic/ennodiagnostic_report.json ;
    - mais aucun DiagnosticRun DB n'est créé ;
    - l'ancien /diagnostic/latest reprenait alors le dernier run DB ancien
      et écrasait l'affichage du nouveau fichier.

    Règle V139 :
    - on compare le rapport fichier et le rapport du dernier run DB ;
    - on expose le plus récent ;
    - le frontend ne choisit plus lui-même.
    """
    bundle = dict(bundle or {})
    file_report = _as_dict(bundle.get("report"))
    file_path = bundle.get("report_path_used")
    file_ts = _report_timestamp(file_report, _path_mtime(file_path))

    run_report: Dict[str, Any] = {}
    run_ts = _run_timestamp(latest_run)
    # Si le fichier officiel est clairement plus récent, sa date suffit pour
    # choisir la source. On évite alors de décompresser/lire un JSONB de run qui
    # peut dépasser 30 Mo juste pour confirmer qu'il est ancien.
    should_read_run_payload = bool(latest_run) and (
        include_run_raw or not file_report or run_ts >= file_ts
    )
    if should_read_run_payload:
        run_report, _run_content = _extract_latest_run_report(latest_run)
        run_ts = _run_timestamp(latest_run, run_report)

    use_run = bool(run_report) and (not file_report or run_ts >= file_ts)

    raw_run_json = None
    if latest_run and (include_run_raw or use_run) and latest_run.raw_result_json:
        if include_run_raw:
            raw_run_json = sanitize_json_value(latest_run.raw_result_json)
            bundle["run_raw_result_json"] = raw_run_json
        else:
            raw_run_json = _as_dict(latest_run.raw_result_json)
            snapshot = _as_dict(raw_run_json.get("diagnostic_snapshot"))
            if snapshot:
                bundle["diagnostic_snapshot"] = snapshot

    if use_run:
        if include_run_raw:
            bundle = _merge_latest_run_into_bundle(bundle, latest_run)
        else:
            bundle["report"] = run_report
        bundle["official_report_source"] = "db_latest_run"
        bundle["official_report_timestamp"] = run_ts
        bundle["official_report_note"] = "Rapport DB plus récent ou aucun rapport fichier disponible."
    else:
        bundle["official_report_source"] = "filesystem_report"
        bundle["official_report_timestamp"] = file_ts
        bundle["official_report_note"] = "Rapport fichier ProjectStore plus récent que le dernier run DB, ou run DB absent."

    bundle["official_report_debug"] = {
        "file_report_path": str(file_path) if file_path else None,
        "file_report_timestamp": file_ts,
        "db_run_id": latest_run.id if latest_run else None,
        "db_run_timestamp": run_ts,
        "used": bundle["official_report_source"],
    }

    return sanitize_json_value(bundle)



# ============================================================
# V132 — extraction officielle des sorties finales EnnoDiagnostic
# ============================================================

def _extract_final_accepted_verrous_from_report(report: dict) -> list[dict[str, Any]]:
    """
    Source officielle pour les verrous à synchroniser en base.

    Important :
    - On prend la liste finale acceptée par EnnoDiagnostic.
    - On ne prend pas uniquement les verrous LLM clean/intermédiaires.
    - Les fallbacks acceptés par le garde final doivent aussi être synchronisés.
    """
    if not isinstance(report, dict):
        return []

    verrou_report = report.get("verrou_synthesis_report") or {}
    if not isinstance(verrou_report, dict):
        verrou_report = {}

    candidates = (
        verrou_report.get("llm_reformulated_verrous")
        or verrou_report.get("final_items")
        or verrou_report.get("accepted_items")
        or verrou_report.get("final_verrous")
        or report.get("llm_reformulated_verrous")
        or report.get("consultant_verrous_cir")
        or report.get("verrous_reformules")
        or report.get("verrous")
        or []
    )

    if not isinstance(candidates, list):
        return []

    final_items: list[dict[str, Any]] = []
    seen: set[str] = set()

    for item in candidates:
        if not isinstance(item, dict):
            continue

        title = _clean(
            item.get("title")
            or item.get("titre")
            or item.get("verrou")
            or item.get("name")
            or item.get("lock_title")
        )

        if not title:
            continue

        key = re.sub(r"\s+", " ", title.lower()).strip()
        if key in seen:
            continue
        seen.add(key)

        justification = _clean(
            item.get("consultant_explanation")
            or item.get("why_agent_found_verrou")
            or item.get("why_not_simple_engineering")
            or item.get("justification")
            or item.get("description")
            or item.get("text")
            or item.get("scientific_lock")
        )

        source_doc = _clean(
            item.get("source_document")
            or item.get("document")
            or item.get("documents")
            or item.get("source")
        )

        source_excerpt = _clean(
            item.get("source_excerpt")
            or item.get("excerpt")
            or item.get("evidence")
            or item.get("preuve")
            or justification
        )

        score = (
            item.get("score")
            or item.get("frascati_score")
            or item.get("confidence_score")
            or item.get("confidence")
            or item.get("rank_score")
        )

        try:
            if score is not None:
                score = float(score)
        except Exception:
            score = None

        final_items.append({
            **item,
            "title": title,
            "titre": title,
            "verrou": title,
            "description": justification,
            "justification": justification,
            "text": justification,
            "score": score,
            "source_document": source_doc,
            "source_excerpt": source_excerpt,
            "consultant_status": item.get("consultant_status") or item.get("status") or "en_attente",
            "needs_human_validation": True,
            "sync_source": "verrou_synthesis_report.llm_reformulated_verrous_final",
        })

    return final_items


def _extract_final_display_sections_from_report(report: dict) -> dict:
    """
    Source officielle pour les sections affichées.
    On privilégie les sections statiques V129/V132 et on évite de parser l'ancien Markdown.
    """
    if not isinstance(report, dict):
        return {
            "diagnostic_sections_by_key": {},
            "diagnostic_sections": {},
            "diagnostic_cards": [],
            "static_diagnostic": {},
        }

    static_diagnostic = report.get("static_diagnostic") or {}
    if not isinstance(static_diagnostic, dict):
        static_diagnostic = {}

    sections_by_key = (
        report.get("diagnostic_sections_by_key")
        or static_diagnostic.get("sections_by_key")
        or {}
    )
    if not isinstance(sections_by_key, dict):
        sections_by_key = {}

    sections = (
        report.get("diagnostic_sections")
        or static_diagnostic.get("sections")
        or {}
    )
    if not isinstance(sections, dict):
        sections = {}

    cards = (
        report.get("diagnostic_cards")
        or static_diagnostic.get("cards")
        or []
    )
    if not isinstance(cards, list):
        cards = []

    return {
        "diagnostic_sections_by_key": sections_by_key,
        "diagnostic_sections": sections,
        "diagnostic_cards": cards,
        "static_diagnostic": static_diagnostic,
    }


def _set_if_has(obj: Any, names: str | list[str], value: Any) -> None:
    if value is None:
        return
    if isinstance(names, str):
        names = [names]
    for name in names:
        if hasattr(obj, name):
            try:
                setattr(obj, name, value)
                return
            except Exception:
                pass


def _model_has_attr(obj: Any, name: str) -> bool:
    return hasattr(obj, name)



def _report_fingerprint(report: Dict[str, Any]) -> str:
    """
    Empreinte stable du rapport réellement affiché.

    Elle permet de savoir si le rapport fichier a déjà été matérialisé en
    DiagnosticRun. On évite ainsi de créer un nouveau run à chaque GET.
    """
    report = _as_dict(report)
    final_items = _extract_final_accepted_verrous_from_report(report)

    payload = {
        "generated_at": report.get("generated_at"),
        "mode": report.get("mode"),
        "verrou_synthesis_version": report.get("verrou_synthesis_version"),
        "output_path": report.get("output_path"),
        "titles": [_clean(item.get("title")) for item in final_items],
    }

    encoded = json.dumps(
        sanitize_json_value(payload),
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")

    return hashlib.sha256(encoded).hexdigest()


def _dict_metadata_from_verrou(verrou: Verrou) -> Dict[str, Any]:
    """
    Récupère les métadonnées riches sauvegardées dans le modèle SQLAlchemy.

    Les noms diffèrent selon les versions du modèle : raw_json,
    metadata_json ou details. On ne retourne que les dictionnaires.
    """
    for name in ["raw_json", "metadata_json", "details", "source_json"]:
        try:
            value = getattr(verrou, name, None)
        except Exception:
            value = None
        if isinstance(value, dict):
            return sanitize_json_value(value)

    return {}


def _find_run_with_same_report(
    db: Session,
    project_id: int,
    report: Dict[str, Any],
    max_runs: int = 3,
) -> DiagnosticRun | None:
    """
    Cherche un DiagnosticRun déjà lié exactement au même rapport fichier.
    """
    expected = _report_fingerprint(report)
    if not expected:
        return None

    # Les runs créés par les versions récentes portent un fingerprint indexable
    # dans leur JSON. Cette recherche évite de rapatrier puis désérialiser les
    # rapports complets des 20 derniers runs (jusqu'à plusieurs centaines de Mo).
    try:
        fingerprint_match = (
            db.query(DiagnosticRun)
            .filter(
                DiagnosticRun.project_id == project_id,
                DiagnosticRun.raw_result_json["report_fingerprint"].as_string()
                == expected,
            )
            .order_by(DiagnosticRun.created_at.desc())
            .first()
        )
        if fingerprint_match is not None:
            return fingerprint_match
    except Exception:
        # Compatibilité SQLite/anciens drivers JSON pendant les tests locaux.
        pass

    runs = (
        db.query(DiagnosticRun)
        .filter(DiagnosticRun.project_id == project_id)
        .order_by(DiagnosticRun.created_at.desc())
        .limit(max_runs)
        .all()
    )

    for run in runs:
        run_report, _content = _extract_latest_run_report(run)
        if run_report and _report_fingerprint(run_report) == expected:
            return run

    return None


def _materialize_filesystem_report_in_db(
    db: Session,
    project: Any,
    bundle: Dict[str, Any],
) -> tuple[DiagnosticRun | None, list[Verrou], bool]:
    """Matérialise atomiquement le rapport fichier complet et ses verrous en DB."""
    report = _as_dict(bundle.get("report"))
    final_items = _extract_final_accepted_verrous_from_report(report)
    if not report or not final_items:
        return None, [], False

    existing_run = _find_run_with_same_report(db, project.id, report)
    if existing_run is not None:
        current = _read_latest_run_verrous_fast(db, existing_run)
        if not current:
            current = sync_verrous_from_diagnostic(db, existing_run)
        return existing_run, current, False

    snapshot = extract_complete_diagnostic_snapshot(report)
    payload = sanitize_json_value({
        "persistence_version": "v143_complete_db_persistence",
        "report_fingerprint": _report_fingerprint(report),
        "saved_at": datetime.utcnow().isoformat(),
        "button": "auto_materialize_filesystem_report_v142",
        "pipeline": "filesystem_report_to_complete_db",
        "report": report,
        "script_or_pipeline_result": {"report": report},
        "diagnostic_snapshot": snapshot,
        "report_markdown": snapshot.get("report_markdown"),
        "report_sections": snapshot.get("canonical_sections") or {},
        "diagnostic_sections_by_key": snapshot.get("sections_by_key") or {},
        "diagnostic_sections": snapshot.get("sections_by_title") or {},
        "diagnostic_cards": snapshot.get("diagnostic_cards") or [],
        "final_verrous_snapshot": snapshot.get("final_verrous") or [],
        "report_fingerprint": _report_fingerprint(report),
        "bundle_metadata": {
            "report_path_used": bundle.get("report_path_used"),
            "nlp_path_used": bundle.get("nlp_path_used"),
            "official_report_source": bundle.get("official_report_source"),
            "official_report_timestamp": bundle.get("official_report_timestamp"),
        },
    })

    run = DiagnosticRun(
        project_id=project.id,
        status="completed_from_filesystem_report_v142",
        report_path=str(bundle.get("report_path_used") or "") or None,
        nlp_result_path=str(bundle.get("nlp_path_used") or "") or None,
        selected_verrous_path=None,
        raw_result_json=payload,
        completed_at=datetime.utcnow(),
    )

    try:
        db.add(run)
        db.flush()
        synced = sync_verrous_from_diagnostic(db, run, commit=False)
        if not synced:
            raise RuntimeError("Aucun verrou final n'a pu être synchronisé depuis le rapport fichier.")
        db.commit()
        db.refresh(run)
        for verrou in synced:
            db.refresh(verrou)
        return run, synced, True
    except Exception:
        db.rollback()
        raise

def _sync_final_verrous_from_run(db: Session, run: DiagnosticRun) -> list[Verrou]:
    """Délègue à la persistance V142 unique du service backend."""
    return sync_verrous_from_diagnostic(db, run)

def _dump_verrous_for_frontend(verrous: list[Verrou]) -> list[dict[str, Any]]:
    """
    Convertit les verrous DB en objets directement utilisables par le frontend.

    V140 :
    - impose un id DB positif ;
    - restitue les métadonnées riches dans source_json ;
    - expose can_decide=true seulement pour un vrai verrou PostgreSQL.
    """
    out: list[dict[str, Any]] = []

    for verrou in verrous or []:
        try:
            data = VerrouRead.model_validate(verrou).model_dump()
        except Exception:
            data = {}

        verrou_id = getattr(verrou, "id", None)
        try:
            verrou_id_int = int(verrou_id) if verrou_id is not None else None
        except Exception:
            verrou_id_int = None

        title = (
            data.get("title")
            or data.get("titre")
            or data.get("verrou")
            or getattr(verrou, "title", "")
            or getattr(verrou, "titre", "")
            or ""
        )

        description = (
            data.get("description")
            or data.get("justification")
            or data.get("text")
            or getattr(verrou, "description", "")
            or getattr(verrou, "justification", "")
            or getattr(verrou, "text", "")
            or ""
        )

        consultant_status = (
            data.get("consultant_status")
            or getattr(verrou, "consultant_status", "")
            or "en_attente"
        )

        rich_source_json = _dict_metadata_from_verrou(verrou)
        existing_source_json = data.get("source_json")
        if not isinstance(existing_source_json, dict):
            existing_source_json = {}

        source_json = {
            **rich_source_json,
            **existing_source_json,
            "db_verrou_id": verrou_id_int,
            "is_db_synced": bool(verrou_id_int and verrou_id_int > 0),
            "can_decide": bool(verrou_id_int and verrou_id_int > 0),
            "sync_source": "backend_v140_report_materialized_to_db",
        }

        data["id"] = verrou_id_int
        data["title"] = title
        data["titre"] = title
        data["verrou"] = title
        data["description"] = description
        data["justification"] = data.get("justification") or description
        data["text"] = data.get("text") or description
        data["consultant_status"] = consultant_status
        data["source_json"] = source_json
        data["is_db_synced"] = bool(verrou_id_int and verrou_id_int > 0)
        data["can_decide"] = bool(verrou_id_int and verrou_id_int > 0)
        data["source"] = "db_synced_verrou"
        data["sync_source"] = "backend_v140_report_materialized_to_db"

        out.append(data)

    return sanitize_json_value(out)


def _js_int32(value: int) -> int:
    value &= 0xFFFFFFFF
    return value - 0x100000000 if value & 0x80000000 else value


def _legacy_frontend_negative_id(title: str, index: int) -> int:
    """Reproduit exactement stableNegativeIdV107() du frontend historique."""
    key = f"{title or 'verrou'}-{index}"
    hash_value = 0
    for char in key:
        hash_value = _js_int32(((hash_value << 5) - hash_value + ord(char)))
    return -abs(hash_value or (index + 1))


def _normalize_title_for_match(value: Any) -> str:
    text = _clean(value).lower()
    text = text.translate(str.maketrans(
        "àâäéèêëîïôöùûüç’'",
        "aaaeeeeiioouuuc__",
    ))
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _resolve_negative_frontend_verrou(
    db: Session,
    project: Any,
    negative_id: int,
) -> Verrou | None:
    """
    Compatibilité V141 avec les anciens écrans qui ont encore un id négatif.

    Le backend retrouve le titre correspondant dans le rapport officiel,
    matérialise ce rapport en DB, puis retourne le vrai verrou PostgreSQL.
    """
    if negative_id >= 0:
        return None

    latest_run = _latest_run_for_project(db, project.id)
    base_bundle = read_diagnostic_bundle(project)
    bundle = _choose_latest_report_source(base_bundle, latest_run)
    report = _as_dict(bundle.get("report"))
    final_items = _extract_final_accepted_verrous_from_report(report)

    matched_item: dict[str, Any] | None = None
    for index, item in enumerate(final_items):
        title = _clean(item.get("title") or item.get("titre") or item.get("verrou"))
        if title and _legacy_frontend_negative_id(title, index) == negative_id:
            matched_item = item
            break

    # Sécurité : essayer également le rapport du dernier run DB si le rapport
    # fichier a changé entre l'affichage et le clic consultant.
    if matched_item is None and latest_run is not None:
        run_report, _content = _extract_latest_run_report(latest_run)
        for index, item in enumerate(_extract_final_accepted_verrous_from_report(run_report)):
            title = _clean(item.get("title") or item.get("titre") or item.get("verrou"))
            if title and _legacy_frontend_negative_id(title, index) == negative_id:
                matched_item = item
                report = run_report
                bundle = {
                    **bundle,
                    "report": run_report,
                    "report_path_used": getattr(latest_run, "report_path", None),
                }
                break

    if matched_item is None:
        return None

    matched_title = _clean(matched_item.get("title") or matched_item.get("titre") or matched_item.get("verrou"))
    normalized_title = _normalize_title_for_match(matched_title)

    # Chercher d'abord un verrou déjà synchronisé, quel que soit le run.
    candidates = (
        db.query(Verrou)
        .join(DiagnosticRun, Verrou.diagnostic_run_id == DiagnosticRun.id)
        .filter(DiagnosticRun.project_id == project.id)
        .order_by(DiagnosticRun.created_at.desc(), Verrou.created_at.desc())
        .all()
    )
    for verrou in candidates:
        if _normalize_title_for_match(getattr(verrou, "title", "")) == normalized_title:
            return verrou

    # Sinon matérialiser le rapport officiel et synchroniser ses verrous.
    materialized_run, synced, _created = _materialize_filesystem_report_in_db(
        db=db,
        project=project,
        bundle=bundle,
    )

    for verrou in synced:
        if _normalize_title_for_match(getattr(verrou, "title", "")) == normalized_title:
            return verrou

    if materialized_run is not None:
        for verrou in _read_latest_run_verrous_fast(db, materialized_run):
            if _normalize_title_for_match(getattr(verrou, "title", "")) == normalized_title:
                return verrou

    return None


def _read_latest_run_verrous_fast(db: Session, latest_run: DiagnosticRun | None) -> list[Verrou]:
    """
    Lecture rapide pour les routes GET.

    Important :
    - aucune resynchronisation ;
    - aucune relecture / réécriture du gros rapport JSON ;
    - utilisé par /diagnostic/latest et /verrous.

    La synchronisation des verrous doit rester dans les routes d'action :
    - POST /diagnostic/run
    - POST /diagnostic/run-agent
    - POST /diagnostic/{run_id}/sync-verrous
    """
    if not latest_run:
        return []

    return (
        db.query(Verrou)
        .filter(Verrou.diagnostic_run_id == latest_run.id)
        .order_by(Verrou.created_at.asc())
        .all()
    )


def _ensure_latest_run_final_verrous_synced(db: Session, latest_run: DiagnosticRun | None) -> list[Verrou]:
    """
    Synchronisation volontaire des verrous finaux.

    Ne pas appeler cette fonction dans une route GET de lecture simple,
    sinon l'interface reste bloquée sur "Chargement du diagnostic".
    """
    if not latest_run:
        return []

    report, _content = _extract_latest_run_report(latest_run)
    final_items = _extract_final_accepted_verrous_from_report(report)

    current = _read_latest_run_verrous_fast(db, latest_run)

    if final_items:
        return _sync_final_verrous_from_run(db, latest_run)

    return current

def _force_display_from_latest_run(display: Dict[str, Any], latest_run: DiagnosticRun | None, project=None) -> Dict[str, Any]:
    display = dict(display or {})
    report, content = _extract_latest_run_report(latest_run)

    if not report and not content:
        return display

    content = _normalize_agent_report_content(content)
    sections = _extract_sections(content) if content else {}

    final_display = _extract_final_display_sections_from_report(report)
    agent_sections = final_display["diagnostic_sections"]
    agent_sections_by_key = final_display["diagnostic_sections_by_key"]
    diagnostic_cards = final_display["diagnostic_cards"]
    static_diagnostic = final_display["static_diagnostic"]
    final_verrous = _extract_final_accepted_verrous_from_report(report)

    def pick(key: str, title: str = "", *legacy_keys: str) -> str:
        if key and agent_sections_by_key.get(key):
            return _clean(agent_sections_by_key.get(key))
        if title and agent_sections.get(title):
            return _clean(agent_sections.get(title))
        for lk in legacy_keys:
            if lk and sections.get(lk):
                return _clean(sections.get(lk))
        return ""

    lecture_frascati = pick("lecture_frascati", "Lecture Frascati du dossier", "lecture_frascati_du_dossier")
    justification_frascati = pick("justification_frascati", "Justification Frascati du score", "justification_frascati_du_score")
    memoire_v2 = pick("memoire_v2", "Mémoire V2", "memoire_v2")
    synthese = pick("synthese_strategique", "Synthèse stratégique", "synthese_strategique_du_projet")
    objectif = pick("objectif_global", "Objectif global", "objectif_global_reformule", "objectif_global_du_projet")
    verrous_section = pick(
        "verrous_rnd",
        "Verrous CIR consolidés",
        "signaux_de_verrous_r_d_candidats",
        "verrous_r_d_signaux_de_verrous",
        "verrous_cir_consolides",
        "verrous_r_d",
        "verrous",
    ) or display.get("verrous_text", "")
    demarche = pick("demarche_detectee", "Démarche détectée", "demarche_experimentale_detectee")
    resultats = pick("resultats_metriques", "Résultats / métriques", "resultats_et_metriques_disponibles")
    parametres = pick("parametres_contraintes", "Paramètres et contraintes techniques", "parametres_et_contraintes_techniques")
    points_validation = pick("points_validation", "Points à valider", "points_a_valider_par_le_consultant")

    display["source"] = "ennodiagnostic_agent_v132_final_static_display"
    display["report_markdown"] = content
    display["summary"] = synthese or display.get("summary", "")
    display["objective"] = objectif or display.get("objective", "")
    display["frascati_text"] = lecture_frascati or display.get("frascati_text", "")
    display["frascati_justification_text"] = justification_frascati or display.get("frascati_justification_text", "")
    display["verrous_text"] = verrous_section

    # Sorties officielles pour le frontend : sections statiques, cartes et verrous finaux.
    display["static_diagnostic"] = static_diagnostic
    display["diagnostic_cards"] = diagnostic_cards
    display["diagnostic_sections_by_key"] = agent_sections_by_key
    display["diagnostic_sections"] = agent_sections
    display["llm_reformulated_verrous"] = final_verrous
    display["consultant_verrous_cir"] = final_verrous
    display["validation_verrous_preview"] = final_verrous

    display["report_sections"] = {
        "lecture_frascati": lecture_frascati,
        "justification_frascati": justification_frascati,
        "memoire_v2": memoire_v2,
        "synthese": synthese,
        "objectif": objectif,
        "verrous": verrous_section,
        "signaux_de_verrous": verrous_section,
        "demarche": demarche,
        "resultats": resultats,
        "parametres": parametres,
        "comparaison_cir": sections.get("comparaison_avec_le_cir_precedent", ""),
        "points_validation": points_validation,
    }

    if report:
        display["frascati_summary"] = report.get("frascati_summary") or display.get("frascati_summary") or {}
        display["inputs_status"] = report.get("inputs_status") or display.get("inputs_status") or {}
        display["chroma_sections"] = report.get("chroma_sections") or display.get("chroma_sections") or {}
        display["verrou_synthesis_report"] = report.get("verrou_synthesis_report") or display.get("verrou_synthesis_report") or {}
        display["consultant_validation_source"] = "verrou_synthesis_report.llm_reformulated_verrous_final"

        pipeline = _as_dict(report.get("pipeline_before_agent"))
        nlp_stats = _as_dict(pipeline.get("nlp_stats"))
        index_report = _as_dict(pipeline.get("index_report"))

        display["pipeline_stats"] = {
            "documents_loaded_count": pipeline.get("documents_loaded_count"),
            "raw_candidates": nlp_stats.get("raw_candidates"),
            "raw_kept": nlp_stats.get("raw_kept"),
            "merged_verrous": nlp_stats.get("merged_verrous"),
            "chunks_prepared": index_report.get("chunks_prepared"),
            "chunks_indexed": index_report.get("chunks_indexed"),
        }

    display["forced_from_latest_run"] = True

    # Comparaison documentaire brute exposée au frontend.
    doc_compare_index = {}
    doc_compare_dir_value = None
    try:
        from services.diagnostic_service import get_project_store

        if project is not None:
            ps = get_project_store(project)
            doc_compare_dir = ps.project_dir / "document_compare"
            doc_compare_dir_value = str(doc_compare_dir)
            index_path = doc_compare_dir / "auto_compare_index.json"

            if index_path.exists():
                try:
                    doc_compare_index = json.loads(index_path.read_text(encoding="utf-8"))
                except Exception:
                    doc_compare_index = {}
    except Exception:
        doc_compare_index = {}

    if isinstance(doc_compare_index, dict) and doc_compare_index.get("ok"):
        display["document_compare"] = doc_compare_index
        display["document_compare_ok"] = True
        display["document_compare_pairs_count"] = doc_compare_index.get("pairs_count", 0)
        display["document_compare_pairs"] = doc_compare_index.get("pairs") or []
        display["document_compare_output_dir"] = doc_compare_index.get("output_dir") or doc_compare_dir_value
    else:
        display["document_compare"] = {}
        display["document_compare_ok"] = False
        display["document_compare_pairs_count"] = 0
        display["document_compare_pairs"] = []
        display["document_compare_output_dir"] = doc_compare_dir_value
        display["document_compare_debug_path"] = doc_compare_dir_value

    # Comparaison avec le CIR précédent exposée au frontend.
    cir_memory_report = report.get("cir_memory_report") if report else {}

    if isinstance(cir_memory_report, dict):
        cir_summary = cir_memory_report.get("summary") or {}
        display["cir_memory"] = cir_memory_report
        display["cir_memory_ok"] = bool(cir_memory_report.get("ok"))
        display["cir_memory_has_previous"] = bool(cir_memory_report.get("has_previous_cir"))
        display["cir_memory_previous_years"] = cir_memory_report.get("previous_cir_years_used") or []
        display["cir_memory_summary"] = cir_summary
        display["cir_memory_project_novelty_score"] = cir_summary.get("project_novelty_score")
        display["cir_memory_signal"] = cir_summary.get("frascati_context_signal")
        display["cir_memory_explanation"] = cir_summary.get("frascati_context_explanation")
        display["cir_memory_new_verrous"] = cir_memory_report.get("new_or_not_found") or []
        display["cir_memory_evolutions"] = cir_memory_report.get("evolution_or_partial_continuity") or []
        display["cir_memory_continuities"] = cir_memory_report.get("continuity_strong") or []
        display["cir_memory_verrou_comparisons"] = cir_memory_report.get("verrou_comparisons") or []
    else:
        display["cir_memory"] = {}
        display["cir_memory_ok"] = False
        display["cir_memory_has_previous"] = False
        display["cir_memory_previous_years"] = []
        display["cir_memory_summary"] = {}
        display["cir_memory_project_novelty_score"] = None
        display["cir_memory_signal"] = None
        display["cir_memory_explanation"] = None
        display["cir_memory_new_verrous"] = []
        display["cir_memory_evolutions"] = []
        display["cir_memory_continuities"] = []
        display["cir_memory_verrou_comparisons"] = []

    # Mémoire rédactionnelle CIR exposée au frontend.
    style_memory_report = report.get("style_memory_report") if report else {}

    if isinstance(style_memory_report, dict):
        display["style_memory"] = style_memory_report
        display["style_memory_ok"] = bool(style_memory_report.get("ok"))
        display["style_memory_examples_count"] = style_memory_report.get("examples_count", 0)
        display["style_memory_roles"] = style_memory_report.get("examples_by_role_count", {})
        display["style_memory_stats"] = style_memory_report.get("stats", {})
    else:
        display["style_memory"] = {}
        display["style_memory_ok"] = False
        display["style_memory_examples_count"] = 0
        display["style_memory_roles"] = {}
        display["style_memory_stats"] = {}

    # Score IA documentaire exposé au frontend.
    ai_report = (
        report.get("ai_detection_report_runtime")
        or report.get("ai_detection_report")
        or {}
    ) if report else {}

    if isinstance(ai_report, dict):
        summary = ai_report.get("summary") or {}
        ai_detection = ai_report.get("ai_detection") or {}
        top_passages = (
            ai_report.get("top_passages")
            or ai_detection.get("suspected_passages")
            or ai_detection.get("passages")
            or []
        )

        display["ai_detection"] = ai_report
        display["ai_score"] = summary.get("average_ai_percentage") or ai_detection.get("global_ai_percentage")
        display["ai_risk_level"] = summary.get("risk_level") or ai_detection.get("risk_level")
        display["ai_suspected_passages"] = top_passages[:10] if isinstance(top_passages, list) else []
    else:
        display["ai_detection"] = {}
        display["ai_score"] = None
        display["ai_risk_level"] = None
        display["ai_suspected_passages"] = []

    return display

def _merge_latest_run_into_bundle(bundle: dict, latest_run: DiagnosticRun | None) -> dict:
    """
    Injecte le rapport du dernier run DB dans le bundle.

    V139 : cette fonction ne doit être appelée qu'après choix de source officielle,
    sinon elle peut écraser un fichier ennodiagnostic_report.json plus récent.
    """
    bundle = dict(bundle or {})

    if latest_run and latest_run.raw_result_json:
        raw = sanitize_json_value(latest_run.raw_result_json)
        bundle["run_raw_result_json"] = raw

        report, _content = _extract_latest_run_report(latest_run)
        if report:
            bundle["report"] = report
            bundle["report_source"] = "db_latest_run"

    return sanitize_json_value(bundle)


@router.get("/projects/{project_id}/diagnostic/latest")
def get_latest_diagnostic(
    project_id: int,
    compact: bool = Query(False),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Retourne une seule sortie officielle au frontend.

    V140 :
    - choisit le rapport le plus récent ;
    - si ce rapport vient du fichier JSON, il est matérialisé une seule fois en DB ;
    - synchronise les verrous avant de les retourner ;
    - le frontend reçoit uniquement des ids positifs et peut enregistrer les décisions.
    """
    project = get_project_for_user(db, project_id, current_user)
    latest_run = _latest_run_for_project(db, project.id)

    base_bundle = read_diagnostic_bundle(project, compact=compact)
    if compact and _as_dict(base_bundle.get("report")):
        # La page n'a besoin que du rapport fichier officiel et des verrous DB.
        # Lire/materialiser le JSONB complet pendant un GET pouvait dépasser
        # 30 secondes et rendait l'écran inutilisable.
        bundle = dict(base_bundle)
        bundle["official_report_source"] = "filesystem_report_compact_read"
        bundle["official_report_timestamp"] = _report_timestamp(
            _as_dict(bundle.get("report")),
            _path_mtime(bundle.get("report_path_used")),
        )
        bundle["official_report_debug"] = {
            "file_report_path": bundle.get("report_path_used"),
            "db_run_id": latest_run.id if latest_run else None,
            "used": "filesystem_report_compact_read",
        }
    else:
        bundle = _choose_latest_report_source(
            base_bundle,
            latest_run,
            include_run_raw=not compact,
        )
    official_source = bundle.get("official_report_source")

    display = build_diagnostic_display(project, bundle)

    latest_verrous: list[Verrou] = []
    auto_materialized = False

    if compact:
        # Lecture pure : aucune resynchronisation ni matérialisation coûteuse
        # ne doit bloquer le premier affichage.
        latest_verrous = get_latest_diagnostic_verrous(db, project.id)
    elif official_source == "filesystem_report":
        materialized_run, latest_verrous, auto_materialized = _materialize_filesystem_report_in_db(
            db=db,
            project=project,
            bundle=bundle,
        )
        if materialized_run is not None:
            latest_run = materialized_run
            official_source = "filesystem_report_materialized_in_db"
    elif latest_run is not None:
        latest_verrous = _read_latest_run_verrous_fast(db, latest_run)
        if not latest_verrous:
            report, _content = _extract_latest_run_report(latest_run)
            if _extract_final_accepted_verrous_from_report(report):
                latest_verrous = _sync_final_verrous_from_run(db, latest_run)

    latest_run_dump = (
        {
            "id": latest_run.id,
            "project_id": latest_run.project_id,
            "status": latest_run.status,
            "report_path": latest_run.report_path,
            "nlp_result_path": latest_run.nlp_result_path,
            "selected_verrous_path": latest_run.selected_verrous_path,
            "created_at": latest_run.created_at,
            "completed_at": latest_run.completed_at,
        }
        if compact and latest_run is not None
        else DiagnosticRead.model_validate(latest_run).model_dump()
        if latest_run is not None
        else None
    )

    db_verrous_dump = _dump_verrous_for_frontend(latest_verrous)

    # Conserve la version JSON pour audit, mais elle ne pilote plus les boutons.
    json_verrous = (
        display.get("validation_verrous")
        or display.get("validation_verrous_preview")
        or display.get("consultant_verrous_cir")
        or display.get("llm_reformulated_verrous")
        or []
    )
    if not isinstance(json_verrous, list):
        json_verrous = []

    if db_verrous_dump:
        display["json_llm_reformulated_verrous"] = json_verrous
        display["validation_verrous"] = db_verrous_dump
        display["validation_verrous_preview"] = db_verrous_dump
        display["consultant_verrous_cir"] = db_verrous_dump
        display["llm_reformulated_verrous"] = db_verrous_dump
        display["consultant_validation_source"] = "db_synced_verrous_v143"
        display["consultant_validation_enabled"] = True

        frontend_verrous = db_verrous_dump
        sync_info = {
            "ok": True,
            "run_id": latest_run.id if latest_run else None,
            "count": len(db_verrous_dump),
            "source": "db_synced_verrous_v143",
            "auto_materialized_from_filesystem": auto_materialized,
            "note": "Tous les verrous affichés possèdent un id PostgreSQL positif.",
        }
    else:
        # Aucun faux id négatif n'est renvoyé comme verrou décisionnel.
        display["validation_verrous"] = []
        display["validation_verrous_preview"] = json_verrous
        display["consultant_verrous_cir"] = []
        display["llm_reformulated_verrous"] = json_verrous
        display["consultant_validation_source"] = "json_preview_without_db_id"
        display["consultant_validation_enabled"] = False

        frontend_verrous = []
        sync_info = {
            "ok": False,
            "run_id": latest_run.id if latest_run else None,
            "count": 0,
            "preview_count": len(json_verrous),
            "source": "json_preview_without_db_id",
            "note": "Aucun verrou décisionnel n'est retourné sans id DB.",
        }

    display["verrous_sync"] = sync_info
    display["official_report_debug"] = bundle.get("official_report_debug") or {}
    display["official_report_source"] = official_source
    display["display_source_policy"] = {
        "single_source_of_truth": "backend",
        "sections": "display.report_sections / display.diagnostic_sections_by_key",
        "verrous": "display.validation_verrous",
        "frontend_rule": "Le frontend ne doit créer aucun id négatif pour une décision consultant.",
    }
    persisted = (
        {}
        if compact
        else _as_dict(getattr(latest_run, "raw_result_json", {}) if latest_run else {})
    )
    persisted_snapshot = _as_dict(
        bundle.get("diagnostic_snapshot") or persisted.get("diagnostic_snapshot")
    )
    display["database_persistence"] = {
        "ok": bool(latest_run and persisted_snapshot and db_verrous_dump),
        "version": persisted.get("persistence_version"),
        "run_id": latest_run.id if latest_run else None,
        "sections_count": persisted_snapshot.get("sections_count", 0),
        "section_titles_count": persisted_snapshot.get("section_titles_count", 0),
        "cards_count": persisted_snapshot.get("diagnostic_cards_count", 0),
        "verrous_count": len(db_verrous_dump),
        "all_sections_saved_in_raw_result_json": bool(persisted_snapshot),
    }

    if compact:
        display = build_compact_diagnostic_display(display)
    response_verrous = (
        display.get("validation_verrous", [])
        if compact
        else frontend_verrous
    )

    return sanitize_json_value(
        {
            "project": {
                "id": project.id,
                "organisme": project.organisme,
                "project_name": project.project_name,
                "year": project.year,
                "domain_label": project.domain_label,
                "status": project.status,
            },
            "latest_run": latest_run_dump,
            "bundle": {} if compact else bundle,
            "display": display,
            "validation_verrous": response_verrous,
            "source_policy": {
                "diagnostic_display_source": "backend_display_service_v143",
                "official_report_source": official_source,
                "validation_source": display.get("consultant_validation_source"),
                "note": "Le backend matérialise le rapport officiel et renvoie des verrous DB décisionnels.",
            },
        }
    )


@router.post("/projects/{project_id}/diagnostic/import-existing", response_model=DiagnosticRead)
def import_existing_diagnostic(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Importe le rapport complet et ses verrous dans une transaction unique."""
    project = get_project_for_user(db, project_id, current_user)
    try:
        return create_diagnostic_run_from_files(db, project)
    except RuntimeError as exc:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=str(exc))




def _load_document_compare_module():
    r"""
    Charge DOCUMENT_COMPARE de manière robuste.

    Pourquoi :
    quand le backend est lancé depuis C:\EnnoSmart\backend_api, Python peut voir
    le package modules mais pas le sous-package DOCUMENT_COMPARE si le dossier
    n'est pas bien déclaré/copié. Cette fonction tente :
    1) import normal modules.DOCUMENT_COMPARE.document_compare
    2) import direct depuis C:\EnnoSmart\modules\DOCUMENT_COMPARE\document_compare.py
    """
    try:
        ensure_ennosmart_imports()
    except Exception:
        pass

    try:
        from modules.DOCUMENT_COMPARE import document_compare as mod
        return mod
    except Exception as import_error:
        # Import direct par chemin fichier.
        candidates = []

        try:
            # C:\EnnoSmart\backend_api\routers\diagnostic.py -> C:\EnnoSmart
            candidates.append(Path(__file__).resolve().parents[2] / "modules" / "DOCUMENT_COMPARE" / "document_compare.py")
        except Exception:
            pass

        candidates.append(Path(r"C:\EnnoSmart\modules\DOCUMENT_COMPARE\document_compare.py"))

        for path in candidates:
            if path.exists():
                spec = importlib.util.spec_from_file_location("ennosmart_document_compare_direct", str(path))
                if spec and spec.loader:
                    mod = importlib.util.module_from_spec(spec)
                    spec.loader.exec_module(mod)
                    return mod

        raise RuntimeError(
            "Module DOCUMENT_COMPARE introuvable. Copie le dossier "
            "C:\\EnnoSmart\\modules\\DOCUMENT_COMPARE avec __init__.py et document_compare.py. "
            f"Erreur import initiale : {import_error}"
        )



def _safe_upload_filename(filename: str) -> str:
    filename = Path(filename or "document").name
    filename = re.sub(r"[^\wÀ-ÿ ._()\\-]+", "_", filename, flags=re.UNICODE)
    filename = re.sub(r"_+", "_", filename).strip("._ ")
    return filename or "document"


def _manual_compare_upload_dir(project) -> Path:
    from services.diagnostic_service import get_project_store

    ps = get_project_store(project)
    root = ps.project_dir / "document_compare" / "manual_uploads"
    root.mkdir(parents=True, exist_ok=True)

    stamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    out = root / stamp
    out.mkdir(parents=True, exist_ok=True)
    return out


def _save_upload_file(upload: UploadFile, target_dir: Path, prefix: str) -> Path:
    if not upload or not upload.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Fichier {prefix} manquant.",
        )

    safe = _safe_upload_filename(upload.filename)
    target = target_dir / f"{prefix}__{safe}"

    with target.open("wb") as buffer:
        shutil.copyfileobj(upload.file, buffer)

    return target


def _resolve_project_raw_file(raw_dir, value: str):
    """
    Accepte soit un chemin complet issu de l'index, soit un nom de fichier.
    Empêche de comparer un fichier hors du dossier raw du projet.
    """
    from pathlib import Path

    if not value:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Chemin de fichier manquant.",
        )

    raw_root = Path(raw_dir).resolve()
    p = Path(value)

    if not p.is_absolute():
        p = raw_root / value

    p = p.resolve()

    try:
        p.relative_to(raw_root)
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Fichier hors du dossier raw du projet.",
        )

    if not p.exists() or not p.is_file():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Fichier introuvable : {p.name}",
        )

    return p


@router.get("/projects/{project_id}/diagnostic/document-compare/auto-pairs")
def get_document_compare_pairs(
    project_id: int,
    min_similarity: float = Query(0.70, ge=0.0, le=1.0),
    include_medium: bool = Query(True),
    force: bool = Query(False),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Retourne l'index des paires de documents bruts comparables.
    Si l'index n'existe pas ou force=true, il le recrée.
    """
    project = get_project_for_user(db, project_id, current_user)

    try:
        from services.diagnostic_service import get_project_store
        doc_compare = _load_document_compare_module()
        auto_compare_project_pairs = doc_compare.auto_compare_project_pairs

        ps = get_project_store(project)
        raw_dir = ps.documents_raw_dir
        output_dir = ps.project_dir / "document_compare"
        index_path = output_dir / "auto_compare_index.json"

        if index_path.exists() and not force:
            try:
                data = json.loads(index_path.read_text(encoding="utf-8"))
                return sanitize_json_value(data)
            except Exception:
                pass

        report = auto_compare_project_pairs(
            project_uploaded_dir=str(raw_dir),
            output_dir=str(output_dir),
            min_similarity=min_similarity,
            include_medium=include_medium,
            force=force,
        )
        return sanitize_json_value(report)

    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Détection des paires documentaires impossible : {exc}",
        )


@router.post("/projects/{project_id}/diagnostic/document-compare/auto-pairs")
def run_document_compare_pairs(
    project_id: int,
    payload: Dict[str, Any] | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Force la détection des paires de documents bruts comparables.
    """
    project = get_project_for_user(db, project_id, current_user)
    payload = payload or {}

    try:
        from services.diagnostic_service import get_project_store
        doc_compare = _load_document_compare_module()
        auto_compare_project_pairs = doc_compare.auto_compare_project_pairs

        ps = get_project_store(project)
        raw_dir = ps.documents_raw_dir
        output_dir = ps.project_dir / "document_compare"

        report = auto_compare_project_pairs(
            project_uploaded_dir=str(raw_dir),
            output_dir=str(output_dir),
            min_similarity=float(payload.get("min_similarity", 0.70)),
            include_medium=bool(payload.get("include_medium", True)),
            force=bool(payload.get("force", True)),
        )
        return sanitize_json_value(report)

    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Détection des paires documentaires impossible : {exc}",
        )



@router.post("/projects/{project_id}/diagnostic/document-compare/upload-pair")
def upload_and_compare_document_pair(
    project_id: int,
    file_a: UploadFile = File(...),
    file_b: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Mode Streamlit-like :
    l'utilisateur charge manuellement 2 documents A/B,
    puis le backend les compare immédiatement.
    """
    project = get_project_for_user(db, project_id, current_user)

    try:
        doc_compare = _load_document_compare_module()
        compare_pair_to_report = doc_compare.compare_pair_to_report

        upload_dir = _manual_compare_upload_dir(project)
        path_a = _save_upload_file(file_a, upload_dir, "A")
        path_b = _save_upload_file(file_b, upload_dir, "B")

        from services.diagnostic_service import get_project_store
        ps = get_project_store(project)
        output_dir = ps.project_dir / "document_compare" / "manual_reports"
        output_dir.mkdir(parents=True, exist_ok=True)

        report = compare_pair_to_report(
            file_a=str(path_a),
            file_b=str(path_b),
            output_dir=str(output_dir),
            force=True,
        )

        report["manual_upload"] = {
            "ok": True,
            "file_a_original": file_a.filename,
            "file_b_original": file_b.filename,
            "file_a_saved": str(path_a),
            "file_b_saved": str(path_b),
            "upload_dir": str(upload_dir),
            "output_dir": str(output_dir),
        }

        return sanitize_json_value(report)

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Comparaison manuelle impossible : {exc}",
        )



@router.post("/projects/{project_id}/diagnostic/document-compare/compare-pair")
def compare_document_pair(
    project_id: int,
    payload: Dict[str, Any],
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Compare deux documents bruts A/B.
    Le body accepte :
    - file_a + file_b
    ou
    - pair_index, en utilisant auto_compare_index.json.
    """
    project = get_project_for_user(db, project_id, current_user)

    try:
        from services.diagnostic_service import get_project_store
        doc_compare = _load_document_compare_module()
        compare_pair_to_report = doc_compare.compare_pair_to_report
        auto_compare_project_pairs = doc_compare.auto_compare_project_pairs

        ps = get_project_store(project)
        raw_dir = ps.documents_raw_dir
        output_dir = ps.project_dir / "document_compare"
        output_dir.mkdir(parents=True, exist_ok=True)

        file_a = payload.get("file_a")
        file_b = payload.get("file_b")

        if (file_a is None or file_b is None) and payload.get("pair_index") is not None:
            index_path = output_dir / "auto_compare_index.json"
            if index_path.exists():
                index = json.loads(index_path.read_text(encoding="utf-8"))
            else:
                index = auto_compare_project_pairs(
                    project_uploaded_dir=str(raw_dir),
                    output_dir=str(output_dir),
                    min_similarity=0.70,
                    include_medium=True,
                    force=False,
                )

            pairs = index.get("pairs") or []
            pair_index = int(payload.get("pair_index"))
            if pair_index < 0 or pair_index >= len(pairs):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="pair_index invalide.",
                )

            pair = pairs[pair_index]
            file_a = pair.get("file_a")
            file_b = pair.get("file_b")

        path_a = _resolve_project_raw_file(raw_dir, str(file_a or ""))
        path_b = _resolve_project_raw_file(raw_dir, str(file_b or ""))

        report = compare_pair_to_report(
            file_a=str(path_a),
            file_b=str(path_b),
            output_dir=str(output_dir),
            force=bool(payload.get("force", True)),
        )

        return sanitize_json_value(report)

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Comparaison documentaire impossible : {exc}",
        )


@router.post("/projects/{project_id}/diagnostic/cir-memory/compare")
def compare_with_previous_cir(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Lance uniquement la comparaison avec le CIR précédent :
    dossier courant NLP/Frascati vs CIR final mémoire N-1.
    """
    project = get_project_for_user(db, project_id, current_user)

    try:
        from services.diagnostic_service import get_project_store
        from modules.CIR_MEMORY.cir_memory import compare_current_raw_with_cir_memory

        ps = get_project_store(project)
        nlp_path = ps.nlp_dir / "nlp_result.json"

        if not nlp_path.exists():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="nlp_result.json introuvable. Lance d'abord prepare-sources.",
            )

        report = compare_current_raw_with_cir_memory(
            organisme=project.organisme,
            project=project.project_name,
            year=str(project.year),
            nlp_result_path=nlp_path,
        )

        return sanitize_json_value(report)

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Comparaison CIR précédent impossible : {exc}",
        )




@router.post("/projects/{project_id}/cir-previous/compare-current")
def compare_current_with_previous_cir_independent(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Lance uniquement la comparaison CIR précédent.

    Cette route ne lance pas EnnoDiagnostic, ne relance pas le LLM diagnostic,
    ne relance pas le score IA et ne refait pas le NLP.
    Elle compare le nlp_result.json courant déjà préparé avec la mémoire CIR N-1.
    """
    project = get_project_for_user(db, project_id, current_user)

    try:
        from services.diagnostic_service import get_project_store
        from modules.CIR_MEMORY.cir_memory import compare_current_raw_with_cir_memory

        ps = get_project_store(project)
        nlp_path = ps.nlp_dir / "nlp_result.json"

        if not nlp_path.exists():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="nlp_result.json introuvable. Lance d'abord Préparer les sources.",
            )

        report = compare_current_raw_with_cir_memory(
            organisme=project.organisme,
            project=project.project_name,
            year=str(project.year),
            nlp_result_path=nlp_path,
        )

        return sanitize_json_value(report)

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Comparaison CIR précédent impossible : {exc}",
        )


@router.get("/projects/{project_id}/cir-previous/comparison-latest")
def get_latest_previous_cir_comparison(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Lit le dernier rapport de comparaison CIR précédent sauvegardé.
    Ne lance aucun calcul.
    """
    project = get_project_for_user(db, project_id, current_user)

    try:
        from modules.CIR_MEMORY.cir_memory import comparison_report_path

        path = comparison_report_path(project.organisme, project.project_name, str(project.year))
        if not path.exists():
            return sanitize_json_value({
                "ok": False,
                "missing": True,
                "has_previous_cir": False,
                "message": "Aucune comparaison CIR précédent sauvegardée pour ce projet.",
                "report_path": str(path),
            })

        try:
            report = json.loads(path.read_text(encoding="utf-8"))
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Rapport CIR précédent illisible : {exc}",
            )

        if isinstance(report, dict):
            report["report_path"] = str(path)
            report["loaded_from_saved_report"] = True

        return sanitize_json_value(report)

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Lecture comparaison CIR précédent impossible : {exc}",
        )

@router.post("/projects/{project_id}/diagnostic/prepare-sources")
def prepare_diagnostic_sources(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Étape 1 :
    upload/raw documents -> extraction -> NLP -> Frascati -> RAG/Chroma.

    Ne lance pas le LLM.
    """
    project = get_project_for_user(db, project_id, current_user)
    return prepare_ennodiagnostic_sources(db, project)


@router.post("/projects/{project_id}/diagnostic/run-agent", response_model=DiagnosticRead)
def run_diagnostic_agent_only(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Lance l'agent puis sauvegarde atomiquement :
    rapport complet, toutes les sections et verrous avec ids DB positifs.
    """
    project = get_project_for_user(db, project_id, current_user)
    try:
        return run_ennodiagnostic_agent_only(db, project)
    except RuntimeError as exc:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=str(exc))


@router.post("/projects/{project_id}/diagnostic/run", response_model=DiagnosticRead)
def run_diagnostic(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Pipeline complet avec persistance DB complète et atomique V142."""
    project = get_project_for_user(db, project_id, current_user)
    try:
        return run_ennodiagnostic(db, project)
    except RuntimeError as exc:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=str(exc))


@router.post("/projects/{project_id}/diagnostic/{run_id}/sync-verrous", response_model=list[VerrouRead])
def sync_verrous(
    project_id: int,
    run_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    project = get_project_for_user(db, project_id, current_user)

    run = (
        db.query(DiagnosticRun)
        .filter(DiagnosticRun.id == run_id, DiagnosticRun.project_id == project.id)
        .first()
    )

    if not run:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Diagnostic run introuvable.",
        )

    return sync_verrous_from_diagnostic(db, run)


@router.get("/projects/{project_id}/verrous", response_model=list[VerrouRead])
def list_verrous(
    project_id: int,
    latest_only: bool = Query(True),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    project = get_project_for_user(db, project_id, current_user)

    query = (
        db.query(Verrou)
        .join(DiagnosticRun, Verrou.diagnostic_run_id == DiagnosticRun.id)
        .filter(DiagnosticRun.project_id == project.id)
    )

    if latest_only:
        # Un GET reste une lecture pure et inclut les ajouts humains historiques.
        current = get_latest_diagnostic_verrous(db, project.id)

        # Retour direct pour garantir les vrais ids du run officiel.
        seen: set[str] = set()
        clean: list[Verrou] = []
        for verrou in current:
            key = _normalize_title_for_match(getattr(verrou, "title", ""))
            if not key or key in seen:
                continue
            seen.add(key)
            clean.append(verrou)
        return clean

    verrous = query.order_by(Verrou.created_at.desc()).all()

    seen: set[str] = set()
    clean: list[Verrou] = []

    for verrou in verrous:
        key = (verrou.title or "").strip().lower()
        if not key or key in seen:
            continue
        seen.add(key)
        clean.append(verrou)

    return clean


@router.post(
    "/projects/{project_id}/verrous/manual",
    response_model=VerrouRead,
    status_code=status.HTTP_201_CREATED,
)
def create_manual_verrou(
    project_id: int,
    payload: VerrouManualCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    project = get_project_for_user(db, project_id, current_user)
    result = create_or_reuse_consultant_verrou(
        db,
        project,
        title=payload.title,
        justification=payload.description,
        supporting_context=payload.description,
        created_by_user_id=int(current_user.id),
        force_create_distinct=payload.force_create_distinct,
        keywords=payload.keywords,
        added_via="ennodiagnostic_manual_form",
    )

    if not result.get("ok"):
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail=result,
        )

    verrou = (
        db.query(Verrou)
        .join(DiagnosticRun, Verrou.diagnostic_run_id == DiagnosticRun.id)
        .filter(
            Verrou.id == int(result["verrou_id"]),
            DiagnosticRun.project_id == project.id,
        )
        .first()
    )
    if verrou is None:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Le verrou manuel a été créé mais ne peut pas être relu.",
        )
    return verrou


@router.patch("/projects/{project_id}/verrous/{verrou_id}/decision", response_model=VerrouRead)
def update_verrou_decision(
    project_id: int,
    verrou_id: int,
    payload: VerrouDecisionRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    project = get_project_for_user(db, project_id, current_user)

    verrou = (
        db.query(Verrou)
        .join(DiagnosticRun, Verrou.diagnostic_run_id == DiagnosticRun.id)
        .filter(Verrou.id == verrou_id, DiagnosticRun.project_id == project.id)
        .first()
    )

    # V141 : compatibilité avec les ids négatifs déjà présents dans un onglet
    # frontend ouvert avant la synchronisation DB. Le backend résout cet id
    # historique vers le titre du rapport officiel, synchronise, puis applique
    # la décision sur le vrai verrou PostgreSQL.
    legacy_negative_id = verrou_id if verrou_id < 0 else None
    if verrou is None and legacy_negative_id is not None:
        verrou = _resolve_negative_frontend_verrou(
            db=db,
            project=project,
            negative_id=legacy_negative_id,
        )

    if not verrou:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=(
                "Verrou introuvable après tentative de synchronisation automatique. "
                "Recharge le diagnostic pour récupérer les ids PostgreSQL."
            ),
        )

    allowed = {"garde", "rejete", "reformuler", "en_attente"}
    if payload.consultant_status not in allowed:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Statut invalide. Valeurs autorisées : {sorted(allowed)}",
        )

    verrou.consultant_status = payload.consultant_status
    db.commit()
    db.refresh(verrou)
    return verrou


# ============================================================
# CIR précédent / mémoire CIR finale
# ============================================================

def _extract_text_from_cir_file(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument
        except Exception as exc:
            raise RuntimeError("python-docx est requis pour lire les fichiers DOCX.") from exc

        doc = DocxDocument(str(path))
        parts: list[str] = []

        for p in doc.paragraphs:
            txt = (p.text or "").strip()
            if txt:
                parts.append(txt)

        for table in doc.tables:
            for row in table.rows:
                cells = [clean for cell in row.cells if (clean := (cell.text or "").strip())]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    if suffix == ".pdf":
        # On essaie d'abord pypdf, puis PyPDF2 si l'environnement l'utilise.
        reader = None
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
        except Exception:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(str(path))
            except Exception as exc:
                raise RuntimeError("pypdf ou PyPDF2 est requis pour lire les fichiers PDF.") from exc

        parts = []
        for page in reader.pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            if txt.strip():
                parts.append(txt.strip())
        return "\n".join(parts)

    raise RuntimeError("Format non supporté. Utilise PDF, DOCX ou TXT pour le CIR final précédent.")


def _role_pack_for_cir_text(text: str) -> str:
    low = text.lower()

    if any(k in low for k in ["verrou", "incertitude", "difficulté", "difficulte", "problématique", "problematique", "non transfér", "non transfer", "non transposable"]):
        return "verrous_rnd_locaux"

    if any(k in low for k in ["objectif", "performances à atteindre", "performances a atteindre", "vise à", "vise a", "débit", "debit", "300 bars", "point de rosée", "point de rosee"]):
        return "objectifs_locaux"

    if any(k in low for k in ["état de l’art", "etat de l'art", "littérature", "litterature", "brevet", "article scientifique", "solutions existantes", "insuffisance"]):
        return "etat_art_local"

    if any(k in low for k in ["essai", "essais", "simulation", "mesure", "relevé", "releve", "modélisation", "modelisation", "calcul", "analyse", "prototype", "développement", "developpement"]):
        return "methodes_locales"

    if any(k in low for k in ["résultat", "resultat", "conclusion", "montré", "montre", "constaté", "constate", "permis", "atteint", "réduit", "reduit", "validé", "valide"]):
        return "resultats_locaux"

    if any(k in low for k in ["contrainte", "exigence", "limite", "non conforme", "insuffisant", "risque"]):
        return "limites_locales"

    if re.search(r"\b\d+(?:[,.]\d+)?\s*(bar|bars|kg|mm|°c|db|hz|rpm|m3/h|%)\b", low):
        return "parametres_locaux"

    return "contributions_locales"


def _split_cir_final_into_items(text: str, filename: str) -> Dict[str, Any]:
    """
    Extraction légère pour mémoire CIR final : pas de Frascati, pas de détection de verrous nouveaux.
    On structure seulement les passages du CIR final N-1 pour que la comparaison N vs N-1 puisse fonctionner.
    """
    pack_keys = [
        "objectifs_locaux",
        "verrous_rnd_locaux",
        "methodes_locales",
        "resultats_locaux",
        "limites_locales",
        "contributions_locales",
        "etat_art_local",
        "parametres_locaux",
    ]
    role_by_pack = {
        "objectifs_locaux": "objectif",
        "verrous_rnd_locaux": "verrou",
        "methodes_locales": "methode",
        "resultats_locaux": "resultat",
        "limites_locales": "limite",
        "contributions_locales": "contribution",
        "etat_art_local": "etat_art",
        "parametres_locaux": "parametre",
    }

    pack: Dict[str, list[dict[str, Any]]] = {k: [] for k in pack_keys}
    items: list[dict[str, Any]] = []

    cleaned = re.sub(r"\r\n?", "\n", text or "")
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    paragraphs = [p.strip() for p in re.split(r"\n{2,}|(?=^\d+(?:\.\d+)*\.\s+)", cleaned, flags=re.M)]

    current_title = "CIR final précédent"
    counter = 0
    seen: set[str] = set()

    for raw in paragraphs:
        p = raw.strip()
        if not p:
            continue

        first_line = p.split("\n", 1)[0].strip()
        if re.match(r"^\d+(?:\.\d+)*\.?\s+.{4,120}$", first_line):
            current_title = first_line[:180]

        # On évite les pages, pieds de page et bouts trop courts.
        p = re.sub(r"(?i).*confidentiel\s*page\s*\d+", "", p).strip()
        p = re.sub(r"(?i)Ce document est la propriété.*", "", p).strip()
        if len(p) < 80:
            continue

        key = re.sub(r"\W+", " ", p.lower())[:260]
        if key in seen:
            continue
        seen.add(key)

        pack_key = _role_pack_for_cir_text(p)
        role = role_by_pack.get(pack_key, "general")
        counter += 1

        item = {
            "id": f"cir_prev_{counter}",
            "role": role,
            "pack_key": pack_key,
            "text": p[:2500],
            "document": filename,
            "section_title": current_title,
            "section_type": "cir_final_precedent",
            "section_label": current_title,
            "source_type": "previous_cir_final_without_frascati",
            "content_origin": "cir_final_uploaded_by_consultant",
            "quality_status": "memory_only_no_frascati",
        }
        pack[pack_key].append(item)
        items.append(item)

        if len(items) >= 180:
            break

    return {"pack": pack, "items": items}


def _roles_count(items: list[dict[str, Any]]) -> Dict[str, int]:
    out: Dict[str, int] = {}
    for item in items:
        role = str(item.get("role") or "unknown")
        out[role] = out.get(role, 0) + 1
    return dict(sorted(out.items()))


@router.post("/projects/{project_id}/cir-previous/upload-final")
def upload_previous_cir_final(
    project_id: int,
    year: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Ajoute un CIR final N-1 comme mémoire CIR.

    Important : ce fichier n'est PAS traité comme document brut de l'année courante.
    Il sert uniquement à comparer le projet courant avec le CIR final précédent.
    """
    project = get_project_for_user(db, project_id, current_user)
    year = str(year or "").strip()

    if not re.fullmatch(r"\d{4}", year):
        raise HTTPException(status_code=400, detail="L'année du CIR précédent doit être au format YYYY, par exemple 2022.")

    if str(year) == str(project.year):
        raise HTTPException(status_code=400, detail="Le CIR précédent doit avoir une année différente de l'année du projet courant.")

    ensure_ennosmart_imports()

    try:
        from modules.CIR_MEMORY.cir_memory import cir_final_dir, cir_final_report_path, comparison_report_path, compare_current_raw_with_cir_memory, write_json
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Module CIR_MEMORY indisponible : {exc}")

    safe_name = _safe_upload_filename(file.filename or f"cir_final_{year}")
    suffix = Path(safe_name).suffix.lower()
    if suffix not in {".pdf", ".docx", ".txt"}:
        raise HTTPException(status_code=400, detail="Format non supporté. Utilise PDF, DOCX ou TXT.")

    raw_dir = cir_final_dir(project.organisme, project.project_name, year) / "raw"
    raw_dir.mkdir(parents=True, exist_ok=True)
    raw_path = raw_dir / safe_name

    with raw_path.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    try:
        text = _extract_text_from_cir_file(raw_path)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Lecture du CIR final impossible : {exc}")

    if len((text or "").strip()) < 200:
        raise HTTPException(status_code=400, detail="Le texte extrait du CIR final est trop court. Vérifie que le fichier n'est pas scanné sans OCR.")

    structured = _split_cir_final_into_items(text, safe_name)
    items = structured["items"]
    pack = structured["pack"]

    if not items:
        raise HTTPException(status_code=400, detail="Aucun passage exploitable n'a été extrait du CIR final précédent.")

    report = {
        "ok": True,
        "version": "cir_previous_upload_front_backend_v40",
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "organisme": project.organisme,
        "project": project.project_name,
        "year": year,
        "current_project_year": str(project.year),
        "cir_final_file": str(raw_path),
        "rule": "CIR final précédent = mémoire CIR uniquement, sans FrascatiGuard et sans injection comme document brut courant.",
        "items_count": len(items),
        "roles": _roles_count(items),
        "items": items,
        "evidence_pack_before_frascati": pack,
    }

    out_path = cir_final_report_path(project.organisme, project.project_name, year)
    write_json(out_path, sanitize_json_value(report))

    # Écrit aussi un nlp_result mémoire pour audit humain.
    nlp_memory_path = cir_final_dir(project.organisme, project.project_name, year) / "cir_final_nlp_memory.json"
    write_json(nlp_memory_path, sanitize_json_value({
        "ok": True,
        "pipeline_type": "cir_final_memory_without_frascati",
        "source_file": str(raw_path),
        "evidence_pack_before_frascati": pack,
        "items": items,
    }))

    comparison = None
    try:
        current_nlp = diagnostic_paths(project)["nlp_result"]
        if current_nlp.exists():
            comparison = compare_current_raw_with_cir_memory(
                organisme=project.organisme,
                project=project.project_name,
                year=str(project.year),
                nlp_result_path=current_nlp,
            )
    except Exception as exc:
        comparison = {
            "ok": False,
            "error": str(exc),
            "note": "Le CIR précédent est enregistré. Relance Préparer les sources puis EnnoDiagnostic pour recalculer la comparaison.",
        }

    return sanitize_json_value({
        "ok": True,
        "message": "CIR final précédent enregistré comme mémoire CIR.",
        "previous_cir_year": year,
        "file": str(raw_path),
        "report_path": str(out_path),
        "items_count": len(items),
        "roles": _roles_count(items),
        "comparison_after_upload": comparison,
    })


@router.get("/projects/{project_id}/cir-previous")
def list_previous_cir_finals(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    project = get_project_for_user(db, project_id, current_user)
    ensure_ennosmart_imports()

    try:
        from modules.CIR_MEMORY.cir_memory import STORAGE_DIR, slug, cir_final_report_path
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Module CIR_MEMORY indisponible : {exc}")

    years_root = STORAGE_DIR / slug(project.organisme) / "projects" / slug(project.project_name) / "years"
    if not years_root.exists():
        return {"ok": True, "items": []}

    items = []
    for year_dir in sorted([p for p in years_root.iterdir() if p.is_dir()], reverse=True):
        year = year_dir.name
        if str(year) == str(project.year):
            continue
        report_path = cir_final_report_path(project.organisme, project.project_name, year)
        if not report_path.exists():
            continue
        try:
            report = json.loads(report_path.read_text(encoding="utf-8"))
        except Exception:
            report = {}
        items.append({
            "year": year,
            "ok": bool(report.get("ok")),
            "items_count": report.get("items_count"),
            "roles": report.get("roles") or {},
            "file": report.get("cir_final_file"),
            "report_path": str(report_path),
            "generated_at": report.get("generated_at"),
        })

    return sanitize_json_value({"ok": True, "items": items})
