# -*- coding: utf-8 -*-
from __future__ import annotations

"""
Router FastAPI V70 - Prévisualisation source avec surlignage

But :
- Le frontend envoie l'extrait cliqué.
- Le backend retrouve le document complet lié à cet extrait.
- Si PDF : il génère une copie PDF avec le passage surligné / encadré.
- Si DOCX/TXT/MD : il génère une page HTML avec le passage encadré.
- Si document introuvable : il renvoie une page HTML fallback avec l'extrait.

Endpoints :
GET  /projects/{project_id}/source-highlight/preview?excerpt=...&source_path=...&source_name=...
POST /projects/{project_id}/source-highlight/preview

Dépendances conseillées :
    pip install pymupdf python-docx

Même sans python-docx, le DOCX est lu via zip/xml en fallback.
"""

import hashlib
import html
import json
import os
import re
import shutil
import unicodedata
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple
from urllib.parse import quote

from fastapi import APIRouter, Body, HTTPException, Query, Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, RedirectResponse
from pydantic import BaseModel

router = APIRouter(tags=["Source preview / highlight"])

# Racine du projet : .../C:\EnnoSmart si ce fichier est dans backend_api/routers
PROJECT_ROOT = Path(__file__).resolve().parents[2]
STORAGE_ROOT = PROJECT_ROOT / "storage"
OUTPUTS_ROOT = PROJECT_ROOT / "outputs"
PREVIEW_ROOT = STORAGE_ROOT / "previews" / "source_highlight"

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".html", ".htm"}
PDF_EXTENSIONS = {".pdf"}
TEXT_EXTENSIONS = {".txt", ".md", ".html", ".htm"}
DOCX_EXTENSIONS = {".docx"}

MAX_EXCERPT_CHARS = 4000
MAX_HTML_TEXT_CHARS = 2_000_000
MAX_SEARCH_FILES = 6000


class SourceHighlightRequest(BaseModel):
    excerpt: str
    source_path: Optional[str] = None
    source_name: Optional[str] = None
    document_name: Optional[str] = None
    title: Optional[str] = None
    role: Optional[str] = None
    organisme: Optional[str] = None
    project_name: Optional[str] = None
    year: Optional[str] = None
    return_json: bool = False


def _ensure_preview_root() -> None:
    PREVIEW_ROOT.mkdir(parents=True, exist_ok=True)


def _strip_accents(value: str) -> str:
    return "".join(
        c for c in unicodedata.normalize("NFD", value or "")
        if unicodedata.category(c) != "Mn"
    )


def normalize_text(value: str) -> str:
    value = _strip_accents(str(value or "")).lower()
    value = value.replace("\u00a0", " ").replace("’", "'")
    value = re.sub(r"[^a-z0-9%.,;:/+\-_' ]+", " ", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def clean_excerpt(value: str) -> str:
    text = str(value or "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"#{1,6}\s*", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()[:MAX_EXCERPT_CHARS]


def short_hash(*parts: Any) -> str:
    raw = "||".join(str(p or "") for p in parts)
    return hashlib.sha256(raw.encode("utf-8", errors="ignore")).hexdigest()[:16]


def is_under_root(path: Path, roots: Iterable[Path]) -> bool:
    try:
        resolved = path.resolve()
    except Exception:
        return False

    for root in roots:
        try:
            resolved.relative_to(root.resolve())
            return True
        except Exception:
            continue
    return False


def allowed_roots() -> List[Path]:
    roots = [STORAGE_ROOT, OUTPUTS_ROOT, PROJECT_ROOT / "backend_api" / "storage"]
    return [root for root in roots if root.exists()]


def safe_candidate_path(raw: Optional[str]) -> Optional[Path]:
    if not raw:
        return None

    value = str(raw).strip().strip('"').strip("'")
    if not value:
        return None

    # Normaliser les chemins Windows / URL encodés simples.
    value = value.replace("\\", os.sep).replace("/", os.sep)
    p = Path(value)

    candidates: List[Path] = []
    if p.is_absolute():
        candidates.append(p)
    else:
        for root in allowed_roots():
            candidates.append(root / p)
        candidates.append(PROJECT_ROOT / p)

    for candidate in candidates:
        try:
            resolved = candidate.resolve()
        except Exception:
            continue
        if resolved.exists() and resolved.is_file() and resolved.suffix.lower() in ALLOWED_EXTENSIONS:
            if is_under_root(resolved, [PROJECT_ROOT, STORAGE_ROOT, OUTPUTS_ROOT]):
                return resolved
    return None


def iter_search_roots(project_id: int, payload: SourceHighlightRequest) -> List[Path]:
    roots: List[Path] = []

    # Chemins les plus probables d'EnnoSmart.
    if payload.organisme and payload.project_name and payload.year:
        org = normalize_slug(payload.organisme)
        proj = normalize_slug(payload.project_name)
        roots.append(STORAGE_ROOT / "organismes" / org / "projects" / proj / "years" / str(payload.year))
        roots.append(OUTPUTS_ROOT / "safe_rag_upload" / str(payload.organisme) / str(payload.project_name) / str(payload.year))

    roots.extend([
        STORAGE_ROOT / "projects" / str(project_id),
        STORAGE_ROOT / "organismes",
        OUTPUTS_ROOT / "safe_rag_upload",
        STORAGE_ROOT,
    ])

    unique: List[Path] = []
    seen = set()
    for root in roots:
        try:
            resolved = root.resolve()
        except Exception:
            continue
        if resolved.exists() and resolved.is_dir() and str(resolved).lower() not in seen:
            unique.append(resolved)
            seen.add(str(resolved).lower())
    return unique


def normalize_slug(value: str) -> str:
    slug = normalize_text(value)
    slug = re.sub(r"[^a-z0-9]+", "_", slug).strip("_")
    return slug or "unknown"


def file_name_score(path: Path, wanted_names: List[str]) -> int:
    if not wanted_names:
        return 0

    name = normalize_text(path.name)
    stem = normalize_text(path.stem)
    score = 0
    for wanted in wanted_names:
        w = normalize_text(wanted)
        if not w:
            continue
        if name == w or stem == w:
            score += 100
        elif w in name or w in stem:
            score += 60
        else:
            w_stem = normalize_text(Path(wanted).stem)
            if w_stem and (w_stem in stem or stem in w_stem):
                score += 40
    return score


def get_candidate_files(roots: List[Path], wanted_names: List[str]) -> List[Path]:
    files: List[Path] = []
    seen = set()

    # D'abord chercher par nom si possible.
    for root in roots:
        count = 0
        try:
            iterator = root.rglob("*")
            for path in iterator:
                if count > MAX_SEARCH_FILES:
                    break
                count += 1
                if not path.is_file() or path.suffix.lower() not in ALLOWED_EXTENSIONS:
                    continue
                key = str(path.resolve()).lower()
                if key in seen:
                    continue
                if wanted_names and file_name_score(path, wanted_names) <= 0:
                    continue
                files.append(path)
                seen.add(key)
        except Exception:
            continue

    # Si rien par nom, prendre un échantillon raisonnable de docs.
    if not files:
        for root in roots:
            count = 0
            try:
                for path in root.rglob("*"):
                    if count > MAX_SEARCH_FILES:
                        break
                    count += 1
                    if path.is_file() and path.suffix.lower() in ALLOWED_EXTENSIONS:
                        key = str(path.resolve()).lower()
                        if key not in seen:
                            files.append(path)
                            seen.add(key)
            except Exception:
                continue

    return sorted(files, key=lambda p: (-file_name_score(p, wanted_names), len(str(p))))


def read_text_file(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=enc, errors="ignore")
        except Exception:
            continue
    return ""


def read_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
        doc = Document(str(path))
        parts: List[str] = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception:
        pass

    # Fallback sans python-docx : lecture XML.
    try:
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        xml = re.sub(r"</w:p>", "\n", xml)
        xml = re.sub(r"<[^>]+>", " ", xml)
        xml = html.unescape(xml)
        xml = re.sub(r"\s+", " ", xml)
        return xml.strip()
    except Exception:
        return ""


def read_pdf_text(path: Path, max_pages: int = 80) -> str:
    try:
        import fitz  # type: ignore
    except Exception:
        return ""

    try:
        doc = fitz.open(str(path))
        chunks: List[str] = []
        for page_index, page in enumerate(doc):
            if page_index >= max_pages:
                break
            chunks.append(page.get_text("text") or "")
        doc.close()
        return "\n".join(chunks)
    except Exception:
        return ""


def read_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return read_text_file(path)
    if suffix in DOCX_EXTENSIONS:
        return read_docx_text(path)
    if suffix in PDF_EXTENSIONS:
        return read_pdf_text(path)
    return ""


def overlap_score(excerpt: str, text: str) -> float:
    ne = normalize_text(excerpt)
    nt = normalize_text(text)
    if not ne or not nt:
        return 0.0
    if ne in nt:
        return 1.0

    words = [w for w in re.findall(r"[a-z0-9]{4,}", ne) if len(w) >= 4]
    if not words:
        return 0.0
    unique_words = list(dict.fromkeys(words))[:80]
    hits = sum(1 for w in unique_words if w in nt)
    return hits / max(1, len(unique_words))


def find_best_document(project_id: int, payload: SourceHighlightRequest) -> Tuple[Optional[Path], str, float]:
    direct = safe_candidate_path(payload.source_path)
    if direct:
        text = read_document_text(direct)
        score = overlap_score(payload.excerpt, text) if text else 0.5
        return direct, text, score

    wanted_names = [
        payload.source_name or "",
        payload.document_name or "",
        Path(payload.source_path or "").name if payload.source_path else "",
    ]
    wanted_names = [x for x in wanted_names if x]

    roots = iter_search_roots(project_id, payload)
    candidates = get_candidate_files(roots, wanted_names)

    best_path: Optional[Path] = None
    best_text = ""
    best_score = -1.0

    for path in candidates:
        text = read_document_text(path)
        name_bonus = min(file_name_score(path, wanted_names) / 200.0, 0.5)
        content_score = overlap_score(payload.excerpt, text) if text else 0.0
        score = content_score + name_bonus
        if score > best_score:
            best_path = path
            best_text = text
            best_score = score
        if content_score >= 0.90 and name_bonus > 0:
            break

    if best_path and best_score > 0:
        return best_path, best_text, best_score
    return None, "", 0.0


def make_search_queries(excerpt: str) -> List[str]:
    text = clean_excerpt(excerpt)
    if not text:
        return []

    candidates: List[str] = []
    candidates.append(text[:250])

    # Phrases et fenêtres plus courtes pour PyMuPDF search_for.
    sentences = re.split(r"(?<=[.!?;:])\s+", text)
    for s in sentences:
        s = s.strip()
        if 25 <= len(s) <= 180:
            candidates.append(s)

    words = text.split()
    for size in (18, 14, 10, 7):
        if len(words) >= size:
            for start in range(0, min(len(words) - size + 1, 30), max(1, size // 2)):
                q = " ".join(words[start:start + size]).strip()
                if len(q) >= 25:
                    candidates.append(q)

    # Dédupliquer.
    unique: List[str] = []
    seen = set()
    for c in candidates:
        c = re.sub(r"\s+", " ", c).strip()
        key = normalize_text(c)
        if c and key not in seen:
            unique.append(c)
            seen.add(key)
    return unique[:25]


def best_pdf_page_by_overlap(doc: Any, excerpt: str) -> int:
    best_page = 0
    best_score = -1.0
    for i, page in enumerate(doc):
        try:
            text = page.get_text("text") or ""
            score = overlap_score(excerpt, text)
            if score > best_score:
                best_score = score
                best_page = i
        except Exception:
            continue
    return best_page


def create_highlighted_pdf(path: Path, excerpt: str, project_id: int) -> Tuple[Optional[Path], bool, str]:
    try:
        import fitz  # type: ignore
    except Exception:
        return None, False, "PyMuPDF n'est pas installé. Lance : pip install pymupdf"

    _ensure_preview_root()
    out = PREVIEW_ROOT / f"source_highlight_p{project_id}_{short_hash(path, excerpt)}.pdf"

    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        return None, False, f"Impossible d'ouvrir le PDF : {exc}"

    found_rects = 0
    found_page = 0
    queries = make_search_queries(excerpt)

    try:
        for page_index, page in enumerate(doc):
            page_found = False
            for query in queries:
                try:
                    rects = page.search_for(query)
                except Exception:
                    rects = []

                if not rects:
                    continue

                for rect in rects[:10]:
                    try:
                        highlight = page.add_highlight_annot(rect)
                        highlight.set_colors(stroke=(1, 0.86, 0.20))
                        highlight.update()
                    except Exception:
                        pass

                    try:
                        box = page.add_rect_annot(rect)
                        box.set_colors(stroke=(1, 0.42, 0.0))
                        box.set_border(width=1.2)
                        box.update()
                    except Exception:
                        pass

                    found_rects += 1
                    page_found = True

                if page_found:
                    found_page = page_index
                    break

            if found_rects > 0:
                break

        # Si aucune occurrence exacte : encadrer la page la plus proche et ajouter une note visuelle.
        if found_rects == 0 and len(doc) > 0:
            found_page = best_pdf_page_by_overlap(doc, excerpt)
            page = doc[found_page]
            rect = page.rect
            banner = fitz.Rect(rect.x0 + 28, rect.y0 + 28, rect.x1 - 28, min(rect.y0 + 135, rect.y1 - 28))
            annot = page.add_rect_annot(banner)
            annot.set_colors(stroke=(1, 0.42, 0.0), fill=(1, 0.94, 0.72))
            annot.set_opacity(0.28)
            annot.set_border(width=2)
            annot.update()

            page.insert_textbox(
                banner,
                "Passage recherché - correspondance approximative sur cette page\n" + clean_excerpt(excerpt)[:450],
                fontsize=9,
                color=(0.20, 0.12, 0.02),
                align=0,
            )

        doc.save(str(out), garbage=4, deflate=True)
        doc.close()
        return out, found_rects > 0, f"PDF généré. Page {found_page + 1}. Occurrences exactes : {found_rects}."
    except Exception as exc:
        try:
            doc.close()
        except Exception:
            pass
        return None, False, f"Erreur génération PDF surligné : {exc}"


def html_document(title: str, body: str, status: str = "ok") -> str:
    safe_title = html.escape(title or "Prévisualisation source")
    status_class = "ok" if status == "ok" else "warn"
    return f"""<!doctype html>
<html lang="fr">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{safe_title}</title>
  <style>
    :root {{ color-scheme: light; }}
    body {{ margin:0; font-family: Inter, Arial, sans-serif; background:#f8fafc; color:#0f172a; }}
    .wrap {{ max-width: 1100px; margin: 0 auto; padding: 24px; }}
    .card {{ background:#fff; border:1px solid #e2e8f0; border-radius:16px; box-shadow:0 8px 28px rgba(15,23,42,.06); overflow:hidden; }}
    .head {{ padding:18px 20px; border-bottom:1px solid #e2e8f0; display:flex; gap:12px; align-items:flex-start; justify-content:space-between; }}
    .title {{ font-size:16px; font-weight:750; margin:0; }}
    .sub {{ font-size:12px; color:#64748b; margin:6px 0 0; }}
    .badge {{ display:inline-flex; align-items:center; border-radius:999px; padding:4px 10px; font-size:12px; font-weight:650; border:1px solid #cbd5e1; }}
    .badge.ok {{ background:#ecfdf5; color:#047857; border-color:#a7f3d0; }}
    .badge.warn {{ background:#fff7ed; color:#c2410c; border-color:#fed7aa; }}
    .content {{ padding:20px; }}
    .doc {{ white-space:pre-wrap; line-height:1.72; font-size:14px; }}
    mark, .highlight {{ background:#fef08a; border:2px solid #f97316; border-radius:6px; padding:2px 4px; color:#111827; }}
    .excerpt {{ border:2px solid #f97316; background:#fff7ed; border-radius:12px; padding:14px; white-space:pre-wrap; line-height:1.7; }}
    .muted {{ color:#64748b; font-size:12px; }}
  </style>
</head>
<body>
  <div class="wrap">
    <div class="card">
      <div class="head">
        <div>
          <p class="title">{safe_title}</p>
          <p class="sub">Prévisualisation générée par EnnoSmart - passage encadré quand il est retrouvé.</p>
        </div>
        <span class="badge {status_class}">{'Document trouvé' if status == 'ok' else 'Fallback'}</span>
      </div>
      <div class="content">
        {body}
      </div>
    </div>
  </div>
</body>
</html>"""


def highlight_html_text(full_text: str, excerpt: str) -> Tuple[str, bool]:
    safe_full = html.escape((full_text or "")[:MAX_HTML_TEXT_CHARS])
    clean = clean_excerpt(excerpt)
    if not full_text:
        return f"<div class='excerpt'>{html.escape(clean or 'Aucun extrait disponible.')}</div>", False

    # Recherche exacte d'abord.
    pos = full_text.lower().find(clean.lower()) if clean else -1
    if pos >= 0:
        before = html.escape(full_text[:pos])
        mid = html.escape(full_text[pos:pos + len(clean)])
        after = html.escape(full_text[pos + len(clean):MAX_HTML_TEXT_CHARS])
        return f"<div class='doc'>{before}<mark>{mid}</mark>{after}</div>", True

    # Recherche par phrase/fenêtre.
    for q in make_search_queries(clean):
        pos = full_text.lower().find(q.lower())
        if pos >= 0:
            before = html.escape(full_text[:pos])
            mid = html.escape(full_text[pos:pos + len(q)])
            after = html.escape(full_text[pos + len(q):MAX_HTML_TEXT_CHARS])
            return f"<div class='doc'>{before}<mark>{mid}</mark>{after}</div>", True

    # Fuzzy fallback : afficher le document complet + extrait au-dessus.
    body = (
        "<p class='muted'>Le document est trouvé, mais le passage exact n'a pas été localisé mot à mot. "
        "L'extrait cliqué est encadré ci-dessous.</p>"
        f"<div class='excerpt'>{html.escape(clean)}</div>"
        "<hr style='border:0;border-top:1px solid #e2e8f0;margin:18px 0;'/>"
        f"<div class='doc'>{safe_full}</div>"
    )
    return body, False


def create_html_preview(path: Optional[Path], full_text: str, excerpt: str, title: str, status: str = "ok") -> HTMLResponse:
    if path and full_text:
        body, found = highlight_html_text(full_text, excerpt)
        label = f"{title} - {path.name}"
        html_page = html_document(label, body, "ok" if found else "warn")
    else:
        body = (
            "<p class='muted'>Le document complet n'a pas été retrouvé. Fallback avec l'extrait cliqué.</p>"
            f"<div class='excerpt'>{html.escape(clean_excerpt(excerpt) or 'Aucun extrait reçu.')}</div>"
        )
        html_page = html_document(title or "Document introuvable", body, "warn")
    return HTMLResponse(html_page)


def build_preview_url(project_id: int, payload: SourceHighlightRequest) -> str:
    params = {
        "excerpt": payload.excerpt,
        "source_path": payload.source_path or "",
        "source_name": payload.source_name or payload.document_name or "",
        "title": payload.title or "",
        "organisme": payload.organisme or "",
        "project_name": payload.project_name or "",
        "year": payload.year or "",
    }
    query = "&".join(f"{k}={quote(str(v))}" for k, v in params.items() if v)
    return f"/projects/{project_id}/source-highlight/preview?{query}"


@router.post("/projects/{project_id}/source-highlight/preview")
async def source_highlight_preview_post(
    project_id: int,
    payload: SourceHighlightRequest = Body(...),
):
    """Retourne soit un JSON avec preview_url, soit directement la prévisualisation."""
    if payload.return_json:
        return JSONResponse({"ok": True, "preview_url": build_preview_url(project_id, payload)})

    return await render_source_highlight(project_id, payload)


@router.get("/projects/{project_id}/source-highlight/preview")
async def source_highlight_preview_get(
    project_id: int,
    request: Request,
    excerpt: str = Query("", description="Extrait cliqué dans le frontend"),
    source_path: Optional[str] = Query(None),
    source_name: Optional[str] = Query(None),
    document_name: Optional[str] = Query(None),
    title: Optional[str] = Query(None),
    role: Optional[str] = Query(None),
    organisme: Optional[str] = Query(None),
    project_name: Optional[str] = Query(None),
    year: Optional[str] = Query(None),
):
    payload = SourceHighlightRequest(
        excerpt=excerpt,
        source_path=source_path,
        source_name=source_name,
        document_name=document_name,
        title=title,
        role=role,
        organisme=organisme,
        project_name=project_name,
        year=year,
    )
    return await render_source_highlight(project_id, payload)


async def render_source_highlight(project_id: int, payload: SourceHighlightRequest):
    excerpt = clean_excerpt(payload.excerpt)
    if not excerpt:
        return create_html_preview(None, "", "Aucun extrait reçu depuis le frontend.", "Extrait manquant", "warn")

    path, full_text, score = find_best_document(project_id, payload)
    title = payload.title or "Prévisualisation source"

    if not path:
        return create_html_preview(None, "", excerpt, title, "warn")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        highlighted_pdf, exact_found, message = create_highlighted_pdf(path, excerpt, project_id)
        if highlighted_pdf and highlighted_pdf.exists():
            headers = {
                "X-EnnoSmart-Source-File": path.name,
                "X-EnnoSmart-Highlight-Exact": "true" if exact_found else "false",
                "X-EnnoSmart-Message": message[:500],
            }
            return FileResponse(
                highlighted_pdf,
                media_type="application/pdf",
                filename=highlighted_pdf.name,
                headers=headers,
            )

        # Si PyMuPDF manque ou erreur PDF, fallback HTML avec extrait.
        return create_html_preview(path, full_text, excerpt, f"{title} - {path.name}", "warn")

    if suffix in DOCX_EXTENSIONS or suffix in TEXT_EXTENSIONS:
        if not full_text:
            full_text = read_document_text(path)
        return create_html_preview(path, full_text, excerpt, f"{title} - {path.name}", "ok")

    return create_html_preview(path, full_text, excerpt, f"{title} - {path.name}", "warn")


@router.get("/projects/{project_id}/source-highlight/health")
async def source_highlight_health(project_id: int):
    try:
        import fitz  # type: ignore
        pymupdf_ok = True
    except Exception:
        pymupdf_ok = False

    try:
        import docx  # type: ignore
        docx_ok = True
    except Exception:
        docx_ok = False

    return {
        "ok": True,
        "project_id": project_id,
        "project_root": str(PROJECT_ROOT),
        "storage_root_exists": STORAGE_ROOT.exists(),
        "outputs_root_exists": OUTPUTS_ROOT.exists(),
        "preview_root": str(PREVIEW_ROOT),
        "pymupdf_ok": pymupdf_ok,
        "python_docx_ok": docx_ok,
    }


# ---------------------------------------------------------------------------
# Compatibilité frontend V70.1
# Le frontend appelle parfois /projects/{id}/cir-source-view/open-passage.
# Ces endpoints redirigent vers la même logique que /source-highlight/preview.
# ---------------------------------------------------------------------------


def _pick_first(payload: Dict[str, Any], keys: List[str], default: str = "") -> str:
    for key in keys:
        value = payload.get(key)
        if value is not None and str(value).strip():
            return str(value)
    return default


def _payload_to_source_request(payload: Dict[str, Any]) -> SourceHighlightRequest:
    """Accepte plusieurs formats possibles envoyés par le frontend."""
    current = payload.get("current_item") if isinstance(payload.get("current_item"), dict) else {}
    previous = payload.get("previous_candidate") if isinstance(payload.get("previous_candidate"), dict) else {}
    best_match = payload.get("best_match") if isinstance(payload.get("best_match"), dict) else {}
    best_previous = best_match.get("previous_candidate") if isinstance(best_match.get("previous_candidate"), dict) else {}

    excerpt = _pick_first(payload, [
        "excerpt", "passage", "text", "highlight", "selected_text", "source_excerpt",
        "current_excerpt", "previous_excerpt", "content",
    ])
    if not excerpt:
        excerpt = _pick_first(current, ["text", "source_text", "excerpt", "description", "title"])
    if not excerpt:
        excerpt = _pick_first(previous, ["text", "source_text", "excerpt", "description", "title"])
    if not excerpt:
        excerpt = _pick_first(best_previous, ["text", "source_text", "excerpt", "description", "title"])

    source_path = _pick_first(payload, [
        "source_path", "path", "file_path", "document_path", "sourceFilePath", "filePath",
        "current_source_path", "previous_source_path",
    ])
    if not source_path:
        source_path = _pick_first(current, ["source_path", "path", "file_path", "document_path"])
    if not source_path:
        source_path = _pick_first(previous, ["source_path", "path", "file_path", "document_path"])
    if not source_path:
        source_path = _pick_first(best_previous, ["source_path", "path", "file_path", "document_path"])

    source_name = _pick_first(payload, [
        "source_name", "document_name", "filename", "file_name", "name", "source", "document",
        "current_document", "previous_document",
    ])
    if not source_name:
        source_name = _pick_first(current, ["source_name", "document_name", "filename", "file_name", "source", "document"])
    if not source_name:
        source_name = _pick_first(previous, ["source_name", "document_name", "filename", "file_name", "source", "document"])
    if not source_name:
        source_name = _pick_first(best_previous, ["source_name", "document_name", "filename", "file_name", "source", "document"])

    return SourceHighlightRequest(
        excerpt=excerpt,
        source_path=source_path or None,
        source_name=source_name or None,
        document_name=_pick_first(payload, ["document_name", "document", "filename", "file_name"]) or None,
        title=_pick_first(payload, ["title", "card_title", "label"], "Prévisualisation source") or None,
        role=_pick_first(payload, ["role", "type", "category"]) or None,
        organisme=_pick_first(payload, ["organisme", "organization", "client"]) or None,
        project_name=_pick_first(payload, ["project_name", "project", "projectName"]) or None,
        year=_pick_first(payload, ["year", "current_year", "currentYear"]) or None,
        return_json=bool(payload.get("return_json", True)),
    )


@router.post("/projects/{project_id}/cir-source-view/open-passage")
async def cir_source_view_open_passage(
    project_id: int,
    payload: Dict[str, Any] = Body(default_factory=dict),
):
    """
    Endpoint compatible avec le frontend CIR précédent.

    Par défaut, il renvoie un JSON avec une URL GET ouvrable dans un nouvel onglet.
    Si le frontend envoie {"return_file": true}, il renvoie directement le PDF/HTML.
    """
    req = _payload_to_source_request(payload or {})

    if bool(payload.get("return_file", False)) or bool(payload.get("direct", False)):
        return await render_source_highlight(project_id, req)

    preview_url = build_preview_url(project_id, req)
    return JSONResponse({
        "ok": True,
        "project_id": project_id,
        "preview_url": preview_url,
        "open_url": preview_url,
        "url": preview_url,
        "display_url": preview_url,
        "method": "GET",
        "message": "Prévisualisation prête. Ouvre preview_url dans un nouvel onglet.",
    })


@router.get("/projects/{project_id}/cir-source-view/open-passage")
async def cir_source_view_open_passage_get(
    project_id: int,
    excerpt: str = Query(""),
    source_path: Optional[str] = Query(None),
    source_name: Optional[str] = Query(None),
    document_name: Optional[str] = Query(None),
    title: Optional[str] = Query(None),
    organisme: Optional[str] = Query(None),
    project_name: Optional[str] = Query(None),
    year: Optional[str] = Query(None),
):
    req = SourceHighlightRequest(
        excerpt=excerpt,
        source_path=source_path,
        source_name=source_name,
        document_name=document_name,
        title=title,
        organisme=organisme,
        project_name=project_name,
        year=year,
    )
    return await render_source_highlight(project_id, req)


@router.get("/projects/{project_id}/cir-source-view/health")
async def cir_source_view_health(project_id: int):
    return await source_highlight_health(project_id)
