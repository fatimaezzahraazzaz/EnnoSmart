# -*- coding: utf-8 -*-
from __future__ import annotations

"""
Router V71 EnnoSmart — ouverture du document source d'un passage CIR.

Routes ajoutées :
- GET  /projects/{project_id}/cir-source-view/health
- POST /projects/{project_id}/cir-source-view/open-passage
- GET  /projects/{project_id}/cir-source-view/open-passage
- GET  /projects/{project_id}/cir-source-view/file/{asset_name}

Objectif :
- Recevoir un extrait cliqué côté frontend.
- Retrouver le document complet lié à cet extrait.
- Pour PDF : générer une copie PDF avec passage surligné/encadré.
- Pour DOCX/TXT : générer une page HTML avec le passage encadré.
- Si document introuvable : fallback HTML avec l'extrait et les métadonnées reçues.

Dépendances conseillées :
    pip install pymupdf python-docx
"""

import html
import json
import re
import shutil
import unicodedata
import zipfile
from dataclasses import dataclass
from datetime import datetime
from difflib import SequenceMatcher
import os
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple
from urllib.parse import quote

from fastapi import APIRouter, Body, HTTPException, Query, Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, RedirectResponse
from pydantic import BaseModel, Field

try:
    import fitz  # PyMuPDF
    PYMUPDF_OK = True
except Exception:
    fitz = None
    PYMUPDF_OK = False

try:
    import docx  # python-docx
    PYTHON_DOCX_OK = True
except Exception:
    docx = None
    PYTHON_DOCX_OK = False

router = APIRouter(prefix="/projects/{project_id}/cir-source-view", tags=["CIR source view"])

ROOT = Path(os.getenv("ENNOSMART_ROOT") or Path(__file__).resolve().parents[2]).resolve()
STORAGE = ROOT / "storage"
OUTPUTS = ROOT / "outputs"
PREVIEW_DIR = ROOT / "storage" / "source_previews"
ALLOWED_SUFFIXES = {".pdf", ".docx", ".doc", ".txt", ".md", ".html", ".htm"}
MAX_SEARCH_FILES = 2500


class OpenPassagePayload(BaseModel):
    # Champs directs
    excerpt: Optional[str] = None
    passage: Optional[str] = None
    text: Optional[str] = None
    highlight: Optional[str] = None
    source_path: Optional[str] = None
    source_name: Optional[str] = None
    source_file: Optional[str] = None
    document: Optional[str] = None
    filename: Optional[str] = None
    title: Optional[str] = None
    side: str = "current"  # current | previous
    comparison_index: Optional[int] = None
    return_json: bool = True

    # Objets complets envoyés par le frontend
    item: Optional[Dict[str, Any]] = None
    comparison_item: Optional[Dict[str, Any]] = None
    current_item: Optional[Dict[str, Any]] = None
    previous_candidate: Optional[Dict[str, Any]] = None
    best_match: Optional[Dict[str, Any]] = None

    # Contexte facultatif
    organisme: Optional[str] = None
    project_name: Optional[str] = None
    year: Optional[str] = None


@dataclass
class SourceRequest:
    excerpt: str = ""
    source_name: str = ""
    source_path: str = ""
    title: str = "Source documentaire"
    side: str = "current"
    raw: Dict[str, Any] = None


@dataclass
class FoundDocument:
    path: Path
    score: float
    reason: str


def _norm(value: Any) -> str:
    text = str(value or "")
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _safe_filename(value: str, default: str = "preview") -> str:
    value = _norm(value).replace(" ", "_")[:90].strip("_")
    return value or default


def _short(value: Any, limit: int = 900) -> str:
    text = str(value or "").strip()
    if len(text) <= limit:
        return text
    return text[:limit].strip() + "…"


def _deep_get(data: Any, paths: Iterable[str]) -> str:
    for path in paths:
        cur = data
        ok = True
        for part in path.split("."):
            if isinstance(cur, dict) and part in cur:
                cur = cur.get(part)
            else:
                ok = False
                break
        if ok and cur not in (None, ""):
            if isinstance(cur, (dict, list)):
                continue
            return str(cur)
    return ""


def _merge_payload(payload: Optional[OpenPassagePayload], request_json: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    data: Dict[str, Any] = {}
    if payload is not None:
        data.update(payload.model_dump(exclude_none=True))
    if isinstance(request_json, dict):
        data.update({k: v for k, v in request_json.items() if v is not None})
    return data


def _extract_from_any_shape(data: Dict[str, Any]) -> SourceRequest:
    """Accepte plusieurs formes de payload pour éviter le body vide/mauvais mapping frontend."""
    item = data.get("comparison_item") or data.get("item") or data
    current = data.get("current_item") or item.get("current_item") or item.get("current") or {}
    best_match = data.get("best_match") or item.get("best_match") or {}
    previous = data.get("previous_candidate") or best_match.get("previous_candidate") or item.get("previous_candidate") or {}

    side = str(data.get("side") or "current").lower().strip()
    source_obj = previous if side.startswith("prev") else current

    excerpt = (
        data.get("excerpt")
        or data.get("passage")
        or data.get("highlight")
        or data.get("text")
        or source_obj.get("excerpt")
        or source_obj.get("source_excerpt")
        or source_obj.get("text_excerpt")
        or source_obj.get("text")
        or source_obj.get("source_text")
        or source_obj.get("description")
        or ""
    )

    source_name = (
        data.get("source_name")
        or data.get("source_file")
        or data.get("document")
        or data.get("filename")
        or source_obj.get("source_name")
        or source_obj.get("source_file")
        or source_obj.get("source_document")
        or source_obj.get("document")
        or source_obj.get("filename")
        or source_obj.get("file_name")
        or source_obj.get("name")
        or ""
    )

    source_path = (
        data.get("source_path")
        or source_obj.get("source_path")
        or source_obj.get("path")
        or source_obj.get("file_path")
        or source_obj.get("absolute_path")
        or ""
    )

    title = (
        data.get("title")
        or item.get("title")
        or current.get("title")
        or source_name
        or "Source documentaire"
    )

    return SourceRequest(
        excerpt=str(excerpt or "").strip(),
        source_name=str(source_name or "").strip(),
        source_path=str(source_path or "").strip(),
        title=str(title or "Source documentaire").strip(),
        side=side,
        raw=data,
    )


def _candidate_roots(project_id: int, data: Dict[str, Any]) -> List[Path]:
    roots = [STORAGE, OUTPUTS]

    organisme = _safe_filename(str(data.get("organisme") or ""), "")
    project_name = _safe_filename(str(data.get("project_name") or ""), "")
    year = str(data.get("year") or "").strip()

    if organisme and project_name and year:
        roots.insert(0, STORAGE / "organismes" / organisme / "projects" / project_name / "years" / year)
        roots.insert(1, OUTPUTS / "safe_rag_upload" / organisme / project_name / year)

    roots.insert(0, STORAGE / "projects" / str(project_id))
    roots.insert(1, STORAGE / "projects" / str(project_id) / "documents")
    roots.insert(2, STORAGE / "projects" / str(project_id) / "raw")

    unique: List[Path] = []
    seen = set()
    for root in roots:
        try:
            key = str(root.resolve()).lower()
        except Exception:
            key = str(root).lower()
        if key not in seen:
            seen.add(key)
            unique.append(root)
    return unique


def _iter_documents(project_id: int, data: Dict[str, Any]) -> Iterable[Path]:
    count = 0
    for root in _candidate_roots(project_id, data):
        if not root.exists():
            continue
        for path in root.rglob("*"):
            if count >= MAX_SEARCH_FILES:
                return
            if path.is_file() and path.suffix.lower() in ALLOWED_SUFFIXES:
                count += 1
                yield path


def _resolve_path(source_path: str) -> Optional[Path]:
    if not source_path:
        return None
    p = Path(source_path)
    if p.exists() and p.is_file():
        return p

    # Sécurité : si chemin relatif, le chercher sous ROOT.
    rel = ROOT / source_path
    if rel.exists() and rel.is_file():
        return rel
    return None


def _read_docx_text(path: Path) -> str:
    if PYTHON_DOCX_OK:
        try:
            d = docx.Document(str(path))
            return "\n".join(p.text for p in d.paragraphs if p.text).strip()
        except Exception:
            pass

    # Fallback XML minimal.
    try:
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        xml = re.sub(r"</w:p>", "\n", xml)
        xml = re.sub(r"<[^>]+>", "", xml)
        return html.unescape(xml).strip()
    except Exception:
        return ""


def _read_text_file(path: Path, limit: int = 2_000_000) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _read_docx_text(path)
    if suffix in {".txt", ".md", ".html", ".htm"}:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")[:limit]
        except Exception:
            try:
                return path.read_text(encoding="latin-1", errors="ignore")[:limit]
            except Exception:
                return ""
    return ""


def _pdf_text(path: Path, max_pages: int = 80) -> str:
    if not PYMUPDF_OK:
        return ""
    try:
        doc = fitz.open(str(path))
        chunks = []
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            chunks.append(page.get_text("text") or "")
        doc.close()
        return "\n".join(chunks)
    except Exception:
        return ""


def _score_file(path: Path, source_name: str, excerpt: str) -> Tuple[float, str]:
    score = 0.0
    reasons: List[str] = []
    n_name = _norm(source_name)
    n_file = _norm(path.name)
    n_excerpt = _norm(excerpt)

    if n_name:
        if n_name == n_file:
            score += 1.0
            reasons.append("nom exact")
        elif n_name in n_file or n_file in n_name:
            score += 0.75
            reasons.append("nom proche")
        else:
            ratio = SequenceMatcher(None, n_name, n_file).ratio()
            if ratio > 0.55:
                score += ratio * 0.55
                reasons.append(f"nom similaire {ratio:.2f}")

    if n_excerpt and len(n_excerpt) > 20:
        content = _pdf_text(path) if path.suffix.lower() == ".pdf" else _read_text_file(path)
        n_content = _norm(content[:200000])
        if n_excerpt and n_excerpt[:150] in n_content:
            score += 1.5
            reasons.append("extrait retrouvé")
        else:
            words = [w for w in n_excerpt.split() if len(w) > 4][:35]
            if words:
                hits = sum(1 for w in words if w in n_content)
                ratio = hits / max(len(words), 1)
                if ratio > 0.15:
                    score += ratio
                    reasons.append(f"mots extrait {ratio:.2f}")

    return score, ", ".join(reasons) or "score faible"


def _find_document(project_id: int, req: SourceRequest) -> Optional[FoundDocument]:
    direct = _resolve_path(req.source_path)
    if direct:
        return FoundDocument(direct, 9.0, "source_path direct")

    best: Optional[FoundDocument] = None
    for path in _iter_documents(project_id, req.raw or {}):
        score, reason = _score_file(path, req.source_name, req.excerpt)
        if best is None or score > best.score:
            best = FoundDocument(path, score, reason)

    if best and best.score >= 0.35:
        return best
    return None


def _find_text_window(full_text: str, excerpt: str, radius: int = 1400) -> Tuple[str, bool]:
    if not full_text:
        return "", False
    if not excerpt:
        return _short(full_text, radius * 2), False

    n_full = _norm(full_text)
    n_excerpt = _norm(excerpt)

    # Chercher une phrase courte représentative plutôt que tout l'extrait long.
    candidates = [n_excerpt[:300], n_excerpt[:180], n_excerpt[:120]]
    candidates += [c for c in re.split(r"[.!?;:\n]+", n_excerpt) if len(c) > 35][:8]

    for cand in candidates:
        cand = cand.strip()
        if len(cand) < 20:
            continue
        pos_norm = n_full.find(cand)
        if pos_norm >= 0:
            # approximation : retrouver par mots dans le texte original
            words = cand.split()[:6]
            pattern = r"\b" + r"\W+".join(map(re.escape, words))
            m = re.search(pattern, full_text, flags=re.IGNORECASE)
            if m:
                start = max(0, m.start() - radius)
                end = min(len(full_text), m.end() + radius)
                return full_text[start:end], True

    # fallback par mots fréquents
    words = [w for w in n_excerpt.split() if len(w) > 5][:12]
    positions = []
    n_full_words = n_full
    for w in words:
        p = n_full_words.find(w)
        if p >= 0:
            positions.append(p)
    if positions:
        # approx sur texte original par premier mot.
        first = words[0]
        m = re.search(re.escape(first), full_text, flags=re.IGNORECASE)
        if m:
            start = max(0, m.start() - radius)
            end = min(len(full_text), m.end() + radius)
            return full_text[start:end], False

    return _short(full_text, radius * 2), False


def _highlight_html(text: str, excerpt: str) -> str:
    safe = html.escape(text or "")
    if not excerpt:
        return f"<pre>{safe}</pre>"

    # Encadrer la fenêtre entière, plus fiable pour DOCX/TXT.
    return f'<div class="highlight-block"><pre>{safe}</pre></div>'


def _write_html_preview(req: SourceRequest, found: Optional[FoundDocument], full_text: str, matched: bool, fallback_reason: str = "") -> Path:
    PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    name = _safe_filename(req.source_name or req.title or "source")
    out = PREVIEW_DIR / f"{stamp}_{name}.html"

    window, window_matched = _find_text_window(full_text, req.excerpt)
    matched = matched or window_matched

    doc_label = found.path.name if found else (req.source_name or "Document non retrouvé")
    doc_path = str(found.path) if found else (req.source_path or "—")
    badge = "Passage retrouvé" if matched else "Fallback"
    badge_class = "ok" if matched else "warn"

    body = f"""<!doctype html>
<html lang="fr">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{html.escape(req.title or 'Prévisualisation source')}</title>
  <style>
    body {{ font-family: Arial, sans-serif; margin: 0; background: #f8fafc; color: #0f172a; }}
    .page {{ max-width: 1100px; margin: 24px auto; padding: 0 16px; }}
    .card {{ background: white; border: 1px solid #e2e8f0; border-radius: 14px; padding: 18px; box-shadow: 0 1px 2px rgba(15,23,42,0.06); }}
    .head {{ display: flex; gap: 10px; align-items: flex-start; justify-content: space-between; flex-wrap: wrap; }}
    h1 {{ font-size: 18px; margin: 0 0 8px; }}
    .meta {{ font-size: 13px; color: #475569; line-height: 1.6; }}
    .badge {{ display: inline-block; border-radius: 999px; padding: 5px 10px; font-size: 12px; font-weight: 700; border: 1px solid; }}
    .ok {{ background: #ecfdf5; color: #047857; border-color: #a7f3d0; }}
    .warn {{ background: #fffbeb; color: #b45309; border-color: #fde68a; }}
    .highlight-block {{ border: 3px solid #f97316; background: #fff7ed; border-radius: 12px; padding: 14px; margin-top: 16px; }}
    pre {{ white-space: pre-wrap; word-break: break-word; font-family: Arial, sans-serif; line-height: 1.7; font-size: 14px; margin: 0; }}
    .excerpt {{ margin-top: 16px; background: #f1f5f9; border: 1px dashed #cbd5e1; border-radius: 12px; padding: 14px; }}
    .label {{ font-size: 12px; text-transform: uppercase; letter-spacing: .06em; font-weight: 700; color: #64748b; margin-bottom: 8px; }}
  </style>
</head>
<body>
  <div class="page">
    <div class="card">
      <div class="head">
        <div>
          <h1>{html.escape(req.title or 'Prévisualisation source')}</h1>
          <div class="meta">
            <strong>Fichier source :</strong> {html.escape(doc_label)}<br />
            <strong>Chemin :</strong> {html.escape(doc_path)}<br />
            <strong>Mode :</strong> {html.escape(req.side)}
          </div>
        </div>
        <span class="badge {badge_class}">{badge}</span>
      </div>

      {('<p class="meta"><strong>Note :</strong> ' + html.escape(fallback_reason) + '</p>') if fallback_reason else ''}

      <div class="excerpt">
        <div class="label">Extrait cliqué reçu depuis le frontend</div>
        <pre>{html.escape(req.excerpt or 'Aucun extrait reçu depuis le frontend.')}</pre>
      </div>

      <div class="label" style="margin-top:18px">Document complet / fenêtre autour du passage</div>
      {_highlight_html(window or req.excerpt or 'Aucun contenu exploitable.', req.excerpt)}
    </div>
  </div>
</body>
</html>"""
    out.write_text(body, encoding="utf-8")
    return out


def _annotate_pdf(path: Path, req: SourceRequest) -> Tuple[Optional[Path], bool]:
    if not PYMUPDF_OK:
        return None, False

    PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    out = PREVIEW_DIR / f"{stamp}_{_safe_filename(path.stem)}_highlight.pdf"
    shutil.copy2(path, out)

    matched = False
    try:
        doc = fitz.open(str(out))
        query_parts = []
        raw = req.excerpt or ""
        query_parts.append(raw[:220])
        query_parts.extend([p.strip() for p in re.split(r"[.!?;:\n]+", raw) if len(p.strip()) > 35][:6])

        for page in doc:
            found_rects = []
            for query in query_parts:
                if not query:
                    continue
                rects = page.search_for(query, quads=False)
                if rects:
                    found_rects.extend(rects[:5])
                    break

            if found_rects:
                matched = True
                for rect in found_rects[:8]:
                    try:
                        annot = page.add_highlight_annot(rect)
                        annot.update()
                    except Exception:
                        pass
                    try:
                        border = fitz.Rect(rect.x0 - 2, rect.y0 - 2, rect.x1 + 2, rect.y1 + 2)
                        shape = page.new_shape()
                        shape.draw_rect(border)
                        shape.finish(color=(1, 0.35, 0), width=1.4)
                        shape.commit()
                    except Exception:
                        pass
                break

        doc.saveIncr()
        doc.close()
        return out, matched
    except Exception:
        return None, False


def _asset_url(project_id: int, path: Path) -> str:
    return f"/projects/{project_id}/cir-source-view/file/{quote(path.name)}"


@router.get("/health")
def health(project_id: int):
    return {
        "ok": True,
        "router": "cir_source_view_v71",
        "project_id": project_id,
        "pymupdf_ok": PYMUPDF_OK,
        "python_docx_ok": PYTHON_DOCX_OK,
        "preview_dir": str(PREVIEW_DIR),
    }


@router.get("/file/{asset_name}")
def get_preview_file(project_id: int, asset_name: str):
    path = PREVIEW_DIR / asset_name
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=404, detail="Prévisualisation introuvable.")
    return FileResponse(str(path))


@router.get("/open-passage")
def open_passage_get(
    project_id: int,
    excerpt: str = Query(""),
    source_path: str = Query(""),
    source_name: str = Query(""),
    title: str = Query("Source documentaire"),
    side: str = Query("current"),
    organisme: str = Query(""),
    project_name: str = Query(""),
    year: str = Query(""),
):
    data = {
        "excerpt": excerpt,
        "source_path": source_path,
        "source_name": source_name,
        "title": title,
        "side": side,
        "organisme": organisme,
        "project_name": project_name,
        "year": year,
    }
    return _open_passage_impl(project_id, data, return_json=False)


@router.post("/open-passage")
async def open_passage_post(
    project_id: int,
    request: Request,
    payload: Optional[OpenPassagePayload] = Body(default=None),
):
    request_json: Optional[Dict[str, Any]] = None
    try:
        request_json = await request.json()
    except Exception:
        request_json = None

    data = _merge_payload(payload, request_json)
    return_json = bool(data.get("return_json", True))
    return _open_passage_impl(project_id, data, return_json=return_json)


def _open_passage_impl(project_id: int, data: Dict[str, Any], return_json: bool):
    req = _extract_from_any_shape(data)
    found = _find_document(project_id, req)

    # Important : si le frontend n'a rien envoyé, on renvoie une page explicite + JSON utile.
    if not req.excerpt:
        out = _write_html_preview(
            req,
            found=None,
            full_text="",
            matched=False,
            fallback_reason="Aucun extrait reçu depuis le frontend. Envoie excerpt/passage/text ou current_item.text.",
        )
        url = _asset_url(project_id, out)
        result = {
            "ok": False,
            "reason": "missing_excerpt",
            "message": "Aucun extrait reçu depuis le frontend. Envoie excerpt/passage/text ou current_item.text.",
            "source_name_received": req.source_name,
            "source_path_received": req.source_path,
            "preview_url": url,
            "open_url": url,
            "url": url,
            "display_url": url,
        }
        return JSONResponse(result, status_code=200) if return_json else RedirectResponse(url)

    if found and found.path.suffix.lower() == ".pdf":
        pdf_out, matched = _annotate_pdf(found.path, req)
        if pdf_out:
            url = _asset_url(project_id, pdf_out)
            result = {
                "ok": True,
                "mode": "pdf",
                "matched": matched,
                "source_name": found.path.name,
                "source_path": str(found.path),
                "match_reason": found.reason,
                "preview_url": url,
                "open_url": url,
                "url": url,
                "display_url": url,
            }
            return JSONResponse(result) if return_json else RedirectResponse(url)

    if found:
        full_text = _read_text_file(found.path)
        if not full_text and found.path.suffix.lower() == ".pdf":
            full_text = _pdf_text(found.path)
        out = _write_html_preview(req, found, full_text, matched=False)
        url = _asset_url(project_id, out)
        result = {
            "ok": True,
            "mode": "html",
            "matched": bool(full_text),
            "source_name": found.path.name,
            "source_path": str(found.path),
            "match_reason": found.reason,
            "preview_url": url,
            "open_url": url,
            "url": url,
            "display_url": url,
        }
        return JSONResponse(result) if return_json else RedirectResponse(url)

    # Fallback avec extrait, même si le document n'est pas retrouvé.
    out = _write_html_preview(
        req,
        found=None,
        full_text=req.excerpt,
        matched=False,
        fallback_reason="Le document complet n'a pas été retrouvé. Fallback avec l'extrait cliqué.",
    )
    url = _asset_url(project_id, out)
    result = {
        "ok": False,
        "reason": "document_not_found",
        "message": "Le document complet n'a pas été retrouvé. Fallback avec l'extrait cliqué.",
        "source_name_received": req.source_name,
        "source_path_received": req.source_path,
        "excerpt_received": _short(req.excerpt, 300),
        "preview_url": url,
        "open_url": url,
        "url": url,
        "display_url": url,
    }
    return JSONResponse(result, status_code=200) if return_json else RedirectResponse(url)
