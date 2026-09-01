from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Callable


DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)

_VISUAL_LINE_RE = re.compile(
    r"^!\[(?P<alt>.*)\]\(ennoscholar-visual://"
    r"(?P<visual_id>[A-Za-z0-9_-]+)\)$"
)
_ORDERED_ITEM_RE = re.compile(r"^\d+[.)]\s+(?P<text>.+)$")
_INLINE_RE = re.compile(
    r"(\*\*[^*\n]+\*\*|`[^`\n]+`|\[[^\]\n]+\]\(https?://[^)\n]+\))"
)
_INVALID_XML_RE = re.compile(
    "[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD]"
)


class StateOfArtDocxExportError(RuntimeError):
    """Raised when the visible preview cannot be faithfully exported."""


def _clean_text(value: object) -> str:
    return _INVALID_XML_RE.sub(" ", str(value or "")).strip()


def _append_inline_markdown(paragraph: object, text: str) -> None:
    cursor = 0
    for match in _INLINE_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            link = re.match(r"^\[([^]]+)\]\((https?://[^)]+)\)$", token)
            if link:
                paragraph.add_run(f"{link.group(1)} ({link.group(2)})")
            else:
                paragraph.add_run(token)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _picture_source(path: Path) -> str | io.BytesIO:
    if path.suffix.casefold() != ".webp":
        return str(path)
    try:
        from PIL import Image

        output = io.BytesIO()
        with Image.open(path) as image:
            image.convert("RGB").save(output, format="PNG")
        output.seek(0)
        return output
    except Exception as exc:
        raise StateOfArtDocxExportError(
            f"La figure {path.name} ne peut pas être convertie pour Word."
        ) from exc


def build_state_of_art_docx(
    *,
    markdown: str,
    title: str,
    resolve_visual: Callable[[str], Path | None],
) -> tuple[bytes, dict[str, int]]:
    text = _clean_text(markdown)
    if not text:
        raise ValueError("L’état de l’art est vide.")

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx est requis pour exporter l’état de l’art."
        ) from exc

    document = Document()
    document.core_properties.title = _clean_text(title) or "État de l’art"
    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    has_main_title = any(
        _clean_text(line).startswith("# ") for line in text.splitlines()
    )
    embedded_visuals = 0
    paragraphs = 0

    if not has_main_title:
        document.add_heading(_clean_text(title) or "État de l’art", level=0)

    for raw_line in text.splitlines():
        line = _clean_text(raw_line)
        if not line:
            continue

        visual_match = _VISUAL_LINE_RE.match(line)
        if visual_match:
            visual_id = visual_match.group("visual_id")
            visual_path = resolve_visual(visual_id)
            if visual_path is None or not visual_path.is_file():
                raise StateOfArtDocxExportError(
                    "Une figure affichée dans l’aperçu est introuvable "
                    f"pour l’export Word ({visual_id})."
                )
            try:
                picture_paragraph = document.add_paragraph()
                picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                picture_paragraph.add_run().add_picture(
                    _picture_source(visual_path),
                    width=Inches(6.15),
                )
            except StateOfArtDocxExportError:
                raise
            except Exception as exc:
                raise StateOfArtDocxExportError(
                    "Une figure affichée dans l’aperçu n’a pas pu être "
                    f"intégrée au document Word ({visual_id})."
                ) from exc
            embedded_visuals += 1
            continue

        if line.startswith("# "):
            paragraph = document.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            paragraph = document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            paragraph = document.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            paragraph = document.add_heading(line[5:].strip(), level=3)
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            _append_inline_markdown(paragraph, line[2:].strip())
        elif (ordered := _ORDERED_ITEM_RE.match(line)) is not None:
            paragraph = document.add_paragraph(style="List Number")
            _append_inline_markdown(paragraph, ordered.group("text").strip())
        elif line.startswith("*") and line.endswith("*") and len(line) > 2:
            paragraph = document.add_paragraph(style="Caption")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _append_inline_markdown(paragraph, line[1:-1].strip())
        elif line in {"---", "***"}:
            continue
        else:
            paragraph = document.add_paragraph()
            _append_inline_markdown(paragraph, line)
        paragraphs += 1

    output = io.BytesIO()
    document.save(output)
    return output.getvalue(), {
        "embedded_visuals": embedded_visuals,
        "paragraphs": paragraphs,
    }
